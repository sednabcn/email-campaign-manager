#!/usr/bin/env python3
# requirements: python-docx
import sys
import csv
from docx import Document
import re
from pathlib import Path

def docx_to_csv(docx_path, csv_path):
    doc = Document(docx_path)
    # If the doc contains tables, use first table
    if doc.tables:
        table = doc.tables[0]
        # assume first row is header
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for r in table.rows[1:]:
            rows.append([cell.text.strip() for cell in r.cells])
        # write to csv
        with open(csv_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(rows)
        print(f"Wrote {csv_path} from table in {docx_path}")
        return True
    else:
        # fallback: try extract lines with emails using regex
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        pattern = re.compile(r'([A-Za-z ,.-]+)\s+([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})')
        rows = []
        for line in lines:
            m = pattern.search(line)
            if m:
                name = m.group(1).strip()
                email = m.group(2).strip()
                rows.append([name, "", email, "", ""])
        if rows:
            with open(csv_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(["Name","Responsibility","Email","Phone","InstitutionAddress"])
                writer.writerows(rows)
            print(f"Wrote {csv_path} (heuristic) from {docx_path}")
            return True
    print(f"No structured data found in {docx_path}")
    return False

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: convert_docx_contacts.py <input.docx> <output.csv>")
        sys.exit(1)
    docx_to_csv(sys.argv[1], sys.argv[2])

### python scripts/convert_docx_contacts.py ~/path/to/EduAdultList.docx
### contacts/adults/EduAdultList.csv
