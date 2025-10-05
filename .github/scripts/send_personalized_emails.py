#!/usr/bin/env python3
import os
import csv
import time
import argparse
from docx import Document
from email.message import EmailMessage
import smtplib
from pathlib import Path

def docx_to_text(path):
    doc = Document(path)
    parts = []
    
    # Process all paragraphs, including empty ones for spacing
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
        else:
            # Empty paragraph = intentional blank line
            # Add empty string to preserve spacing
            if parts:  # Only if we already have content
                parts.append("")
    
    # Include tables content (concatenate rows)
    for t in doc.tables:
        for row in t.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    
    # Join with single newlines (empty strings will create blank lines)
    return "\n".join(parts)
    
def replace_placeholders(template_text, mapping):
    out = template_text
    for k, v in mapping.items():
        out = out.replace(f"[{k}]", v or "")
    return out

def send_email(smtp_server, smtp_port, smtp_user, smtp_pass, mail_from, mail_to, subject, body, dry_run=False):
    msg = EmailMessage()
    msg["From"] = mail_from
    msg["To"] = mail_to
    msg["Subject"] = subject
    msg.set_content(body)
    
    if dry_run:
        fname = f"dryrun_{mail_to.replace('@','_at_')}.txt"
        Path("outbox").mkdir(exist_ok=True)
        with open(Path("outbox")/fname, "w", encoding="utf-8") as f:
            f.write("-----HEADERS-----\n")
            f.write(f"From: {mail_from}\nTo: {mail_to}\nSubject: {subject}\n\n")
            f.write(body)
        print(f"DRY RUN: wrote {fname}")
        return True
    
    # real send
    try:
        with smtplib.SMTP(smtp_server, int(smtp_port), timeout=30) as s:
            s.ehlo()
            s.starttls()
            s.login(smtp_user, smtp_pass)
            s.send_message(msg)
        print(f"✓ Sent to {mail_to}")
        return True
    except smtplib.SMTPAuthenticationError as e:
        print(f"✗ SMTP Authentication failed: {e}")
        raise
    except smtplib.SMTPException as e:
        print(f"✗ SMTP error: {e}")
        raise
    except Exception as e:
        print(f"✗ Unexpected error: {e}")
        raise

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--contacts", required=True, help="CSV contacts file")
    parser.add_argument("--template", required=True, help="Docx template path")
    parser.add_argument("--subject", required=True, help="Email subject")
    parser.add_argument("--dry-run", action="store_true", help="Do not send, write files")
    parser.add_argument("--delay", type=float, default=1.0, help="Seconds between sends")
    args = parser.parse_args()
    
    # Get SMTP configuration from environment
    smtp_server = os.environ.get("SMTP_SERVER")
    smtp_port = os.environ.get("SMTP_PORT", "587")
    smtp_user = os.environ.get("SMTP_USERNAME")
    smtp_pass = os.environ.get("SMTP_PASSWORD")
    mail_from = os.environ.get("MAIL_FROM")
    
    # Validate configuration
    if not all([smtp_server, smtp_port, smtp_user, smtp_pass, mail_from]):
        print("ERROR: Missing SMTP environment variables")
        print(f"SMTP_SERVER: {'✓' if smtp_server else '✗'}")
        print(f"SMTP_PORT: {'✓' if smtp_port else '✗'}")
        print(f"SMTP_USERNAME: {'✓' if smtp_user else '✗'}")
        print(f"SMTP_PASSWORD: {'✓' if smtp_pass else '✗'}")
        print(f"MAIL_FROM: {'✓' if mail_from else '✗'}")
        raise SystemExit("Set SMTP_SERVER, SMTP_PORT, SMTP_USERNAME, SMTP_PASSWORD, MAIL_FROM")
    
    print(f"SMTP Configuration:")
    print(f"  Server: {smtp_server}:{smtp_port}")
    print(f"  Server length: {len(smtp_server)} chars")
    print(f"  Server has spaces: {' ' in smtp_server}")
    print(f"  User: {smtp_user}")
    print(f"  From: {mail_from}")
    print(f"  Mode: {'DRY RUN' if args.dry_run else 'LIVE'}")
    print()
    
    # Load template
    print(f"Loading template: {args.template}")
    template_text = docx_to_text(args.template)
    print(f"Template loaded ({len(template_text)} characters)")
    print()
    
    # Process contacts
    print(f"Loading contacts: {args.contacts}")
    with open(args.contacts, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        sent = 0
        failed = 0
        
        # Convert to list to count total and iterate
        contacts = list(reader)
        total_contacts = len(contacts)
        print(f"Found {total_contacts} contacts to process")
        print()
        
        if total_contacts == 0:
            print("WARNING: No contacts found in CSV file!")
            print("Make sure the CSV has data rows beyond the header")
        
        for i, row in enumerate(contacts, 1):
            # Make case-insensitive by converting all keys to lowercase
            row_lower = {k.lower().strip(): v for k, v in row.items()}
            
            # Build mapping with flexible key matching
            mapping = {
      
                "Recipient Name": row_lower.get("name", ""),
                "Name": row_lower.get("name", ""),

                "Title / Position": row_lower.get("rank/title", 
                                    row_lower.get("position", "")),

                "Responsibility": row_lower.get("responsibility", 
                                               row_lower.get("rank/title", 
                                               row_lower.get("position", ""))),
                "Department / School": row_lower.get("position", 
                                                    row_lower.get("department", "")),
                "University / Institution": row_lower.get("organization", ""),
                "Address or Campus, if needed": row_lower.get("address", ""),
                "Email": row_lower.get("email", ""),
                "Phone": row_lower.get("phone", ""),
                "InstitutionAddress": row_lower.get("institutionaddress", 
                                                   row_lower.get("address", 
                                                   row_lower.get("organization", ""))),
                "BodyText": row_lower.get("bodytext", "")
            }
            
            to_addr = mapping["Email"].strip()
            if not to_addr or '@' not in to_addr:
                print(f"[{i}] ⚠️  SKIPPING - No valid email for: {mapping.get('Name', 'Unknown')}")
                failed += 1
                continue
            
            print(f"[{i}] Preparing email to {to_addr} ({mapping.get('Recipient Name', 'Unknown')})")
            
            # Replace placeholders
            body = replace_placeholders(template_text, mapping)
            
            print(f"[{i}] Preparing email to {to_addr} ({mapping['Name']})")
            
            try:
                success = send_email(smtp_server, smtp_port, smtp_user, smtp_pass, 
                                    mail_from, to_addr, args.subject, body, 
                                    dry_run=args.dry_run)
                if success:
                    sent += 1
                else:
                    failed += 1
            except Exception as e:
                print(f"ERROR sending to {to_addr}: {e}")
                failed += 1
            
            # Rate limiting between sends (not after last one)
            if i < total_contacts:
                time.sleep(args.delay)
        
        print()
        print("="*50)
        print(f"SUMMARY: {sent} sent, {failed} failed")
        print("="*50)

if __name__ == "__main__":
    main()
