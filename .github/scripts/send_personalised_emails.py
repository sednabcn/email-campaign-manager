#!/usr/bin/env python3
import os
import csv
import time
import argparse
from docx import Document
from email.message import EmailMessage
import smtplib
from pathlib import Path

def docx_to_text(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # also include tables content (concatenate rows)
    for t in doc.tables:
        for row in t.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)

def replace_placeholders(template_text, mapping):
    out = template_text
    for k, v in mapping.items():
        out = out.replace(f"[{k}]", v or "")
    return out

def send_email(smtp_server, smtp_port, smtp_user, smtp_pass, mail_from, mail_to, subject, body, dry_run=False):
    msg = EmailMessage()
    msg["From"] = mail_from
    msg["To"] = mail_to
    msg["Subject"] = subject
    msg.set_content(body)
    if dry_run:
        fname = f"dryrun_{mail_to.replace('@','_at_')}.txt"
        Path("outbox").mkdir(exist_ok=True)
        with open(Path("outbox")/fname, "w", encoding="utf-8") as f:
            f.write("-----HEADERS-----\n")
            f.write(f"From: {mail_from}\nTo: {mail_to}\nSubject: {subject}\n\n")
            f.write(body)
        print(f"DRY RUN: wrote {fname}")
        return True
    # real send
    with smtplib.SMTP(smtp_server, int(smtp_port)) as s:
        s.ehlo()
        s.starttls()
        s.login(smtp_user, smtp_pass)
        s.send_message(msg)
    return True

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--contacts", required=True, help="CSV contacts file")
    parser.add_argument("--template", required=True, help="Docx template path")
    parser.add_argument("--subject", required=True, help="Email subject")
    parser.add_argument("--dry-run", action="store_true", help="Do not send, write files")
    parser.add_argument("--delay", type=float, default=1.0, help="Seconds between sends")
    args = parser.parse_args()

    smtp_server = os.environ.get("SMTP_SERVER")
    smtp_port = os.environ.get("SMTP_PORT","587")
    smtp_user = os.environ.get("SMTP_USERNAME")
    smtp_pass = os.environ.get("SMTP_PASSWORD")
    mail_from = os.environ.get("MAIL_FROM")

    if not all([smtp_server, smtp_port, smtp_user, smtp_pass, mail_from]):
        raise SystemExit("Missing SMTP env vars. Set SMTP_SERVER, SMTP_PORT, SMTP_USERNAME, SMTP_PASSWORD, MAIL_FROM")

    template_text = docx_to_text(args.template)

    with open(args.contacts, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        sent=0
        for row in reader:
            # Normalize keys: ensure CSV header names match placeholders
            mapping = {
                "Name": row.get("Name",""),
                "Responsibility": row.get("Responsibility",""),
                "Email": row.get("Email",""),
                "Phone": row.get("Phone",""),
                "InstitutionAddress": row.get("InstitutionAddress",""),
                "BodyText": row.get("BodyText","")
            }
            body = replace_placeholders(template_text, mapping)
            to_addr = mapping["Email"]
            if not to_addr:
                print("Skipping row with no Email:", row)
                continue
            print(f"Preparing email to {to_addr} ({mapping['Name']})")
            try:
                send_email(smtp_server, smtp_port, smtp_user, smtp_pass, mail_from, to_addr, args.subject, body, dry_run=args.dry_run)
                sent += 1
            except Exception as e:
                print("ERROR sending to", to_addr, ":", e)
            time.sleep(args.delay)
        print(f"Done. Attempts: {sent}")

if __name__ == "__main__":
    main()
