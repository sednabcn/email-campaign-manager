#!/usr/bin/env python3
"""
Enhanced Domain and Template Analysis Script for GitHub Actions
Analyzes campaign templates and sets GitHub Actions outputs
"""

import os
import sys
import json
import re
from pathlib import Path
from datetime import datetime

def extract_template_variables(content):
    """Extract template variables like {{name}}, {{email}} from content"""
    pattern = r'\{\{([^}]+)\}\}'
    matches = re.findall(pattern, content)
    return [match.strip() for match in matches]

def analyze_text_template(file_path):
    """Analyze a text-based template file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        variables = extract_template_variables(content)
        return {
            'file': str(file_path),
            'type': file_path.suffix,
            'variables': variables,
            'variable_count': len(variables),
            'size': file_path.stat().st_size
        }
    except Exception as e:
        print(f"Error analyzing template {file_path}: {e}")
        return None

def analyze_docx_template(file_path):
    """Analyze a DOCX template file"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        content = '\n'.join([para.text for para in doc.paragraphs])
        variables = extract_template_variables(content)
        
        return {
            'file': str(file_path),
            'type': '.docx',
            'variables': variables,
            'variable_count': len(variables),
            'size': file_path.stat().st_size,
            'paragraphs': len(doc.paragraphs)
        }
    except Exception as e:
        print(f"Error analyzing DOCX template {file_path}: {e}")
        return None

def analyze_domains():
    """Main domain and template analysis function"""
    templates_dir = os.environ.get('TEMPLATES_DIR', 'campaign-templates')
    scheduled_dir = os.environ.get('SCHEDULED_DIR', 'scheduled-campaigns')
    target_domain_filter = os.environ.get('TARGET_DOMAIN_FILTER', '')
    
    print(f"Analyzing templates from: {templates_dir}")
    print(f"Analyzing scheduled campaigns from: {scheduled_dir}")
    print(f"Target domain filter: {target_domain_filter or 'None'}")
    
    domain_templates = {}
    scheduled_campaigns = []
    all_template_variables = {}
    template_count = 0
    
    # Analyze campaign-templates directory structure
    templates_path = Path(templates_dir)
    if templates_path.exists():
        print(f"\nAnalyzing enhanced template structure...")
        
        # Expected domain directories
        domains = ['education', 'finance', 'healthcare', 'industry', 'technology', 'government']
        
        for domain in domains:
            domain_path = templates_path / domain
            if domain_path.exists():
                # Count templates by type
                docx_count = len(list(domain_path.glob('*.docx')))
                json_count = len(list(domain_path.glob('*.json')))
                txt_count = len(list(domain_path.glob('*.txt')))
                html_count = len(list(domain_path.glob('*.html')))
                
                total = docx_count + json_count + txt_count + html_count
                
                if target_domain_filter and domain != target_domain_filter:
                    print(f"  - {domain}: {total} templates (FILTERED OUT)")
                    continue
                
                print(f"  - {domain}: {total} templates ({docx_count} DOCX, {json_count} JSON, {txt_count} TXT, {html_count} HTML)")
                
                if total > 0:
                    domain_templates[domain] = {
                        'total': total,
                        'docx': docx_count,
                        'json': json_count,
                        'txt': txt_count,
                        'html': html_count,
                        'templates': []
                    }
                    template_count += total
                    
                    # Analyze individual templates
                    for template_file in domain_path.glob('*'):
                        if template_file.is_file():
                            if template_file.suffix == '.docx':
                                template_info = analyze_docx_template(template_file)
                            elif template_file.suffix in ['.txt', '.html', '.json']:
                                template_info = analyze_text_template(template_file)
                            else:
                                continue
                            
                            if template_info:
                                domain_templates[domain]['templates'].append(template_info)
                                for var in template_info['variables']:
                                    if var not in all_template_variables:
                                        all_template_variables[var] = []
                                    all_template_variables[var].append(str(template_file))
    
    # Analyze scheduled-campaigns directory
    scheduled_path = Path(scheduled_dir)
    if scheduled_path.exists():
        for campaign_file in scheduled_path.glob('*'):
            if campaign_file.is_file():
                if campaign_file.suffix == '.docx':
                    campaign_info = analyze_docx_template(campaign_file)
                elif campaign_file.suffix in ['.txt', '.html', '.json', '.md']:
                    campaign_info = analyze_text_template(campaign_file)
                else:
                    continue
                
                if campaign_info:
                    scheduled_campaigns.append(campaign_info)
                    for var in campaign_info['variables']:
                        if var not in all_template_variables:
                            all_template_variables[var] = []
                        all_template_variables[var].append(str(campaign_file))
    
    scheduled_count = len(scheduled_campaigns)
    
    print(f"Set GITHUB_OUTPUT campaigns={template_count}")
    print(f"Enhanced template analysis completed: {template_count} templates, {scheduled_count} scheduled campaigns")
    
    # Show template variables found
    if all_template_variables:
        print("\nTemplate variables detected in:")
        for var, files in list(all_template_variables.items())[:5]:
            print(f"  - {files[0]}: {[var] + list(all_template_variables.keys())[1:4]}...")
            break
    
    # CRITICAL: Set GitHub Actions output
    github_output = os.environ.get('GITHUB_OUTPUT')
    if github_output:
        with open(github_output, 'a') as f:
            f.write(f'campaigns={template_count}\n')
        print(f"Set GITHUB_OUTPUT campaigns={template_count}")
    else:
        print("WARNING: GITHUB_OUTPUT not set, cannot set output variable")
    
    # Save detailed analysis to JSON
    analysis_data = {
        'template_count': template_count,
        'domain_count': len(domain_templates),
        'scheduled_campaigns': scheduled_count,
        'domain_templates': domain_templates,
        'scheduled_campaign_list': scheduled_campaigns,
        'template_variables_found': all_template_variables,
        'target_domain_filter': target_domain_filter,
        'analysis_timestamp': datetime.now().isoformat(),
        'templates_dir': templates_dir,
        'scheduled_dir': scheduled_dir
    }
    
    with open('domain_analysis.json', 'w') as f:
        json.dump(analysis_data, f, indent=2)
    
    print("Domain analysis completed")
    
    return template_count

if __name__ == '__main__':
    try:
        template_count = analyze_domains()
        sys.exit(0)
    except Exception as e:
        print(f"ERROR in domain analysis: {e}")
        import traceback
        traceback.print_exc()
        
        # Set campaigns to 0 on error
        github_output = os.environ.get('GITHUB_OUTPUT')
        if github_output:
            with open(github_output, 'a') as f:
                f.write('campaigns=0\n')
        
        sys.exit(1)
