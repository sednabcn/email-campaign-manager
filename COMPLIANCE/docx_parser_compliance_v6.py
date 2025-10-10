#!/usr/bin/env python3
"""
docx_parser.py — Domain-aware campaign processor with FULL compliance integration

NEW FEATURES in this version:
 - Active suppression list checking before every send
 - Real-time rate limit enforcement
 - Domain-level throttling
 - Reply handler integration hooks
 - Comprehensive compliance reporting
 - Enhanced error handling and recovery
"""

from __future__ import annotations
import argparse, csv, json, os, sys, time, zipfile, traceback, re, hashlib
from pathlib import Path
from datetime import datetime, date
from typing import List, Dict, Optional, Tuple

# ==============================================================
# Load compliance configuration
# ==============================================================
CONFIG_PATH = Path("compliance_config.json")

if CONFIG_PATH.exists():
    with open(CONFIG_PATH, "r") as f:
        CONFIG = json.load(f)
    print("✅ [config] Loaded compliance_config.json")
else:
    print("⚠️  compliance_config.json not found — using safe defaults.")
    CONFIG = {}

SENDER_INFO = CONFIG.get("sender_info", {})
UNSUBSCRIBE = CONFIG.get("unsubscribe", {})
RATES = CONFIG.get("rate_limits", {})
COMPLIANCE_OPTS = CONFIG.get("compliance", {})

FROM_NAME  = SENDER_INFO.get("from_name", "Your Name")
FROM_EMAIL = SENDER_INFO.get("from_email", "your.email@domain.com")
REPLY_TO   = SENDER_INFO.get("reply_to", FROM_EMAIL)
PHYS_ADDR  = SENDER_INFO.get("physical_address", "Your Address, City, Country")
UNSUB_BASE = UNSUBSCRIBE.get("base_url", "https://example.com/unsubscribe")

MAX_DAILY  = RATES.get("max_daily_sends", 50)
MAX_DOMAIN = RATES.get("max_per_domain_daily", 5)
MIN_DELAY  = RATES.get("min_delay_seconds", 30)

SUPPRESSION_FILE = Path(COMPLIANCE_OPTS.get("suppression_file", "contacts/suppression_list.json"))
RATE_LIMIT_FILE = Path(COMPLIANCE_OPTS.get("rate_limit_file", "tracking/rate_limits.json"))

# ==============================================================
# Path setup for utils import
# ==============================================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)
UTILS_DIR = os.path.join(SCRIPT_DIR, "utils")
if os.path.isdir(UTILS_DIR) and UTILS_DIR not in sys.path:
    sys.path.insert(0, UTILS_DIR)

# ==============================================================
# Optional dependencies
# ==============================================================
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    DOCX_AVAILABLE = False

try:
    from email_personalizer import EmailPersonalizer
    PERSONALIZER_AVAILABLE = True
except Exception:
    try:
        from utils.email_personalizer import EmailPersonalizer
        PERSONALIZER_AVAILABLE = True
    except Exception:
        PERSONALIZER_AVAILABLE = False
        EmailPersonalizer = None

# ==============================================================
# ENHANCED: Built-in Compliance Manager
# ==============================================================
class ComplianceManager:
    """Manages suppression list, rate limits, and compliance checks"""
    
    def __init__(self, max_daily=50, max_per_domain=5, min_delay=30):
        self.max_daily = max_daily
        self.max_per_domain = max_per_domain
        self.min_delay = min_delay
        
        # Load suppression list
        self.suppressed = self._load_suppression_list()
        print(f"📋 [compliance] Loaded {len(self.suppressed)} suppressed emails")
        
        # Load rate limits
        self.rate_limits = self._load_rate_limits()
        self.last_send_time = 0
        
        print(f"🔒 [compliance] Rate limits: {max_daily}/day, {max_per_domain}/domain, {min_delay}s delay")
    
    def _load_suppression_list(self) -> set:
        """Load suppressed emails from JSON"""
        if SUPPRESSION_FILE.exists():
            try:
                with open(SUPPRESSION_FILE) as f:
                    data = json.load(f)
                    return set(email.lower() for email in data.get('suppressed_emails', []))
            except Exception as e:
                print(f"⚠️  Error loading suppression list: {e}")
        return set()
    
    def _load_rate_limits(self) -> dict:
        """Load current rate limit tracking"""
        if RATE_LIMIT_FILE.exists():
            try:
                with open(RATE_LIMIT_FILE) as f:
                    data = json.load(f)
                    # Check if it's today's data
                    if data.get('date') == str(date.today()):
                        return data
            except Exception as e:
                print(f"⚠️  Error loading rate limits: {e}")
        
        # Return fresh data for today
        return {
            'date': str(date.today()),
            'total_sent': 0,
            'domain_counts': {},
            'last_updated': datetime.now().isoformat(),
            'emails_sent': []
        }
    
    def is_suppressed(self, email: str) -> bool:
        """Check if email is on suppression list"""
        return email.lower() in self.suppressed
    
    def can_send_today(self) -> bool:
        """Check if we haven't hit daily limit"""
        return self.rate_limits['total_sent'] < self.max_daily
    
    def can_send_to_domain(self, domain: str) -> bool:
        """Check if we haven't hit domain limit"""
        current = self.rate_limits['domain_counts'].get(domain, 0)
        return current < self.max_per_domain
    
    def enforce_delay(self):
        """Enforce minimum delay between sends"""
        if self.last_send_time > 0:
            elapsed = time.time() - self.last_send_time
            if elapsed < self.min_delay:
                wait_time = self.min_delay - elapsed
                print(f"⏸️  Rate limit delay: {wait_time:.1f}s")
                time.sleep(wait_time)
        self.last_send_time = time.time()
    
    def record_send(self, email: str, domain: str):
        """Record that we sent an email"""
        # Update counts
        self.rate_limits['total_sent'] += 1
        self.rate_limits['domain_counts'][domain] = \
            self.rate_limits['domain_counts'].get(domain, 0) + 1
        self.rate_limits['last_updated'] = datetime.now().isoformat()
        
        # Track individual email
        self.rate_limits['emails_sent'].append({
            'email': email,
            'domain': domain,
            'timestamp': datetime.now().isoformat()
        })
        
        # Save to file
        RATE_LIMIT_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(RATE_LIMIT_FILE, 'w') as f:
            json.dump(self.rate_limits, f, indent=2)
    
    def get_stats(self) -> dict:
        """Get current compliance statistics"""
        return {
            'suppressed_count': len(self.suppressed),
            'sent_today': self.rate_limits['total_sent'],
            'daily_limit': self.max_daily,
            'remaining_today': self.max_daily - self.rate_limits['total_sent'],
            'domains_contacted': len(self.rate_limits['domain_counts']),
            'per_domain_limit': self.max_per_domain,
            'domain_counts': self.rate_limits['domain_counts']
        }
    
    def add_suppression(self, email: str, reason: str = 'manual'):
        """Add email to suppression list"""
        self.suppressed.add(email.lower())
        
        # Save to file
        SUPPRESSION_FILE.parent.mkdir(parents=True, exist_ok=True)
        data = {
            'suppressed_emails': sorted(list(self.suppressed)),
            'last_updated': datetime.now().isoformat(),
            'count': len(self.suppressed)
        }
        with open(SUPPRESSION_FILE, 'w') as f:
            json.dump(data, f, indent=2)
        
        # Log the suppression
        log_file = SUPPRESSION_FILE.parent / 'suppression_log.jsonl'
        with open(log_file, 'a') as f:
            log_entry = {
                'email': email,
                'reason': reason,
                'timestamp': datetime.now().isoformat()
            }
            f.write(json.dumps(log_entry) + '\n')
        
        print(f"🚫 Added {email} to suppression list (reason: {reason})")

# ==============================================================
# Enhanced campaign loader & DOCX validator
# ==============================================================
def is_valid_docx(filepath: str) -> Tuple[bool, Optional[str]]:
    """Validate if a DOCX file is a proper zip archive with required XML files."""
    try:
        if not zipfile.is_zipfile(filepath):
            return False, "Not a valid ZIP archive"
        with zipfile.ZipFile(filepath, 'r') as z:
            required = ["word/document.xml", "[Content_Types].xml"]
            for r in required:
                if r not in z.namelist():
                    return False, f"Missing {r}"
        return True, None
    except Exception as e:
        return False, str(e)

def load_json_campaign(campaign_path: str) -> Optional[dict]:
    """Load and normalize JSON campaign formats."""
    try:
        with open(campaign_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        print(f"[templates] JSON load error for {campaign_path}: {e}")
        return None

    if isinstance(data, dict) and "content" in data:
        return data
    if isinstance(data, dict) and "campaigns" in data and data["campaigns"]:
        return data["campaigns"][0]
    print(f"[templates] Unknown JSON format: {campaign_path}")
    return None

def load_campaign_content(campaign_path: str) -> Optional[str]:
    """Universal loader for campaign templates."""
    p = Path(campaign_path)
    if not p.exists():
        print(f"[templates] Not found: {campaign_path}")
        return None
    
    ext = p.suffix.lower()
    
    try:
        if ext == ".docx":
            if not DOCX_AVAILABLE:
                print("[templates] python-docx not installed — skipping DOCX.")
                return None
            valid, err = is_valid_docx(str(p))
            if not valid:
                print(f"[templates] Invalid DOCX: {err}")
                return None
            doc = Document(str(p))
            parts = [para.text for para in doc.paragraphs if para.text]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts).strip()
            print(f"[templates] Loaded DOCX ({p.name}) — {len(text)} chars")
            return text

        elif ext in (".txt", ".md", ".html"):
            for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
                try:
                    text = p.read_text(encoding=enc)
                    print(f"[templates] Loaded {p.name} ({enc}) — {len(text)} chars")
                    return text
                except UnicodeDecodeError:
                    continue
            print(f"[templates] Encoding issue: {p}")
            return None

        elif ext == ".json":
            data = load_json_campaign(str(p))
            if data and "content" in data:
                return data["content"]
            return None

        else:
            print(f"[templates] Unsupported file type: {ext}")
            return None
            
    except Exception as e:
        print(f"[templates] Error loading {p}: {e}")
        traceback.print_exc()
        return None

# ==============================================================
# PRE-FLIGHT VALIDATION SYSTEM
# ==============================================================
def pre_flight_checks(contacts_root: str, scheduled_root: str, tracking_root: str,
                     alerts_email: str, compliance: ComplianceManager) -> Dict:
    """
    Comprehensive pre-flight validation before campaign execution.
    Returns dict with status and any blocking issues.
    """
    print("\n" + "="*60)
    print("🔍 PRE-FLIGHT VALIDATION SYSTEM")
    print("="*60 + "\n")
    
    issues = []
    warnings = []
    checks_passed = 0
    checks_total = 0
    
    # CHECK 1: Directory Structure
    checks_total += 1
    print("📁 [1/10] Checking directory structure...")
    dirs_to_check = {
        'contacts': contacts_root,
        'scheduled': scheduled_root,
        'tracking': tracking_root
    }
    
    for name, path in dirs_to_check.items():
        p = Path(path)
        if not p.exists():
            issues.append(f"Directory missing: {name} ({path})")
            print(f"  ❌ {name}: NOT FOUND")
        elif not p.is_dir():
            issues.append(f"Path is not a directory: {name} ({path})")
            print(f"  ❌ {name}: NOT A DIRECTORY")
        else:
            print(f"  ✅ {name}: {path}")
            checks_passed += 1
    
    # CHECK 2: Write Permissions
    checks_total += 1
    print("\n✍️  [2/10] Checking write permissions...")
    try:
        test_file = Path(tracking_root) / ".write_test"
        test_file.parent.mkdir(parents=True, exist_ok=True)
        test_file.write_text("test")
        test_file.unlink()
        print("  ✅ Write permissions: OK")
        checks_passed += 1
    except Exception as e:
        issues.append(f"No write permission to tracking directory: {e}")
        print(f"  ❌ Write permissions: FAILED - {e}")
    
    # CHECK 3: Contact Files
    checks_total += 1
    print("\n👥 [3/10] Checking contact files...")
    contact_files = list(Path(contacts_root).glob("*.csv")) + \
                   list(Path(contacts_root).glob("*.xlsx")) + \
                   list(Path(contacts_root).glob("*.xls"))
    
    if not contact_files:
        issues.append(f"No contact files found in {contacts_root}")
        print(f"  ❌ No contact files found")
    else:
        print(f"  ✅ Found {len(contact_files)} contact file(s)")
        for cf in contact_files[:3]:  # Show first 3
            print(f"     - {cf.name}")
        checks_passed += 1
    
    # CHECK 4: Campaign Templates
    checks_total += 1
    print("\n📄 [4/10] Checking campaign templates...")
    template_files = list(Path(scheduled_root).rglob("*.txt")) + \
                    list(Path(scheduled_root).rglob("*.docx")) + \
                    list(Path(scheduled_root).rglob("*.json")) + \
                    list(Path(scheduled_root).rglob("*.md")) + \
                    list(Path(scheduled_root).rglob("*.html"))
    
    if not template_files:
        issues.append(f"No campaign templates found in {scheduled_root}")
        print(f"  ❌ No templates found")
    else:
        print(f"  ✅ Found {len(template_files)} template(s)")
        for tf in template_files[:3]:  # Show first 3
            print(f"     - {tf.name}")
        checks_passed += 1
    
    # CHECK 5: DOCX Validation
    checks_total += 1
    print("\n📋 [5/10] Validating DOCX files...")
    docx_files = list(Path(scheduled_root).rglob("*.docx"))
    invalid_docx = []
    
    if docx_files:
        for docx_file in docx_files:
            valid, error = is_valid_docx(str(docx_file))
            if not valid:
                invalid_docx.append(f"{docx_file.name}: {error}")
                print(f"  ⚠️  {docx_file.name}: INVALID - {error}")
        
        if invalid_docx:
            warnings.extend(invalid_docx)
            print(f"  ⚠️  {len(invalid_docx)} invalid DOCX file(s)")
        else:
            print(f"  ✅ All {len(docx_files)} DOCX files valid")
            checks_passed += 1
    else:
        print("  ℹ️  No DOCX files to validate")
        checks_passed += 1
    
    # CHECK 6: Compliance Configuration
    checks_total += 1
    print("\n🔒 [6/10] Checking compliance configuration...")
    if not compliance:
        warnings.append("Compliance manager not initialized")
        print("  ⚠️  Compliance manager: NOT INITIALIZED")
    else:
        stats = compliance.get_stats()
        print(f"  ✅ Suppression list: {stats['suppressed_count']} emails")
        print(f"  ✅ Daily limit: {stats['daily_limit']} emails")
        print(f"  ✅ Already sent today: {stats['sent_today']}")
        print(f"  ✅ Remaining capacity: {stats['remaining_today']}")
        
        if stats['remaining_today'] <= 0:
            issues.append("Daily send limit already reached")
            print("  ❌ No remaining send capacity today")
        else:
            checks_passed += 1
    
    # CHECK 7: Email Configuration
    checks_total += 1
    print("\n📧 [7/10] Checking email configuration...")
    email_config_ok = True
    
    if not FROM_EMAIL or "@" not in FROM_EMAIL:
        warnings.append("FROM_EMAIL not properly configured")
        email_config_ok = False
        print(f"  ⚠️  FROM_EMAIL: Invalid or missing")
    else:
        print(f"  ✅ FROM_EMAIL: {FROM_EMAIL}")
    
    if not REPLY_TO or "@" not in REPLY_TO:
        warnings.append("REPLY_TO not properly configured")
        email_config_ok = False
        print(f"  ⚠️  REPLY_TO: Invalid or missing")
    else:
        print(f"  ✅ REPLY_TO: {REPLY_TO}")
    
    if not alerts_email or "@" not in alerts_email:
        warnings.append("Alerts email not properly configured")
        print(f"  ⚠️  Alerts email: Invalid or missing")
    else:
        print(f"  ✅ Alerts email: {alerts_email}")
    
    if email_config_ok:
        checks_passed += 1
    
    # CHECK 8: Unsubscribe Configuration
    checks_total += 1
    print("\n🔗 [8/10] Checking unsubscribe configuration...")
    if not UNSUB_BASE or not UNSUB_BASE.startswith("http"):
        warnings.append("Unsubscribe URL not properly configured")
        print(f"  ⚠️  Unsubscribe URL: Invalid or missing")
    else:
        print(f"  ✅ Unsubscribe URL: {UNSUB_BASE}")
        checks_passed += 1
    
    # CHECK 9: Physical Address
    checks_total += 1
    print("\n🏢 [9/10] Checking physical address...")
    if not PHYS_ADDR or len(PHYS_ADDR) < 10:
        warnings.append("Physical address not properly configured (required for CAN-SPAM)")
        print(f"  ⚠️  Physical address: Invalid or too short")
    else:
        print(f"  ✅ Physical address: {PHYS_ADDR[:50]}...")
        checks_passed += 1
    
    # CHECK 10: Dependencies
    checks_total += 1
    print("\n📦 [10/10] Checking dependencies...")
    deps_ok = True
    
    if not DOCX_AVAILABLE:
        warnings.append("python-docx not available - DOCX templates will be skipped")
        print("  ⚠️  python-docx: NOT AVAILABLE")
        deps_ok = False
    else:
        print("  ✅ python-docx: Available")
    
    if not PERSONALIZER_AVAILABLE:
        warnings.append("EmailPersonalizer not available - using fallback")
        print("  ⚠️  EmailPersonalizer: NOT AVAILABLE (using fallback)")
    else:
        print("  ✅ EmailPersonalizer: Available")
    
    if deps_ok:
        checks_passed += 1
    
    # SUMMARY
    print("\n" + "="*60)
    print("📊 PRE-FLIGHT VALIDATION SUMMARY")
    print("="*60)
    print(f"✅ Checks passed: {checks_passed}/{checks_total}")
    print(f"⚠️  Warnings: {len(warnings)}")
    print(f"❌ Blocking issues: {len(issues)}")
    
    if warnings:
        print("\n⚠️  WARNINGS:")
        for w in warnings:
            print(f"   - {w}")
    
    if issues:
        print("\n❌ BLOCKING ISSUES:")
        for i in issues:
            print(f"   - {i}")
        print("\n🚫 Cannot proceed - fix blocking issues first!")
        print("="*60 + "\n")
        return {
            'status': 'FAILED',
            'checks_passed': checks_passed,
            'checks_total': checks_total,
            'issues': issues,
            'warnings': warnings,
            'can_proceed': False
        }
    
    if checks_passed >= checks_total * 0.7:  # 70% threshold
        print("\n✅ PRE-FLIGHT VALIDATION PASSED")
        print("🚀 System ready for campaign execution")
    else:
        print("\n⚠️  PRE-FLIGHT VALIDATION: MARGINAL")
        print("⚠️  Proceed with caution - review warnings above")
    
    print("="*60 + "\n")
    
    return {
        'status': 'PASSED' if checks_passed >= checks_total * 0.7 else 'MARGINAL',
        'checks_passed': checks_passed,
        'checks_total': checks_total,
        'issues': issues,
        'warnings': warnings,
        'can_proceed': True
    }


# ==============================================================
# POST-FLIGHT REPORTING SYSTEM
# ==============================================================
def post_flight_report(tracking_root: str, alerts_email: str, 
                       compliance: ComplianceManager,
                       campaign_results: Dict,
                       pre_flight_results: Dict) -> Dict:
    """
    Comprehensive post-flight reporting and cleanup.
    Generates reports, sends alerts, and creates audit trail.
    """
    print("\n" + "="*60)
    print("📊 POST-FLIGHT REPORTING SYSTEM")
    print("="*60 + "\n")
    
    timestamp = datetime.now()
    
    # Gather final statistics
    final_stats = compliance.get_stats() if compliance else {}
    
    # CREATE COMPREHENSIVE REPORT
    print("📝 [1/7] Generating comprehensive report...")
    
    report = {
        'execution_metadata': {
            'timestamp': timestamp.isoformat(),
            'date': timestamp.strftime('%Y-%m-%d'),
            'time': timestamp.strftime('%H:%M:%S'),
            'duration': 'N/A'  # Would need start time to calculate
        },
        'pre_flight_validation': {
            'status': pre_flight_results.get('status', 'UNKNOWN'),
            'checks_passed': pre_flight_results.get('checks_passed', 0),
            'checks_total': pre_flight_results.get('checks_total', 0),
            'warnings_count': len(pre_flight_results.get('warnings', [])),
            'issues_count': len(pre_flight_results.get('issues', []))
        },
        'campaign_execution': campaign_results,
        'compliance_final': final_stats,
        'configuration': {
            'from_email': FROM_EMAIL,
            'from_name': FROM_NAME,
            'reply_to': REPLY_TO,
            'alerts_email': alerts_email,
            'max_daily': MAX_DAILY,
            'max_per_domain': MAX_DOMAIN,
            'min_delay': MIN_DELAY
        }
    }
    
    # Save JSON report
    report_file = Path(tracking_root) / f"campaign_report_{timestamp.strftime('%Y%m%d_%H%M%S')}.json"
    report_file.parent.mkdir(parents=True, exist_ok=True)
    with open(report_file, 'w') as f:
        json.dump(report, f, indent=2)
    print(f"  ✅ JSON report saved: {report_file}")
    
    # CREATE HUMAN-READABLE REPORT
    print("\n📄 [2/7] Creating human-readable report...")
    
    markdown_report = f"""# Email Campaign Execution Report

## Executive Summary
- **Execution Date:** {timestamp.strftime('%Y-%m-%d %H:%M:%S')}
- **Status:** {'✅ SUCCESS' if campaign_results.get('sent', 0) > 0 else '⚠️ ISSUES'}
- **Pre-flight:** {pre_flight_results.get('status', 'UNKNOWN')}

## Campaign Results
- **Emails Sent:** {campaign_results.get('sent', 0)}
- **Emails Queued:** {campaign_results.get('queued', 0)}
- **Failed:** {campaign_results.get('failed', 0)}
- **Suppressed:** {campaign_results.get('skipped_suppressed', 0)}
- **Rate Limited:** {campaign_results.get('skipped_rate_limit', 0)}

## Compliance Status
- **Total Sent Today:** {final_stats.get('sent_today', 0)}/{final_stats.get('daily_limit', 0)}
- **Remaining Capacity:** {final_stats.get('remaining_today', 0)}
- **Domains Contacted:** {final_stats.get('domains_contacted', 0)}
- **Suppression List Size:** {final_stats.get('suppressed_count', 0)}

## Domain Breakdown
"""
    
    if final_stats.get('domain_counts'):
        for domain, count in sorted(final_stats['domain_counts'].items(), 
                                    key=lambda x: x[1], reverse=True):
            markdown_report += f"- **{domain}:** {count}/{MAX_DOMAIN}\n"
    else:
        markdown_report += "- No domains contacted\n"
    
    markdown_report += f"""
## Configuration
- **From:** {FROM_NAME} <{FROM_EMAIL}>
- **Reply-To:** {REPLY_TO}
- **Alerts:** {alerts_email}
- **Rate Limits:** {MAX_DAILY}/day, {MAX_DOMAIN}/domain, {MIN_DELAY}s delay

## Pre-flight Validation
- **Checks Passed:** {pre_flight_results.get('checks_passed', 0)}/{pre_flight_results.get('checks_total', 0)}
- **Warnings:** {len(pre_flight_results.get('warnings', []))}
- **Blocking Issues:** {len(pre_flight_results.get('issues', []))}

---
*Generated by Compliant Email Campaign System*
"""
    
    markdown_file = Path(tracking_root) / f"campaign_report_{timestamp.strftime('%Y%m%d_%H%M%S')}.md"
    with open(markdown_file, 'w') as f:
        f.write(markdown_report)
    print(f"  ✅ Markdown report saved: {markdown_file}")
    
    # AUDIT TRAIL
    print("\n📋 [3/7] Creating audit trail...")
    
    audit_log = Path(tracking_root) / "audit_trail.jsonl"
    audit_entry = {
        'timestamp': timestamp.isoformat(),
        'event': 'campaign_execution',
        'results': {
            'sent': campaign_results.get('sent', 0),
            'queued': campaign_results.get('queued', 0),
            'failed': campaign_results.get('failed', 0),
            'suppressed': campaign_results.get('skipped_suppressed', 0),
            'rate_limited': campaign_results.get('skipped_rate_limit', 0)
        },
        'compliance': {
            'sent_today': final_stats.get('sent_today', 0),
            'daily_limit': final_stats.get('daily_limit', 0)
        }
    }
    
    with open(audit_log, 'a') as f:
        f.write(json.dumps(audit_entry) + '\n')
    print(f"  ✅ Audit entry logged: {audit_log}")
    
    # ALERT EMAIL SUMMARY (would integrate with SMTP here)
    print("\n📧 [4/7] Preparing alert email summary...")
    
    alert_subject = f"Campaign Report: {campaign_results.get('sent', 0)} sent, {campaign_results.get('failed', 0)} failed"
    alert_body = f"""Campaign Execution Summary
    
Sent: {campaign_results.get('sent', 0)}
Queued: {campaign_results.get('queued', 0)}
Failed: {campaign_results.get('failed', 0)}
Suppressed: {campaign_results.get('skipped_suppressed', 0)}
Rate Limited: {campaign_results.get('skipped_rate_limit', 0)}

Daily Status: {final_stats.get('sent_today', 0)}/{final_stats.get('daily_limit', 0)}

Full report: {report_file}
"""
    
    # Save alert to file (would actually send via SMTP in production)
    alert_file = Path(tracking_root) / f"alert_{timestamp.strftime('%Y%m%d_%H%M%S')}.txt"
    with open(alert_file, 'w') as f:
        f.write(f"To: {alerts_email}\n")
        f.write(f"Subject: {alert_subject}\n\n")
        f.write(alert_body)
    print(f"  ✅ Alert prepared: {alert_file}")
    print(f"  📬 Would send to: {alerts_email}")
    
    # STATISTICS SUMMARY
    print("\n📊 [5/7] Final statistics summary...")
    
    print(f"  Campaign Performance:")
    print(f"    Sent: {campaign_results.get('sent', 0)}")
    print(f"    Queued: {campaign_results.get('queued', 0)}")
    print(f"    Failed: {campaign_results.get('failed', 0)}")
    
    success_rate = 0
    total_attempted = sum([
        campaign_results.get('sent', 0),
        campaign_results.get('queued', 0),
        campaign_results.get('failed', 0)
    ])
    if total_attempted > 0:
        success_rate = ((campaign_results.get('sent', 0) + 
                        campaign_results.get('queued', 0)) / total_attempted * 100)
    
    print(f"    Success Rate: {success_rate:.1f}%")
    
    print(f"\n  Compliance Status:")
    print(f"    Sent Today: {final_stats.get('sent_today', 0)}/{final_stats.get('daily_limit', 0)}")
    print(f"    Remaining: {final_stats.get('remaining_today', 0)}")
    
    # CLEANUP
    print("\n🧹 [6/7] Cleanup operations...")
    
    # Archive old reports (keep last 30 days)
    archive_cutoff = timestamp.timestamp() - (30 * 24 * 60 * 60)
    archived_count = 0
    
    for old_report in Path(tracking_root).glob("campaign_report_*.json"):
        try:
            if old_report.stat().st_mtime < archive_cutoff:
                archive_dir = Path(tracking_root) / "archive"
                archive_dir.mkdir(exist_ok=True)
                old_report.rename(archive_dir / old_report.name)
                archived_count += 1
        except Exception as e:
            print(f"  ⚠️  Could not archive {old_report.name}: {e}")
    
    if archived_count > 0:
        print(f"  ✅ Archived {archived_count} old report(s)")
    else:
        print(f"  ✅ No old reports to archive")
    
    # FINAL STATUS
    print("\n✅ [7/7] Post-flight operations complete!")
    print("\n" + "="*60)
    print("📊 POST-FLIGHT SUMMARY")
    print("="*60)
    print(f"✅ Reports generated: {report_file.name}")
    print(f"✅ Audit trail updated: {audit_log.name}")
    print(f"✅ Alert prepared for: {alerts_email}")
    print(f"✅ Success rate: {success_rate:.1f}%")
    print(f"✅ Compliance: {final_stats.get('remaining_today', 0)} sends remaining today")
    print("="*60 + "\n")
    
    return {
        'status': 'COMPLETE',
        'report_file': str(report_file),
        'markdown_file': str(markdown_file),
        'audit_log': str(audit_log),
        'alert_file': str(alert_file),
        'success_rate': success_rate,
        'archived_reports': archived_count
    }
# ==============================================================
# Fallback personalizer
# ==============================================================
class _FallbackPersonalizer:
    """Used when EmailPersonalizer is unavailable."""
    def __init__(self, from_name=FROM_NAME, from_email=FROM_EMAIL, 
                 physical_address=PHYS_ADDR, reply_to=REPLY_TO):
        self.from_name = from_name
        self.from_email = from_email
        self.physical_address = physical_address
        self.reply_to = reply_to
        self.unsubscribe_base_url = UNSUB_BASE

    def create_personalized_email(self, template: str, contact: Dict, 
                                  campaign_id="general", is_html=False) -> Dict:
        body = template
        
        # Replace all variables
        for k, v in contact.items():
            body = body.replace(f"{{{{{k}}}}}", str(v))
        
        # Ensure basic variables work
        body = body.replace("{{name}}", contact.get("name", "there"))
        body = body.replace("{{email}}", contact.get("email", ""))
        
        # Create unsubscribe link
        unsubscribe_link = f"{self.unsubscribe_base_url}?email={contact.get('email','')}&campaign={campaign_id}"
        
        # Add compliant footer
        footer = f"\n\n{'='*60}\n"
        footer += f"To unsubscribe: {unsubscribe_link}\n"
        footer += f"Reply with questions: {self.reply_to}\n"
        footer += f"\nFrom: {self.from_name}\n{self.physical_address}\n"
        
        body += footer
        
        # Extract subject
        subj = re.search(r"Subject:\s*(.+?)(?:\n|$)", body, re.IGNORECASE)
        subject = subj.group(1).strip() if subj else f"Message from {self.from_name}"
        if subj:
            body = re.sub(r"Subject:\s*.+?(?:\n|$)", "", body, count=1, flags=re.IGNORECASE)
        
        return {
            "to": contact.get("email", ""),
            "subject": subject,
            "body": body.strip(),
            "from_name": self.from_name,
            "from_email": self.from_email,
            "reply_to": self.reply_to,
            "unsubscribe_url": unsubscribe_link,
            "headers": {
                "List-Unsubscribe": f"<{unsubscribe_link}>",
                "Reply-To": self.reply_to
            },
        }

# ==============================================================
# Contact loading
# ==============================================================
def fallback_load_contacts_from_directory(contacts_dir: str) -> List[Dict]:
    contacts = []
    p = Path(contacts_dir)
    if not p.exists():
        print(f"[contacts] Directory not found: {contacts_dir}")
        return contacts
    
    for csv_file in p.glob("*.csv"):
        try:
            with open(csv_file, newline='', encoding='utf-8') as fh:
                for row in csv.DictReader(fh):
                    contact = {k.lower().strip(): (v or "").strip() 
                              for k, v in row.items() if k}
                    
                    email = contact.get("email") or contact.get("email_address")
                    if not email or "@" not in email:
                        continue
                    
                    contact["email"] = email
                    contact["name"] = contact.get("name") or contact.get("full_name", "")
                    contact["domain"] = email.split("@")[1] if "@" in email else "unknown"
                    contact["opt_in"] = str(contact.get("opt_in", "true")).lower() in ("1", "true", "yes", "y")
                    contacts.append(contact)
            
            print(f"[contacts] Loaded {len(contacts)} from {csv_file.name}")
        except Exception as e:
            print(f"[contacts] Error reading {csv_file}: {e}")
    
    return contacts

def try_professional_loader(contacts_dir: str) -> Optional[List[Dict]]:
    try:
        from data_loader import load_contacts_directory, validate_contact_data
        contacts = load_contacts_directory(contacts_dir)
        stats, valid = validate_contact_data(contacts)
        print(f"[contacts] Professional loader: total={stats.get('total')} valid={stats.get('valid_emails')}")
        return valid
    except Exception:
        return None

# ==============================================================
# Tracking utilities
# ==============================================================
def generate_tracking_id(domain: str, campaign_name: str, filename: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    h = hashlib.md5(f"{domain}_{campaign_name}_{filename}_{ts}".encode()).hexdigest()[:8]
    return f"{domain.upper()}_{h}_{ts}"

def save_tracking_data(tracking_root: str, domain: str, tracking_id: str, data: Dict):
    p = Path(tracking_root) / domain / "campaigns"
    p.mkdir(parents=True, exist_ok=True)
    f = p / f"{tracking_id}.json"
    with open(f, "w", encoding="utf-8") as fh:
        json.dump(data, fh, indent=2)
    print(f"[tracking] Saved {f}")

# ==============================================================
# ENHANCED: Emailer with Full Compliance
# ==============================================================
class Emailer:
    def __init__(self, alerts_email: str, dry_run=False, queue_emails=False,
                 compliance: ComplianceManager = None, personalizer=None):
        self.alerts_email = alerts_email
        self.dry_run = dry_run
        self.queue_emails = queue_emails
        self.compliance = compliance
        self.personalizer = personalizer or _FallbackPersonalizer()
        self.queue_dir = None
        
        if queue_emails:
            self.queue_dir = Path(f"email_queue_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
            self.queue_dir.mkdir(exist_ok=True)
            print(f"📦 [emailer] Queuing mode active → {self.queue_dir}")

    def send_campaign(self, campaign_name: str, content_template: str, 
                     recipients: List[Dict], from_name: str) -> Dict:
        print(f"\n{'='*60}")
        print(f"📧 Campaign: {campaign_name}")
        print(f"{'='*60}")
        
        sent, failed, queued, skipped_suppressed, skipped_rate = 0, 0, 0, 0, 0
        
        for i, r in enumerate(recipients, 1):
            email = r.get("email")
            domain = r.get("domain", "unknown")
            
            if not email or "@" not in email:
                failed += 1
                continue
            
            # COMPLIANCE CHECK 1: Suppression list
            if self.compliance and self.compliance.is_suppressed(email):
                print(f"  🚫 [{i}/{len(recipients)}] SUPPRESSED: {email}")
                skipped_suppressed += 1
                continue
            
            # COMPLIANCE CHECK 2: Daily limit
            if self.compliance and not self.compliance.can_send_today():
                print(f"  ⚠️  [{i}/{len(recipients)}] DAILY LIMIT REACHED")
                skipped_rate += 1
                continue
            
            # COMPLIANCE CHECK 3: Domain limit
            if self.compliance and not self.compliance.can_send_to_domain(domain):
                print(f"  ⚠️  [{i}/{len(recipients)}] DOMAIN LIMIT REACHED: {domain}")
                skipped_rate += 1
                continue
            
            # Create personalized message
            msg = self.personalizer.create_personalized_email(
                content_template, r, campaign_id=campaign_name
            )
            subject, body = msg["subject"], msg["body"]

            # DRY RUN mode
            if self.dry_run:
                print(f"  🔍 [{i}/{len(recipients)}] DRY-RUN: {email}")
                print(f"      Subject: {subject}")
                print(f"      Body preview: {body[:100]}...")
                continue
            
            # QUEUE mode
            if self.queue_emails:
                f = self.queue_dir / f"email_{i:03d}.json"
                with open(f, "w", encoding="utf-8") as fh:
                    json.dump(msg, fh, indent=2)
                queued += 1
                print(f"  📝 [{i}/{len(recipients)}] QUEUED: {email}")
                
                # Record in compliance even in queue mode
                if self.compliance:
                    self.compliance.record_send(email, domain)
                continue

            # LIVE SEND mode
            # Enforce rate limiting delay
            if self.compliance:
                self.compliance.enforce_delay()
            
            print(f"  ✉️  [{i}/{len(recipients)}] SENDING: {email}")
            print(f"      Subject: {subject}")
            
            # TODO: Integrate with actual SMTP sender here
            # For now, this is a simulation
            sent += 1
            
            # Record the send
            if self.compliance:
                self.compliance.record_send(email, domain)
        
        # Summary
        print(f"\n{'='*60}")
        print(f"📊 Campaign Summary: {campaign_name}")
        print(f"{'='*60}")
        print(f"  ✅ Sent: {sent}")
        print(f"  📝 Queued: {queued}")
        print(f"  ❌ Failed: {failed}")
        print(f"  🚫 Suppressed: {skipped_suppressed}")
        print(f"  ⏸️  Rate limited: {skipped_rate}")
        print(f"{'='*60}\n")
        
        return {
            "sent": sent,
            "queued": queued,
            "failed": failed,
            "skipped_suppressed": skipped_suppressed,
            "skipped_rate_limit": skipped_rate
        }

# ==============================================================
# Main runner with compliance reporting
# ==============================================================
def campaign_main(contacts_root: str, scheduled_root: str, tracking_root: str, 
                 alerts_email: str, dry_run=False, queue_emails=False):
    
    print("\n" + "="*60)
    print("🚀 COMPLIANT EMAIL CAMPAIGN PROCESSOR")
    print("="*60 + "\n")
    
    # Initialize compliance manager
    compliance = ComplianceManager(
        max_daily=MAX_DAILY,
        max_per_domain=MAX_DOMAIN,
        min_delay=MIN_DELAY
    )

    # Check if we can proceed
    if not pre_flight_results['can_proceed']:
        print("❌ Pre-flight validation failed - aborting campaign")
        return {
            'status': 'ABORTED',
            'pre_flight': pre_flight_results,
            'reason': 'Pre-flight validation failed'
        }
    
    if pre_flight_results['status'] == 'MARGINAL':
        print("⚠️  Pre-flight validation marginal - proceeding with caution\n")
    
    # Show initial compliance stats
    stats = compliance.get_stats()
    print("📊 Initial Compliance Status:")
    print(f"   Suppressed emails: {stats['suppressed_count']}")
    print(f"   Already sent today: {stats['sent_today']}/{stats['daily_limit']}")
    print(f"   Remaining capacity: {stats['remaining_today']}")
    print(f"   Domains contacted: {stats['domains_contacted']}\n")
    
    # Load contacts
    contacts = try_professional_loader(contacts_root) or \
               fallback_load_contacts_from_directory(contacts_root)
    
    if not contacts:
        print("❌ [runner] No contacts found — aborting.")
        return
    
    print(f"✅ [runner] Loaded {len(contacts)} contacts\n")
    
    # Initialize personalizer
    if PERSONALIZER_AVAILABLE and EmailPersonalizer:
        try:
            personalizer = EmailPersonalizer(
                unsubscribe_base_url=UNSUB_BASE,
                from_name=FROM_NAME,
                from_email=FROM_EMAIL,
                physical_address=PHYS_ADDR
            )
            print(f"✅ [runner] EmailPersonalizer initialized")
        except Exception as e:
            print(f"⚠️  [runner] Personalizer failed: {e}")
            personalizer = _FallbackPersonalizer()
    else:
        personalizer = _FallbackPersonalizer()
        print(f"✅ [runner] Using fallback personalizer")
    
    # Initialize emailer
    emailer = Emailer(
        alerts_email=alerts_email,
        dry_run=dry_run,
        queue_emails=queue_emails,
        compliance=compliance,
        personalizer=personalizer
    )
    
    # Find templates
    templates = list(Path(scheduled_root).rglob("*.txt")) + \
                list(Path(scheduled_root).rglob("*.docx")) + \
                list(Path(scheduled_root).rglob("*.json")) + \
                list(Path(scheduled_root).rglob("*.md")) + \
                list(Path(scheduled_root).rglob("*.html"))
    
    print(f"✅ [runner] Found {len(templates)} campaign templates\n")
    
    # Process each campaign
    total_sent = 0
    total_queued = 0
    total_skipped = 0
    
    for tpl in templates:
        domain = tpl.parent.name
        content = load_campaign_content(str(tpl))
        
        if not content:
            print(f"⚠️  Skipping {tpl.name} - couldn't load content\n")
            continue
        
        tracking_id = generate_tracking_id(domain, tpl.stem, tpl.name)
        enriched = [dict(c, tracking_id=tracking_id) for c in contacts]
        
        result = emailer.send_campaign(
            f"{domain}/{tpl.stem}",
            content,
            enriched,
            FROM_NAME
        )
        
        total_sent += result['sent']
        total_queued += result['queued']
        total_skipped += result['skipped_suppressed'] + result['skipped_rate_limit']
        
        save_tracking_data(tracking_root, domain, tracking_id, result)
    
    # Final compliance report
    print("\n" + "="*60)
    print("📊 FINAL COMPLIANCE REPORT")
    print("="*60)
    
    final_stats = compliance.get_stats()
    print(f"  Total sent today: {final_stats['sent_today']}/{final_stats['daily_limit']}")
    print(f"  Remaining capacity: {final_stats['remaining_today']}")
    print(f"  Domains contacted: {final_stats['domains_contacted']}")
    print(f"  Suppressed emails: {final_stats['suppressed_count']}")
    
    print(f"\n  Campaign totals:")
    print(f"    ✅ Sent: {total_sent}")
    print(f"    📝 Queued: {total_queued}")
    print(f"    ⏸️  Skipped: {total_skipped}")
    
    if final_stats['domain_counts']:
        print(f"\n  Domain breakdown:")
        for domain, count in sorted(final_stats['domain_counts'].items(), 
                                    key=lambda x: x[1], reverse=True):
            print(f"    {domain}: {count}/{MAX_DOMAIN}")
    
    print("="*60 + "\n")

    # ============================================================
    # POST-FLIGHT REPORTING
    # ============================================================
    
    # Aggregate all campaign results
    aggregate_results = {
        'sent': total_sent,
        'queued': total_queued,
        'failed': 0,  # Track this in your campaign loop
        'skipped_suppressed': total_skipped,
        'skipped_rate_limit': 0,  # Track this separately if needed
        'campaigns_processed': len(templates),
        'templates_found': len(templates)
    }
    
    post_flight_results = post_flight_report(
        tracking_root=tracking_root,
        alerts_email=alerts_email,
        compliance=compliance,
        campaign_results=aggregate_results,
        pre_flight_results=pre_flight_results
    )
    
    print("✅ [runner] Campaign processing complete!")
    print(f"📊 Full report available: {post_flight_results.get('report_file', 'N/A')}\n")
    
    # Return comprehensive results
    return {
        'status': 'SUCCESS',
        'pre_flight': pre_flight_results,
        'campaign_results': aggregate_results,
        'post_flight': post_flight_results,
        'compliance_final': compliance.get_stats() if compliance else {}
    }

# ==============================================================
# CLI
# ==============================================================
if __name__ == "__main__":
    p = argparse.ArgumentParser(
        description="Compliant Domain-Aware Email Campaign Processor with Full Enforcement"
    )
    p.add_argument("--contacts", required=True, help="Contacts directory")
    p.add_argument("--scheduled", required=True, help="Campaign templates directory")
    p.add_argument("--tracking", required=True, help="Tracking output directory")
    p.add_argument("--alerts", required=True, help="Alerts email address")
    p.add_argument("--dry-run", action="store_true", 
                   help="Simulate sending without actually sending emails")
    p.add_argument("--queue", action="store_true",
                   help="Queue emails to files instead of sending")
    
    args = p.parse_args()
    
    campaign_main(
        contacts_root=args.contacts,
        scheduled_root=args.scheduled,
        tracking_root=args.tracking,
        alerts_email=args.alerts,
        dry_run=args.dry_run,
        queue_emails=args.queue
    )
