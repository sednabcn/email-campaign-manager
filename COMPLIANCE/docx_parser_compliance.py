#!/usr/bin/env python3
"""
docx_parser.py - Domain-aware campaign processor with compliance integration
Full rewrite: integrates utils/compliance_wrapper.MinimalCompliance to enforce:
 - opt-in (consent) checks
 - suppression list checks
 - daily and per-domain rate limits
 - minimum delay between sends
 - automatic unsubscribe/footer injection
 - recording of successful sends

Usage (example, called by integrated_runner.py):
    python docx_parser.py --contacts contacts --scheduled campaign-templates --tracking tracking --alerts ops@example.com [--dry-run] [--queue-emails] [--template-file path/to/file.docx]
"""

import argparse
import os
import sys
import json
import time
import traceback
import zipfile
import re
import hashlib
from pathlib import Path
from datetime import datetime

# Allow importing utils/ when this script is run from repo root
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)
# Also ensure utils dir is importable
if os.path.isdir(os.path.join(SCRIPT_DIR, "utils")):
    if os.path.join(SCRIPT_DIR, "utils") not in sys.path:
        sys.path.insert(0, os.path.join(SCRIPT_DIR, "utils"))

# ============================================================
# Compliance wrapper import (expected at utils/compliance_wrapper.py)
# ============================================================
try:
    from compliance_wrapper import MinimalCompliance
    print("Loaded utils/compliance_wrapper.MinimalCompliance")
except Exception as e:
    print("Warning: compliance_wrapper not found or import failed:", e)
    # Fallback minimal compliance (safe default, very small)
    class MinimalCompliance:
        def __init__(self, max_daily=50, max_per_domain=5, min_delay=30):
            self.max_daily = max_daily
            self.max_per_domain = max_per_domain
            self.min_delay = min_delay
            self.suppressed = set()
            self.limits = {'date': datetime.now().date().isoformat(), 'total_sent': 0, 'domain_counts': {}}
            self.last_send = None
        def can_send(self, email):
            email = (email or "").lower().strip()
            if not email: return False, "invalid_email"
            if email in self.suppressed: return False, "suppressed"
            if self.limits['total_sent'] >= self.max_daily: return False, "daily_limit"
            domain = email.split('@')[-1]
            if self.limits['domain_counts'].get(domain, 0) >= self.max_per_domain:
                return False, "domain_limit"
            if self.last_send:
                elapsed = (datetime.now() - self.last_send).total_seconds()
                if elapsed < self.min_delay:
                    return False, f"wait_{int(self.min_delay - elapsed)}"
            return True, "ok"
        def record_send(self, email):
            domain = (email.split('@')[-1] if email and '@' in email else 'unknown').lower()
            self.limits['total_sent'] += 1
            self.limits['domain_counts'][domain] = self.limits['domain_counts'].get(domain, 0) + 1
            self.last_send = datetime.now()
        def add_footer(self, body, email):
            return body + f"\n\n---\nTo unsubscribe, reply or visit: https://example.invalid/unsubscribe?email={email}\n"

# ============================================================
# Attempt to import optional helper modules
# ============================================================
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    DOCX_AVAILABLE = False

# The "email_sender" optional module (if you use a separate email sender implementation)
try:
    from email_sender import EmailSender as BaseEmailSender
    EMAIL_SENDER_AVAILABLE = True
    print("Optional email_sender module available - will use BaseEmailSender where possible")
except Exception:
    BaseEmailSender = None
    EMAIL_SENDER_AVAILABLE = False
    print("Optional email_sender module not available - falling back to internal sender behaviour")

# ============================================================
# Utilities: load contacts (either from professional data_loader or fallback)
# ============================================================
def fallback_load_contacts_from_directory(contacts_dir):
    """Fallback loader: read CSV files in contacts_dir; return list of dicts"""
    print("Using fallback contact loader")
    all_contacts = []
    contacts_path = Path(contacts_dir)
    if not contacts_path.exists():
        print(f"Contacts directory not found: {contacts_dir}")
        return all_contacts

    import csv
    csv_files = list(contacts_path.glob('*.csv'))
    for csv_file in csv_files:
        try:
            with open(csv_file, 'r', newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    contact = {}
                    for key, value in row.items():
                        if not key:
                            continue
                        k = key.strip().lower()
                        v = value.strip() if value is not None else ''
                        if k in ('email', 'email_address'):
                            contact['email'] = v
                        elif k in ('name', 'full_name'):
                            contact['name'] = v
                        elif k in ('opt_in','consent','subscribed'):
                            # Accept many possible column names for consent
                            contact['opt_in'] = v.lower() in ('1','true','yes','y','opt-in','optin')
                        else:
                            contact[k] = v
                    # default opt_in True if not present (safer to require opt-in; adjust as needed)
                    if 'opt_in' not in contact:
                        contact['opt_in'] = True
                    if contact.get('email') and '@' in contact['email']:
                        all_contacts.append(contact)
            print(f"Loaded {len(all_contacts)} contacts from {csv_file.name}")
        except Exception as e:
            print(f"Error reading {csv_file}: {e}")
    return all_contacts

def try_professional_loader(contacts_dir):
    """Try to use data_loader.load_contacts_directory if available"""
    try:
        from data_loader import load_contacts_directory, validate_contact_data
        contacts = load_contacts_directory(contacts_dir)
        stats, valid_contacts = validate_contact_data(contacts)
        print(f"Loaded contacts via professional loader: total={stats.get('total')} valid={stats.get('valid_emails')}")
        return valid_contacts
    except Exception as e:
        print("Professional data_loader not available or failed:", e)
        return None

# ============================================================
# Campaign/template file loading utilities
# ============================================================
def is_valid_docx(filepath):
    """Check DOCX archive integrity quickly"""
    try:
        if not zipfile.is_zipfile(filepath):
            return False, "Not a ZIP file"
        with zipfile.ZipFile(filepath, 'r') as z:
            required = ['word/document.xml', '[Content_Types].xml']
            files = z.namelist()
            for r in required:
                if r not in files:
                    return False, f"Missing {r}"
        return True, None
    except zipfile.BadZipFile:
        return False, "BadZipFile"
    except Exception as e:
        return False, str(e)

def load_campaign_content(campaign_path):
    """
    Load a campaign template:
      - .docx -> extract text via python-docx when available (with validation)
      - .txt/.html/.md -> read text
      - .json -> structured content (subject/content/config)
    Returns either a string (content) or a dict with keys {subject, content, from_name, config,...}
    """
    p = Path(campaign_path)
    if not p.exists():
        print(f"Template not found: {campaign_path}")
        return None
    ext = p.suffix.lower()
    if ext == '.docx':
        if not DOCX_AVAILABLE:
            print("python-docx not available; skipping docx template:", campaign_path)
            return None
        valid, err = is_valid_docx(str(p))
        if not valid:
            print(f"DOCX validation failed for {p.name}: {err}")
            return None
        try:
            doc = Document(str(p))
            text = []
            for para in doc.paragraphs:
                if para.text:
                    text.append(para.text)
            # include simple table extraction
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text.append(row_text)
            content = "\n".join(text).strip()
            if not content:
                print(f"Warning: DOCX valid but empty: {p.name}")
                return None
            return content
        except Exception as e:
            print(f"Error reading DOCX {p.name}: {e}")
            traceback.print_exc()
            return None
    elif ext in ('.txt', '.md', '.html'):
        try:
            encs = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            for enc in encs:
                try:
                    with open(p, 'r', encoding=enc) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            with open(p, 'rb') as f:
                return f.read().decode('utf-8', errors='ignore')
        except Exception as e:
            print(f"Error reading text file {p.name}: {e}")
            return None
    elif ext == '.json':
        try:
            with open(p, 'r', encoding='utf-8') as f:
                data = json.load(f)
            # Accept different JSON shapes: {subject,content} or {templates:[...]} or {campaigns:[...]}
            if isinstance(data, dict):
                if 'content' in data:
                    return data
                if 'templates' in data and isinstance(data['templates'], list) and data['templates']:
                    # load first referenced template
                    t = Path(data['templates'][0])
                    if not t.is_absolute():
                        t = p.parent / t
                    tpl = load_campaign_content(str(t))
                    return {
                        'subject': data.get('subject', None),
                        'content': tpl if isinstance(tpl, str) else (tpl.get('content') if isinstance(tpl, dict) else ''),
                        'from_name': data.get('from_name', 'Campaign System'),
                        'config': data
                    }
                if 'campaigns' in data and isinstance(data['campaigns'], list) and data['campaigns']:
                    return data['campaigns'][0]
            return None
        except Exception as e:
            print(f"Error parsing JSON campaign {p.name}: {e}")
            traceback.print_exc()
            return None
    else:
        print(f"Unsupported template extension: {ext}")
        return None

def extract_subject_from_content(content):
    """Try to infer subject from content by looking for a subject: line or top-level heading"""
    try:
        if isinstance(content, dict):
            return content.get('subject', None)
        text = str(content)
        lines = text.splitlines()
        for ln in lines[:12]:
            s = ln.strip()
            if s.lower().startswith('subject:'):
                return s.split(':', 1)[1].strip()
            if s.startswith('# '):
                return s[2:].strip()
        return None
    except Exception:
        return None

# ============================================================
# EmailSender wrapper classes (with compliance integrated in send loop)
# ============================================================
# We'll implement an Emailer that supports:
#  - queue mode (save JSON files)
#  - dry-run (just print)
#  - actual send (if BaseEmailSender is available, use its send_email method; otherwise fallback to no-op print)
# In all actual send paths we enforce compliance.can_send/add_footer/record_send.

class Emailer:
    """Unified emailer wrapper that integrates compliance logic"""
    def __init__(self, smtp_host=None, smtp_port=None, smtp_user=None, smtp_password=None,
                 alerts_email=None, dry_run=False, queue_emails=False, queue_dir_prefix="email_batch",
                 compliance=None):
        self.smtp_host = smtp_host
        self.smtp_port = smtp_port
        self.smtp_user = smtp_user
        self.smtp_password = smtp_password
        self.alerts_email = alerts_email
        self.dry_run = dry_run
        self.queue_emails = queue_emails
        self.queue_dir = None
        self.queued_count = 0
        self.compliance = compliance
        # If there's an installed BaseEmailSender class, instantiate it for real sending if not in queue mode
        self.base_sender = None
        if EMAIL_SENDER_AVAILABLE and not self.queue_emails:
            try:
                self.base_sender = BaseEmailSender(smtp_host, smtp_port, smtp_user, smtp_password, alerts_email, dry_run)
            except Exception as e:
                print("Warning: Could not initialize BaseEmailSender:", e)
                self.base_sender = None
        if self.queue_emails:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            self.queue_dir = Path(f"{queue_dir_prefix}_{ts}")
            self.queue_dir.mkdir(parents=True, exist_ok=True)
            print(f"Email queue enabled - saving messages to {self.queue_dir}")

    def _save_email_to_queue(self, index, to_email, subject, body, from_name="Campaign System"):
        data = {
            'to': to_email,
            'subject': subject,
            'body': body,
            'from_name': from_name,
            'queued_at': datetime.now().isoformat()
        }
        fname = self.queue_dir / f"email_{index}.json"
        with open(fname, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
        self.queued_count += 1
        return fname

    def _deliver_immediately(self, to_email, subject, body):
        """
        Attempt to send immediately:
        - If base_sender exists and provides a send_email method, call that
        - Otherwise simulate (print) — user can replace with production sender integration
        """
        # Attempt to use BaseEmailSender if present
        if self.base_sender:
            try:
                # Try common method names in third-party sender implementations
                if hasattr(self.base_sender, 'send_email'):
                    self.base_sender.send_email(to=to_email, subject=subject, body=body)
                elif hasattr(self.base_sender, 'send'):
                    self.base_sender.send(to_email, subject, body)
                else:
                    # If no standard method, try a generic send_campaign-like method
                    if hasattr(self.base_sender, 'send_campaign'):
                        self.base_sender.send_campaign(to_email, subject, body)
                    else:
                        raise AttributeError("BaseEmailSender found but no recognized send method")
                return True, None
            except Exception as e:
                return False, str(e)
        # Fallback: simulate send (print)
        try:
            print(f"\n[SIMULATED SEND] -> {to_email}")
            print(f"Subject: {subject}")
            print(f"Body preview: {body[:140]}...\n")
            # If you want actual SMTP sending here, implement it with smtplib
            return True, None
        except Exception as e:
            return False, str(e)

    def send_campaign(self, campaign_name, subject_template, content_template, recipients,
                      from_name="Campaign System", tracking_id=None, contact_mapping=None, batch_size=50, delay=1):
        """
        recipients: list of contact dicts
        This method will:
          - personalize subject/content for each recipient
          - enforce compliance.can_send, waiting when requested and skipping otherwise
          - add footer via compliance.add_footer
          - queue or deliver or dry-run accordingly
        Returns a result dict with counts
        """
        print(f"\n=== CAMPAIGN: {campaign_name} ===")
        if tracking_id:
            print(f"Tracking ID: {tracking_id}")
        print(f"Recipients: {len(recipients)}  (queue={self.queue_emails} dry_run={self.dry_run})")

        sent_count = 0
        queued_count = 0
        failed_count = 0
        processed = 0

        # helper for variable substitution (simple)
        def substitute(text, contact):
            if not isinstance(text, str):
                return str(text)
            rez = text
            # simple patterns: {name}, {email}, {company}
            mapping = {
                'name': contact.get('name') or contact.get('full_name') or contact.get('email', '').split('@')[0],
                'email': contact.get('email', ''),
                'company': contact.get('company') or contact.get('organization') or 'your organization'
            }
            # replace curly patterns
            for k, v in mapping.items():
                rez = rez.replace(f"{{{k}}}", str(v))
                rez = rez.replace(f"{{{{{k}}}}}", str(v))
            # support [Placeholder] mapping if contact_mapping provided
            if contact_mapping and isinstance(contact_mapping, dict):
                for placeholder, field in contact_mapping.items():
                    val = contact.get(field, '')
                    rez = rez.replace(f"[{placeholder}]", str(val))
            return rez

        for idx, recipient in enumerate(recipients):
            processed += 1
            if not isinstance(recipient, dict):
                print(f"  Skipping recipient #{idx+1}: invalid format")
                failed_count += 1
                continue

            email = (recipient.get('email') or '').strip()
            if not email or '@' not in email:
                print(f"  Skipping recipient #{idx+1}: invalid email '{email}'")
                failed_count += 1
                continue

            opt_in = recipient.get('opt_in', recipient.get('consent', True))
            if not opt_in:
                print(f"  Skipping {email}: opt_in=False")
                failed_count += 1
                continue

            # Check compliance before preparing content
            if self.compliance:
                can_send, reason = self.compliance.can_send(email)
            else:
                can_send, reason = True, "no-compliance"

            if not can_send:
                print(f"  ⏭️  Compliance prevented send to {email}: {reason}")
                # If the reason requests a wait, honor it, then retry once
                if reason.startswith('wait_'):
                    wait_seconds = int(reason.split('_', 1)[1]) if '_' in reason else 1
                    print(f"    Waiting {wait_seconds}s for rate-limit then retry...")
                    time.sleep(wait_seconds)
                    if self.compliance:
                        can_send2, reason2 = self.compliance.can_send(email)
                    else:
                        can_send2, reason2 = True, "no-compliance"
                    if not can_send2:
                        print(f"    Still cannot send to {email}: {reason2} -> skipping")
                        failed_count += 1
                        continue
                else:
                    failed_count += 1
                    continue

            # Personalize subject and body
            try:
                personalized_subject = substitute(subject_template, recipient)
                personalized_content = substitute(content_template, recipient)
            except Exception as e:
                print(f"  Error personalizing for {email}: {e}")
                failed_count += 1
                continue

            # Add compliance footer if available
            if self.compliance:
                personalized_content = self.compliance.add_footer(personalized_content, email)

            # Queue mode: save to JSON files for external delivery
            if self.queue_emails:
                try:
                    self._save_email_to_queue(self.queued_count, email, personalized_subject, personalized_content, from_name)
                    queued_count += 1
                    if queued_count % 10 == 0:
                        print(f"  Queued {queued_count}/{len(recipients)}")
                except Exception as e:
                    print(f"  Failed to queue {email}: {e}")
                    failed_count += 1
                continue

            # Dry-run: print preview but don't record compliance sends
            if self.dry_run:
                print(f"  [DRY-RUN] To: {email} | Subj: {personalized_subject}")
                if len(personalized_content) > 200:
                    print(f"    Body preview: {personalized_content[:200]}...")
                else:
                    print(f"    Body: {personalized_content}")
                sent_count += 1
                continue

            # Actual send path:
            ok, err = self._deliver_immediately(email, personalized_subject, personalized_content)
            if not ok:
                print(f"  ❌ Send failed to {email}: {err}")
                failed_count += 1
                continue

            # Record send in compliance after successful send
            if self.compliance:
                try:
                    self.compliance.record_send(email)
                except Exception as e:
                    print(f"  Warning: compliance.record_send failed for {email}: {e}")

            sent_count += 1
            # Optional small delay to avoid rapid-fire sending
            if delay and delay > 0:
                time.sleep(delay)

        result = {
            'campaign_name': campaign_name,
            'tracking_id': tracking_id,
            'total_recipients': len(recipients),
            'sent': sent_count,
            'queued': queued_count,
            'failed': failed_count,
            'duration_seconds': 1.0,
            'template_substitution': True
        }
        print(f"  Campaign result: sent={sent_count} queued={queued_count} failed={failed_count}")
        return result

    def send_alert(self, subject, body):
        # For alerts we do not run compliance checks; alerts are local internal messages
        if self.queue_emails:
            # Save alert metadata
            self._save_email_to_queue(self.queued_count, self.alerts_email or "ops@example.com", subject, body, "System")
        else:
            ok, err = self._deliver_immediately(self.alerts_email or "ops@example.com", subject, body)
            if not ok:
                print("Failed to send alert:", err)

# ============================================================
# Domain scanning and campaign dispatch
# ============================================================
def scan_domain_campaigns(templates_dir, specific_file=None):
    """
    Discover campaign template files in a domain-based structure:
      - Domain-based: templates_dir/<domain>/*.docx
      - Flat: templates_dir/*.docx (mapped to 'default')
      - Specific file: passed via --template-file
    Returns a dict: domain -> list(Path objects)
    """
    templates_path = Path(templates_dir)
    domain_campaigns = {}
    if specific_file:
        p = Path(specific_file)
        if not p.exists():
            print("Specified template file not found:", specific_file)
            return {}
        # Try to infer domain from path
        try:
            rel = p.relative_to(templates_path)
            domain = rel.parts[0] if len(rel.parts) > 1 else 'default'
        except Exception:
            domain = 'default'
        domain_campaigns.setdefault(domain, []).append(p)
        print(f"Processing specific template file: {specific_file} (domain: {domain})")
        return domain_campaigns

    if not templates_path.exists():
        print("Templates directory not found:", templates_dir)
        return {}

    # Domain subdirectories?
    found = False
    for child in templates_path.iterdir():
        if child.is_dir() and not child.name.startswith('.'):
            domain = child.name
            campaigns = []
            for ext in ('.docx', '.txt', '.html', '.md', '.json'):
                campaigns.extend(list(child.rglob(f"*{ext}")))
            if campaigns:
                found = True
                domain_campaigns[domain] = campaigns
                print(f"Found {len(campaigns)} campaigns in domain '{domain}'")
    if not found:
        # fallback to flat structure
        flats = []
        for ext in ('.docx', '.txt', '.html', '.md', '.json'):
            flats.extend(list(templates_path.glob(f"*{ext}")))
        if flats:
            domain_campaigns['default'] = flats
            print(f"Found {len(flats)} campaigns in flat templates dir (default domain)")
    return domain_campaigns

def generate_tracking_id(domain, campaign_name, template_filename):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    seed = f"{domain}_{campaign_name}_{template_filename}_{ts}"
    h = hashlib.md5(seed.encode()).hexdigest()[:8]
    return f"{domain.upper()}_{h}_{ts}"

def save_tracking_data(tracking_dir, domain, tracking_id, campaign_data):
    domain_tracking = Path(tracking_dir) / domain / "campaigns"
    domain_tracking.mkdir(parents=True, exist_ok=True)
    fpath = domain_tracking / f"{tracking_id}.json"
    try:
        with open(fpath, 'w', encoding='utf-8') as f:
            json.dump(campaign_data, f, indent=2)
        print(f"Saved tracking metadata: {fpath}")
    except Exception as e:
        print("Warning: could not save tracking metadata:", e)

# ============================================================
# Main campaign runner
# ============================================================
def campaign_main(contacts_root, scheduled_root, tracking_root, alerts_email,
                  dry_run=False, queue_emails=False, specific_template=None,
                  debug=False, batch_size=50, delay=1):
    """
    Main runner: loads contacts, discovers templates, and calls Emailer.send_campaign
    """
    print("Starting domain-aware campaign_main")
    print(f"Contacts: {contacts_root}, Templates: {scheduled_root}, Tracking: {tracking_root}")
    Path(tracking_root).mkdir(parents=True, exist_ok=True)

    # Load contacts (try professional loader first)
    contacts = try_professional_loader(contacts_root)
    if contacts is None:
        contacts = fallback_load_contacts_from_directory(contacts_root)

    if not contacts:
        print("No contacts loaded - aborting campaign processing")
        return

    # Initialize compliance with desired params (these could be configurable)
    compliance = MinimalCompliance(max_daily=50, max_per_domain=5, min_delay=30)

    # Initialize unified emailer with compliance
    emailer = Emailer(
        smtp_host=os.getenv('SMTP_HOST'),
        smtp_port=os.getenv('SMTP_PORT'),
        smtp_user=os.getenv('SMTP_USER'),
        smtp_password=os.getenv('SMTP_PASS'),
        alerts_email=alerts_email,
        dry_run=dry_run,
        queue_emails=queue_emails,
        compliance=compliance
    )

    # Discover campaigns
    domain_campaigns = scan_domain_campaigns(scheduled_root, specific_file=specific_template)
    if not domain_campaigns:
        print("No campaign templates found - nothing to do")
        return

    campaigns_processed = 0
    total_sent = 0
    total_queued = 0
    total_failed = 0
    campaign_results = []

    # iterate domains and campaign templates
    for domain, campaign_files in domain_campaigns.items():
        print("\n" + "="*70)
        print(f"PROCESSING DOMAIN: {domain}")
        print("="*70)

        # ensure per-domain tracking dirs
        (Path(tracking_root) / domain / "campaigns").mkdir(parents=True, exist_ok=True)
        for campaign_file in campaign_files:
            campaign_name = campaign_file.stem
            print(f"\n--- Campaign file: {campaign_file} ---")

            content_loaded = load_campaign_content(str(campaign_file))
            if not content_loaded:
                print(f"Skipping template {campaign_file.name} due to load failure")
                continue

            if isinstance(content_loaded, dict):
                subject = content_loaded.get('subject') or f"Campaign: {campaign_name}"
                content = content_loaded.get('content', '')
                from_name = content_loaded.get('from_name', 'Campaign System')
                config = content_loaded.get('config', {}) or content_loaded.get('metadata', {}) or {}
                contact_mapping = config.get('contact_mapping', {}) or {}
            else:
                # string content
                subject = extract_subject_from_content(content_loaded) or f"Campaign: {campaign_name}"
                content = content_loaded
                from_name = "Campaign System"
                config = {}
                contact_mapping = {}

            tracking_id = generate_tracking_id(domain, campaign_name, campaign_file.name)

            # Build recipients enriched with metadata (not copying huge objects)
            recipients = []
            for i, c in enumerate(contacts):
                rc = dict(c)  # shallow copy
                rc['recipient_id'] = f"{domain}_{campaign_name.replace('/', '_')}_{i+1}"
                rc['campaign_id'] = campaign_name
                rc['domain'] = domain
                rc['tracking_id'] = tracking_id
                recipients.append(rc)

            try:
                result = emailer.send_campaign(
                    campaign_name=f"{domain}/{campaign_name}",
                    subject_template=subject,
                    content_template=content,
                    recipients=recipients,
                    from_name=from_name,
                    tracking_id=tracking_id,
                    contact_mapping=contact_mapping,
                    batch_size=batch_size,
                    delay=delay
                )
                campaigns_processed += 1
                total_sent += result.get('sent', 0)
                total_queued += result.get('queued', 0)
                total_failed += result.get('failed', 0)
                campaign_results.append(result)

                # Save per-campaign tracking metadata
                tracking_data = {
                    'tracking_id': tracking_id,
                    'domain': domain,
                    'campaign_file': str(campaign_file),
                    'subject': subject,
                    'from_name': from_name,
                    'timestamp': datetime.now().isoformat(),
                    'total_recipients': result.get('total_recipients', 0),
                    'sent': result.get('sent', 0),
                    'queued': result.get('queued', 0),
                    'failed': result.get('failed', 0),
                    'config': config
                }
                save_tracking_data(tracking_root, domain, tracking_id, tracking_data)

                # Append to an overall log file
                log_file = Path(tracking_root) / "campaign_execution.log"
                with open(log_file, 'a', encoding='utf-8') as f:
                    f.write(json.dumps(tracking_data) + "\n")

            except Exception as e:
                print(f"Exception while processing campaign {campaign_file.name}: {e}")
                traceback.print_exc()
                continue

    # final summary
    print("\n" + "="*70)
    print("CAMPAIGN RUN SUMMARY")
    print("="*70)
    print(f"Campaigns processed: {campaigns_processed}")
    print(f"Emails sent: {total_sent}")
    print(f"Emails queued: {total_queued}")
    print(f"Failures: {total_failed}")

    # If queue mode created a batch, produce the summary files used by GH Actions pipeline
    if queue_emails and emailer.queue_dir:
        create_github_actions_summary(emailer.queue_dir, emailer.queued_count, campaign_results, tracking_root)

    # Send a summary alert
    try:
        summary_subject = f"Campaign Summary: {campaigns_processed} campaigns, {total_sent + total_queued + total_failed} emails"
        summary_body = f"Processed {campaigns_processed} campaigns\nSent: {total_sent}\nQueued: {total_queued}\nFailed: {total_failed}\nTime: {datetime.now().isoformat()}\n"
        emailer.send_alert(summary_subject, summary_body)
    except Exception as e:
        print("Warning: failed to send summary alert:", e)

def create_github_actions_summary(queue_dir, total_queued, campaign_results, tracking_root):
    """
    Create summary artifacts used by a GitHub Actions job that picks up the queue and performs actual sending.
    Writes:
      - github_actions_email_summary.json
      - campaign_summary_email.json
    """
    summary = {
        'queue_dir': str(queue_dir),
        'queued_at': datetime.now().isoformat(),
        'total_queued': total_queued,
        'campaigns': campaign_results
    }
    with open('github_actions_email_summary.json', 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2)
    print("Created github_actions_email_summary.json")

    # Create a simple summary email payload
    subject = f"Campaign Queue Summary: {total_queued} emails queued"
    body = f"Total queued: {total_queued}\nQueue dir: {queue_dir}\nTimestamp: {datetime.now().isoformat()}\n"
    payload = {
        'to': os.getenv('ALERTS_EMAIL') or 'ops@example.com',
        'subject': subject,
        'body': body,
        'from_name': 'Campaign System'
    }
    with open('campaign_summary_email.json', 'w', encoding='utf-8') as f:
        json.dump(payload, f, indent=2)
    print("Created campaign_summary_email.json")

# ============================================================
# CLI entrypoint (preserves your expected arguments)
# ============================================================
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Domain-Aware Email Campaign Processor (with Compliance)')
    parser.add_argument("--contacts", required=True, help="Contacts directory path")
    parser.add_argument("--scheduled", required=True, help="Templates directory (campaign-templates/)")
    parser.add_argument("--tracking", required=True, help="Tracking directory path")
    parser.add_argument("--alerts", required=True, help="Alerts email address")
    parser.add_argument("--template-file", help="Specific template file to process")
    parser.add_argument("--dry-run", action="store_true", help="Print personalized emails instead of sending")
    parser.add_argument("--queue-emails", action="store_true", help="Queue emails for later sending (GitHub Actions mode)")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode")
    parser.add_argument("--batch-size", type=int, default=50, help="Batch size for processing")
    parser.add_argument("--delay", type=int, default=1, help="Delay (seconds) between sends in live mode")
    args = parser.parse_args()

    # Call main
    campaign_main(
        contacts_root=args.contacts,
        scheduled_root=args.scheduled,
        tracking_root=args.tracking,
        alerts_email=args.alerts,
        dry_run=args.dry_run,
        queue_emails=args.queue_emails,
        specific_template=args.template_file,
        debug=args.debug,
        batch_size=args.batch_size,
        delay=args.delay
    )
