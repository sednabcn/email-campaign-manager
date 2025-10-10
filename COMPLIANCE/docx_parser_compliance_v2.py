#!/usr/bin/env python3
"""
docx_parser.py - Domain-aware campaign processor with personalization & compliance

Features:
 - Loads contacts (professional loader or CSV fallback)
 - Loads campaign templates (.docx/.txt/.md/.html/.json)
 - Integrates EmailPersonalizer for compliant footers and unsubscribe links
 - Integrates MinimalCompliance / ContactValidator and SmartRateLimiter when available
 - Supports dry-run, queue mode, and simulated sending (or an external BaseEmailSender)
 - Safe fallbacks if optional modules are missing
"""

from __future__ import annotations

import argparse
import csv
import json
import os
import sys
import time
import zipfile
import traceback
import re
import hashlib
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Tuple

from pathlib import Path
import json

# ==============================================================
# Load compliance configuration
# ==============================================================
CONFIG_PATH = Path("compliance_config.json")

if CONFIG_PATH.exists():
    with open(CONFIG_PATH, "r") as f:
        CONFIG = json.load(f)
else:
    print("⚠️  compliance_config.json not found — using internal defaults.")
    CONFIG = {}

# Extract sections with sensible fallbacks
SENDER_INFO   = CONFIG.get("sender_info", {})
UNSUBSCRIBE   = CONFIG.get("unsubscribe", {})
RATES         = CONFIG.get("rate_limits", {})
COMPLIANCE_OPTS = CONFIG.get("compliance", {})

FROM_NAME  = SENDER_INFO.get("from_name", "Your Name")
FROM_EMAIL = SENDER_INFO.get("from_email", "your.email@domain.com")
PHYS_ADDR  = SENDER_INFO.get("physical_address", "Your Address, City, Country")
UNSUB_BASE = UNSUBSCRIBE.get("base_url", "https://your-domain.com/unsubscribe")

MAX_DAILY   = RATES.get("max_daily_sends", 50)
MAX_DOMAIN  = RATES.get("max_per_domain_daily", 5)
MIN_DELAY   = RATES.get("min_delay_seconds", 30)


# Make repo root and utils importable
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)
if os.path.isdir(os.path.join(SCRIPT_DIR, "utils")) and os.path.join(SCRIPT_DIR, "utils") not in sys.path:
    sys.path.insert(0, os.path.join(SCRIPT_DIR, "utils"))

# Optional dependencies
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    DOCX_AVAILABLE = False

# Try to import the EmailPersonalizer provided by user
try:
    from email_personalizer import EmailPersonalizer
    PERSONALIZER_AVAILABLE = True
except Exception:
    # try utils subfolder
    try:
        from utils.email_personalizer import EmailPersonalizer
        PERSONALIZER_AVAILABLE = True
    except Exception:
        PERSONALIZER_AVAILABLE = False
        EmailPersonalizer = None

# Try to import compliance / validator / rate modules (optional)
try:
    from compliance_wrapper import MinimalCompliance
    COMPLIANCE_AVAILABLE = True
except Exception:
    MinimalCompliance = None
    COMPLIANCE_AVAILABLE = False

try:
    from contact_validator import ContactValidator
    CONTACT_VALIDATOR_AVAILABLE = True
except Exception:
    try:
        from utils.contact_validator import ContactValidator
        CONTACT_VALIDATOR_AVAILABLE = True
    except Exception:
        ContactValidator = None
        CONTACT_VALIDATOR_AVAILABLE = False

try:
    from smart_rate_limits import SmartRateLimiter, TargetingOptimizer
    SMART_RATE_AVAILABLE = True
except Exception:
    try:
        from utils.smart_rate_limits import SmartRateLimiter, TargetingOptimizer
        SMART_RATE_AVAILABLE = True
    except Exception:
        SmartRateLimiter = None
        TargetingOptimizer = None
        SMART_RATE_AVAILABLE = False

# Optional external EmailSender integration
try:
    from email_sender import EmailSender as BaseEmailSender
    EMAIL_SENDER_AVAILABLE = True
except Exception:
    BaseEmailSender = None
    EMAIL_SENDER_AVAILABLE = False

# --------------------------
# Fallback simple personalizer
# --------------------------
class _FallbackPersonalizer:
    """Minimal personalizer used when real EmailPersonalizer isn't available."""
    def __init__(self, from_name="Campaign System", from_email="no-reply@example.invalid", physical_address=""):
        self.from_name = from_name
        self.from_email = from_email
        self.physical_address = physical_address
        self.unsubscribe_base_url = "https://example.invalid/unsubscribe"

    def create_personalized_email(self, template: str, contact: Dict, campaign_id: str = "general", is_html: bool = False) -> Dict:
        # naive replacements for common placeholders
        body = template
        for k, v in contact.items():
            body = body.replace(f"{{{{{k}}}}}", str(v))
        # defaults
        body = body.replace("{{name}}", contact.get("name", "there"))
        # add simple footer
        unsubscribe_link = f"{self.unsubscribe_base_url}?email={contact.get('email','')}&campaign={campaign_id}"
        footer = f"\n\n---\nTo unsubscribe: {unsubscribe_link}\nFrom: {self.from_name} <{self.from_email}>\n{self.physical_address}"
        body = body + footer
        # subject extraction
        subj = None
        m = re.search(r"Subject:\s*(.+?)(?:\n|$)", body, flags=re.IGNORECASE)
        if m:
            subj = m.group(1).strip()
            body = re.sub(r"Subject:\s*.+?(?:\n|$)", "", body, count=1, flags=re.IGNORECASE)
        if not subj:
            subj = f"Message from {self.from_name}"
        return {
            "to": contact.get("email", ""),
            "to_name": contact.get("name", ""),
            "subject": subj,
            "body": body.strip(),
            "from_name": self.from_name,
            "from_email": self.from_email,
            "campaign_id": campaign_id,
            "unsubscribe_url": unsubscribe_link,
            "headers": {
                "List-Unsubscribe": f"<{unsubscribe_link}>"
            },
            "metadata": {"personalized_at": datetime.now().isoformat()}
        }

# --------------------------
# Utility: contact loaders
# --------------------------
def fallback_load_contacts_from_directory(contacts_dir: str) -> List[Dict]:
    """Read CSV files from contacts_dir and return list of contact dicts."""
    contacts: List[Dict] = []
    p = Path(contacts_dir)
    if not p.exists():
        print(f"[contacts] Directory not found: {contacts_dir}")
        return contacts
    csv_files = list(p.glob("*.csv"))
    for csv_file in csv_files:
        try:
            with open(csv_file, newline='', encoding='utf-8') as fh:
                reader = csv.DictReader(fh)
                rows_before = len(contacts)
                for row in reader:
                    contact = {}
                    for k, v in row.items():
                        if not k:
                            continue
                        kk = k.strip().lower()
                        vv = (v or "").strip()
                        if kk in ("email", "email_address"):
                            contact["email"] = vv
                        elif kk in ("name", "full_name"):
                            contact["name"] = vv
                        elif kk in ("opt_in", "consent", "subscribed"):
                            contact["opt_in"] = vv.lower() in ("1", "true", "yes", "y", "opt-in", "optin")
                        else:
                            contact[kk] = vv
                    if "opt_in" not in contact:
                        contact["opt_in"] = True
                    if contact.get("email") and "@" in contact.get("email"):
                        contacts.append(contact)
                print(f"[contacts] Loaded {len(contacts)-rows_before} from {csv_file.name}")
        except Exception as e:
            print(f"[contacts] Error reading {csv_file}: {e}")
    return contacts

def try_professional_loader(contacts_dir: str) -> Optional[List[Dict]]:
    """Attempt to use professional data loader if available."""
    try:
        from data_loader import load_contacts_directory, validate_contact_data
        contacts = load_contacts_directory(contacts_dir)
        stats, valid = validate_contact_data(contacts)
        print(f"[contacts] Professional loader: total={stats.get('total')} valid={stats.get('valid_emails')}")
        return valid
    except Exception:
        return None

# --------------------------
# Template loading
# --------------------------
def is_valid_docx(filepath: str) -> Tuple[bool, Optional[str]]:
    try:
        if not zipfile.is_zipfile(filepath):
            return False, "Not a ZIP archive"
        with zipfile.ZipFile(filepath, "r") as z:
            required = ["word/document.xml", "[Content_Types].xml"]
            files = z.namelist()
            for r in required:
                if r not in files:
                    return False, f"Missing {r}"
        return True, None
    except zipfile.BadZipFile:
        return False, "BadZipFile"
    except Exception as e:
        return False, str(e)

def load_campaign_content(campaign_path: str) -> Optional[object]:
    p = Path(campaign_path)
    if not p.exists():
        print(f"[templates] Not found: {campaign_path}")
        return None
    ext = p.suffix.lower()
    if ext == ".docx":
        if not DOCX_AVAILABLE:
            print("[templates] python-docx not installed; cannot read .docx templates.")
            return None
        ok, err = is_valid_docx(str(p))
        if not ok:
            print(f"[templates] DOCX validation failed for {p.name}: {err}")
            return None
        try:
            doc = Document(str(p))
            parts: List[str] = []
            for para in doc.paragraphs:
                if para.text:
                    parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    parts.append(row_text)
            text = "\n".join(parts).strip()
            if not text:
                print(f"[templates] Warning: DOCX {p.name} empty")
                return None
            return text
        except Exception as e:
            print(f"[templates] Error reading DOCX {p.name}: {e}")
            traceback.print_exc()
            return None
    elif ext in (".txt", ".md", ".html"):
        encs = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        for enc in encs:
            try:
                with open(p, "r", encoding=enc) as fh:
                    return fh.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"[templates] Error reading {p.name} with {enc}: {e}")
                break
        try:
            with open(p, "rb") as fh:
                return fh.read().decode("utf-8", errors="ignore")
        except Exception as e:
            print(f"[templates] Failed to read {p.name}: {e}")
            return None
    elif ext == ".json":
        try:
            with open(p, "r", encoding="utf-8") as fh:
                data = json.load(fh)
            if isinstance(data, dict):
                if "content" in data:
                    return data
                if "templates" in data and isinstance(data["templates"], list) and data["templates"]:
                    t = Path(data["templates"][0])
                    if not t.is_absolute():
                        t = p.parent / t
                    tpl = load_campaign_content(str(t))
                    return {
                        "subject": data.get("subject"),
                        "content": tpl if isinstance(tpl, str) else (tpl.get("content") if isinstance(tpl, dict) else ""),
                        "from_name": data.get("from_name", "Campaign System"),
                        "config": data
                    }
                if "campaigns" in data and isinstance(data["campaigns"], list) and data["campaigns"]:
                    return data["campaigns"][0]
            return None
        except Exception as e:
            print(f"[templates] Error parsing JSON {p.name}: {e}")
            traceback.print_exc()
            return None
    else:
        print(f"[templates] Unsupported extension {ext} for {p.name}")
        return None

def extract_subject_from_content(content: object) -> Optional[str]:
    try:
        if isinstance(content, dict):
            return content.get("subject")
        text = str(content)
        for ln in text.splitlines()[:12]:
            s = ln.strip()
            if s.lower().startswith("subject:"):
                return s.split(":", 1)[1].strip()
            if s.startswith("# "):
                return s[2:].strip()
        return None
    except Exception:
        return None

# --------------------------
# Emailer
# --------------------------
class Emailer:
    def __init__(self,
                 smtp_host: Optional[str] = None,
                 smtp_port: Optional[int] = None,
                 smtp_user: Optional[str] = None,
                 smtp_password: Optional[str] = None,
                 alerts_email: Optional[str] = None,
                 dry_run: bool = False,
                 queue_emails: bool = False,
                 compliance=None,
                 smart_limiter=None,
                 optimizer=None,
                 personalizer=None):
        self.smtp_host = smtp_host or os.getenv("SMTP_HOST")
        self.smtp_port = smtp_port or os.getenv("SMTP_PORT")
        self.smtp_user = smtp_user or os.getenv("SMTP_USER")
        self.smtp_password = smtp_password or os.getenv("SMTP_PASS")
        self.alerts_email = alerts_email
        self.dry_run = dry_run
        self.queue_emails = queue_emails
        self.queue_dir: Optional[Path] = None
        self.queued_count = 0

        # Use provided compliance or fallback
        if compliance is not None:
            self.compliance = compliance
        elif MinimalCompliance is not None:
            # prefer wrapper if available
            try:
                self.compliance = MinimalCompliance()
            except Exception:
                self.compliance = None
        else:
            # no compliance wrapper: try ContactValidator if present
            if CONTACT_VALIDATOR_AVAILABLE and ContactValidator is not None:
                try:
                    self.compliance = ContactValidator()
                except Exception:
                    self.compliance = None
            else:
                self.compliance = None

        # Smart limiter and optimizer
        self.smart_limiter = smart_limiter
        self.optimizer = optimizer

        # Personalizer
        if personalizer is not None:
            self.personalizer = personalizer
        elif PERSONALIZER_AVAILABLE and EmailPersonalizer is not None:
            try:
                # Provide sensible defaults for the personalizer if not configured
                self.personalizer = EmailPersonalizer()
            except Exception:
                self.personalizer = _FallbackPersonalizer()
        else:
            self.personalizer = _FallbackPersonalizer()

        # External sender
        self.base_sender = None
        if EMAIL_SENDER_AVAILABLE and not self.queue_emails:
            try:
                self.base_sender = BaseEmailSender(self.smtp_host, self.smtp_port, self.smtp_user, self.smtp_password, self.alerts_email, self.dry_run)
            except Exception as e:
                print("[emailer] Warning: BaseEmailSender init failed:", e)
                self.base_sender = None

        if self.queue_emails:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            self.queue_dir = Path(f"email_batch_{ts}")
            self.queue_dir.mkdir(parents=True, exist_ok=True)
            print(f"[emailer] Queue enabled: {self.queue_dir}")

    def _save_email_to_queue(self, index: int, to_email: str, subject: str, body: str, from_name: str = "Campaign System") -> Path:
        payload = {
            "to": to_email,
            "subject": subject,
            "body": body,
            "from_name": from_name,
            "queued_at": datetime.now().isoformat()
        }
        fname = (self.queue_dir or Path(".")) / f"email_{index:04d}.json"
        with open(fname, "w", encoding="utf-8") as fh:
            json.dump(payload, fh, indent=2)
        self.queued_count += 1
        return fname

    def _deliver_immediately(self, to_email: str, subject: str, body: str) -> Tuple[bool, Optional[str]]:
        # Prefer external BaseEmailSender if configured
        if self.base_sender:
            try:
                if hasattr(self.base_sender, "send_email"):
                    self.base_sender.send_email(to=to_email, subject=subject, body=body)
                elif hasattr(self.base_sender, "send"):
                    self.base_sender.send(to_email, subject, body)
                else:
                    raise AttributeError("BaseEmailSender has no known send method")
                return True, None
            except Exception as e:
                return False, str(e)
        # Fallback: simulated send (print)
        try:
            print(f"\n[SIMULATED SEND] -> {to_email}")
            print(f"  Subject: {subject}")
            preview = body[:400] + ("..." if len(body) > 400 else "")
            print(f"  Body preview: {preview}\n")
            return True, None
        except Exception as e:
            return False, str(e)

    def send_campaign(self,
                      campaign_name: str,
                      subject_template: str,
                      content_template: str,
                      recipients: List[Dict],
                      from_name: str = "Campaign System",
                      tracking_id: Optional[str] = None,
                      contact_mapping: Optional[Dict] = None,
                      batch_size: int = 50,
                      delay: int = 1) -> Dict:
        """
        Send a campaign, integrating personalization and compliance.
        Returns a result dict.
        """
        print(f"\n=== CAMPAIGN: {campaign_name} ===")
        if tracking_id:
            print(f" Tracking ID: {tracking_id}")
        print(f" Recipients: {len(recipients)} (queue={self.queue_emails} dry_run={self.dry_run})")

        sent_count = 0
        queued_count = 0
        failed_count = 0

        def substitute_basic(text: str, contact: Dict) -> str:
            rez = text or ""
            for k, v in (contact or {}).items():
                rez = rez.replace(f"{{{k}}}", str(v))
                rez = rez.replace(f"{{{{{k}}}}}", str(v))
            return rez

        for idx, recipient in enumerate(recipients):
            if not isinstance(recipient, dict):
                print(f"  Skipping #{idx+1}: invalid recipient format")
                failed_count += 1
                continue

            email_addr = (recipient.get("email") or "").strip()
            if not email_addr or "@" not in email_addr:
                print(f"  Skipping #{idx+1}: invalid email '{email_addr}'")
                failed_count += 1
                continue

            if not recipient.get("opt_in", True):
                print(f"  Skipping {email_addr}: opt_in=False")
                failed_count += 1
                continue

            # 1. Hard compliance / validator
            if self.compliance:
                try:
                    can_send, reason = self.compliance.can_send(email_addr)
                except Exception as e:
                    print(f"  [compliance] can_send error for {email_addr}: {e}")
                    can_send, reason = False, "compliance_error"
            else:
                can_send, reason = True, "no_compliance"

            if not can_send:
                print(f"  ⏭️  Compliance prevented send to {email_addr}: {reason}")
                # support wait_x semantics (some wrappers return 'wait_30')
                if isinstance(reason, str) and reason.startswith("wait_"):
                    try:
                        wait_seconds = int(reason.split("_", 1)[1])
                        print(f"    Waiting {wait_seconds}s due to compliance signal then retrying once...")
                        time.sleep(wait_seconds)
                        if self.compliance:
                            can_send2, reason2 = self.compliance.can_send(email_addr)
                        else:
                            can_send2, reason2 = True, "no_compliance"
                        if not can_send2:
                            print(f"    Still cannot send to {email_addr}: {reason2} -> skipping")
                            failed_count += 1
                            continue
                    except Exception:
                        failed_count += 1
                        continue
                else:
                    failed_count += 1
                    continue

            # 2. Targeting optimizer: skip if recently contacted, opted-out historically, etc.
            if self.optimizer:
                try:
                    campaign_id_for_opt = recipient.get("campaign_id") or campaign_name
                    should_contact, reason_opt = self.optimizer.should_contact(email_addr, campaign_id_for_opt)
                    if not should_contact:
                        print(f"  ⏭️ Skipping {email_addr}: {reason_opt}")
                        failed_count += 1
                        continue
                except Exception as e:
                    print(f"  [optimizer] error for {email_addr}: {e}")

            # 3. Smart limiter: wait if necessary
            if self.smart_limiter:
                try:
                    ready = self.smart_limiter.wait_if_needed(email_addr)
                    if not ready:
                        print(f"  ⏭️ Smart rate limiter blocked send to {email_addr}")
                        failed_count += 1
                        continue
                except Exception as e:
                    print(f"  [smart_limiter] error: {e}")

            # 4. Personalize content using EmailPersonalizer
            try:
                is_html = isinstance(content_template, str) and content_template.strip().startswith("<")
                # Use personalizer API if available
                personalized = self.personalizer.create_personalized_email(
                    template=content_template,
                    contact=recipient,
                    campaign_id=(recipient.get("campaign_id") or campaign_name),
                    is_html=is_html
                )
                personalized_subject = personalized.get("subject") or substitute_basic(subject_template, recipient)
                personalized_body = personalized.get("body") or substitute_basic(content_template, recipient)
            except Exception as e:
                print(f"  [personalizer] error for {email_addr}: {e}")
                # fall back to simple substitution
                try:
                    personalized_subject = substitute_basic(subject_template, recipient)
                    personalized_body = substitute_basic(content_template, recipient)
                    # ensure footer
                    footer = "\n\n---\nTo unsubscribe: reply or visit: https://example.invalid/unsubscribe"
                    personalized_body += footer
                except Exception as e2:
                    print(f"  [fallback personalizer] failed for {email_addr}: {e2}")
                    failed_count += 1
                    continue

            # 5. Queue mode
            if self.queue_emails:
                try:
                    fname = self._save_email_to_queue(queued_count, email_addr, personalized_subject, personalized_body, from_name)
                    queued_count += 1
                    if queued_count % 10 == 0:
                        print(f"  Queued {queued_count}/{len(recipients)}")
                except Exception as e:
                    print(f"  Failed to queue {email_addr}: {e}")
                    failed_count += 1
                continue

            # 6. Dry-run mode
            if self.dry_run:
                print(f"  [DRY-RUN] To: {email_addr} | Subj: {personalized_subject}")
                preview = personalized_body[:300] + ("..." if len(personalized_body) > 300 else "")
                print(f"    Body preview: {preview}")
                sent_count += 1
                # update optimizer simulation
                if self.optimizer:
                    try:
                        self.optimizer.record_contact(email_addr, recipient.get("campaign_id") or campaign_name, result="sent")
                    except Exception:
                        pass
                continue

            # 7. Actual send (or simulated)
            ok, err = self._deliver_immediately(email_addr, personalized_subject, personalized_body)
            if not ok:
                print(f"  ❌ Send failed to {email_addr}: {err}")
                failed_count += 1
                # notify optimizer about bounce if available
                if self.optimizer:
                    try:
                        self.optimizer.record_contact(email_addr, recipient.get("campaign_id") or campaign_name, result="bounced")
                    except Exception:
                        pass
                continue

            # 8. On success: record compliance and rate tracking
            if self.compliance:
                try:
                    self.compliance.record_send(email_addr)
                except Exception as e:
                    print(f"  [compliance] record_send failed for {email_addr}: {e}")

            if self.smart_limiter:
                try:
                    self.smart_limiter.record_send(email_addr, recipient.get("campaign_id") or campaign_name)
                except Exception as e:
                    print(f"  [smart_limiter] record_send failed: {e}")

            if self.optimizer:
                try:
                    self.optimizer.record_contact(email_addr, recipient.get("campaign_id") or campaign_name, result="sent")
                except Exception as e:
                    print(f"  [optimizer] record_contact failed: {e}")

            sent_count += 1
            # Respect small delay between sends to avoid bursts
            if delay and delay > 0:
                time.sleep(delay)

        result = {
            "campaign_name": campaign_name,
            "tracking_id": tracking_id,
            "total_recipients": len(recipients),
            "sent": sent_count,
            "queued": queued_count,
            "failed": failed_count,
            "timestamp": datetime.now().isoformat()
        }
        print(f"  Campaign result: sent={sent_count} queued={queued_count} failed={failed_count}")
        return result

    def send_alert(self, subject: str, body: str):
        if self.queue_emails:
            if not self.queue_dir:
                self.queue_dir = Path("email_alerts")
                self.queue_dir.mkdir(parents=True, exist_ok=True)
            self._save_email_to_queue(self.queued_count, self.alerts_email or "ops@example.com", subject, body, "System")
        else:
            ok, err = self._deliver_immediately(self.alerts_email or "ops@example.com", subject, body)
            if not ok:
                print("[alert] Failed to send alert:", err)

# --------------------------
# Domain campaign discovery
# --------------------------
def scan_domain_campaigns(templates_dir: str, specific_file: Optional[str] = None) -> Dict[str, List[Path]]:
    templates_path = Path(templates_dir)
    domain_campaigns: Dict[str, List[Path]] = {}
    if specific_file:
        p = Path(specific_file)
        if not p.exists():
            print(f"[templates] Specific template not found: {specific_file}")
            return {}
        try:
            rel = p.relative_to(templates_path)
            domain = rel.parts[0] if len(rel.parts) > 1 else "default"
        except Exception:
            domain = "default"
        domain_campaigns.setdefault(domain, []).append(p)
        print(f"[templates] Processing specific template: {specific_file} (domain: {domain})")
        return domain_campaigns

    if not templates_path.exists():
        print(f"[templates] Templates dir not found: {templates_dir}")
        return {}

    found = False
    for child in templates_path.iterdir():
        if child.is_dir() and not child.name.startswith("."):
            campaigns = []
            for ext in (".docx", ".txt", ".html", ".md", ".json"):
                campaigns.extend(list(child.rglob(f"*{ext}")))
            if campaigns:
                domain_campaigns[child.name] = campaigns
                found = True
                print(f"[templates] Found {len(campaigns)} campaigns in domain '{child.name}'")
    if not found:
        flats: List[Path] = []
        for ext in (".docx", ".txt", ".html", ".md", ".json"):
            flats.extend(list(templates_path.glob(f"*{ext}")))
        if flats:
            domain_campaigns["default"] = flats
            print(f"[templates] Found {len(flats)} campaigns in flat templates dir (default)")
    return domain_campaigns

# --------------------------
# Tracking helpers
# --------------------------
def generate_tracking_id(domain: str, campaign_name: str, template_filename: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    seed = f"{domain}_{campaign_name}_{template_filename}_{ts}"
    h = hashlib.md5(seed.encode()).hexdigest()[:8]
    return f"{domain.upper()}_{h}_{ts}"

def save_tracking_data(tracking_dir: str, domain: str, tracking_id: str, campaign_data: Dict):
    domain_tracking = Path(tracking_dir) / domain / "campaigns"
    domain_tracking.mkdir(parents=True, exist_ok=True)
    fpath = domain_tracking / f"{tracking_id}.json"
    try:
        with open(fpath, "w", encoding="utf-8") as fh:
            json.dump(campaign_data, fh, indent=2)
        print(f"[tracking] Saved metadata: {fpath}")
    except Exception as e:
        print(f"[tracking] Could not save metadata: {e}")

def create_github_actions_summary(queue_dir: Path, total_queued: int, campaign_results: List[Dict], tracking_root: str):
    summary = {
        "queue_dir": str(queue_dir),
        "queued_at": datetime.now().isoformat(),
        "total_queued": total_queued,
        "campaigns": campaign_results
    }
    with open("github_actions_email_summary.json", "w", encoding="utf-8") as fh:
        json.dump(summary, fh, indent=2)
    print("[gha] Created github_actions_email_summary.json")
    # small email payload
    subject = f"Campaign Queue Summary: {total_queued} emails queued"
    body = f"Total queued: {total_queued}\nQueue dir: {queue_dir}\nTimestamp: {datetime.now().isoformat()}\n"
    payload = {"to": os.getenv("ALERTS_EMAIL") or "ops@example.com", "subject": subject, "body": body, "from_name": "Campaign System"}
    with open("campaign_summary_email.json", "w", encoding="utf-8") as fh:
        json.dump(payload, fh, indent=2)
    print("[gha] Created campaign_summary_email.json")

# --------------------------
# Main runner
# --------------------------
def campaign_main(contacts_root: str,
                  scheduled_root: str,
                  tracking_root: str,
                  alerts_email: str,
                  dry_run: bool = False,
                  queue_emails: bool = False,
                  specific_template: Optional[str] = None,
                  debug: bool = False,
                  batch_size: int = 50,
                  delay: int = 1,
                  process_replies_before_run: bool = False):
    print("[runner] Starting campaign runner")
    print(f"[runner] Contacts: {contacts_root}  Templates: {scheduled_root}  Tracking: {tracking_root}")
    Path(tracking_root).mkdir(parents=True, exist_ok=True)
    Path("contacts").mkdir(parents=True, exist_ok=True)

    # Optionally run reply_handler to update suppression list (if available)
    if process_replies_before_run:
        try:
            from reply_handler import ReplyHandler
            rh = ReplyHandler()
            print("[runner] Running reply handler to ingest unsubscribes...")
            rh.check_replies(days_back=7, mark_read=False)
        except Exception as e:
            print("[runner] reply_handler not available or failed:", e)

    # Load contacts
    contacts = try_professional_loader(contacts_root)
    if contacts is None:
        contacts = fallback_load_contacts_from_directory(contacts_root)
    if not contacts:
        print("[runner] No contacts loaded - aborting")
        return

    # Prepare compliance / validator
    compliance = None
    if MinimalCompliance is not None:
        try:
            compliance = MinimalCompliance(max_daily=50, max_per_domain=5, min_delay=30)
            print("[runner] Using MinimalCompliance wrapper")
        except Exception as e:
            print("[runner] MinimalCompliance init failed:", e)
            compliance = None
    elif CONTACT_VALIDATOR_AVAILABLE and ContactValidator is not None:
        try:
            compliance = ContactValidator()
            print("[runner] Using ContactValidator fallback")
        except Exception as e:
            print("[runner] ContactValidator init failed:", e)
            compliance = None
    else:
        print("[runner] No compliance module available - proceed with caution")

    # Smart limiter + optimizer (single instance per run)
    limiter = None
    optimizer = None
    if SMART_RATE_AVAILABLE and SmartRateLimiter is not None and TargetingOptimizer is not None:
        try:
            limiter = SmartRateLimiter(max_hourly=10, max_daily=50, max_per_domain_daily=5, min_delay_seconds=30)
            optimizer = TargetingOptimizer()
            print("[runner] SmartRateLimiter & TargetingOptimizer initialized")
        except Exception as e:
            print("[runner] smart rate init failed:", e)
            limiter = None
            optimizer = None

    # Personalizer instance
    if PERSONALIZER_AVAILABLE and EmailPersonalizer is not None:
        try:
            personalizer = EmailPersonalizer()
        except Exception:
            personalizer = _FallbackPersonalizer()
    else:
        personalizer = _FallbackPersonalizer()

    # Emailer
    emailer = Emailer(
        smtp_host=os.getenv("SMTP_HOST"),
        smtp_port=os.getenv("SMTP_PORT"),
        smtp_user=os.getenv("SMTP_USER"),
        smtp_password=os.getenv("SMTP_PASS"),
        alerts_email=alerts_email,
        dry_run=dry_run,
        queue_emails=queue_emails,
        compliance=compliance,
        smart_limiter=limiter,
        optimizer=optimizer,
        personalizer=personalizer
    )

    # Discover campaign templates
    domain_campaigns = scan_domain_campaigns(scheduled_root, specific_file=specific_template)
    if not domain_campaigns:
        print("[runner] No campaigns found - nothing to do")
        return

    campaigns_processed = 0
    total_sent = 0
    total_queued = 0
    total_failed = 0
    campaign_results = []

    for domain, campaign_files in domain_campaigns.items():
        print("\n" + "=" * 60)
        print(f"[runner] PROCESSING DOMAIN: {domain}")
        print("=" * 60)
        (Path(tracking_root) / domain / "campaigns").mkdir(parents=True, exist_ok=True)

        for campaign_file in campaign_files:
            campaign_name = campaign_file.stem
            print(f"\n[runner] Campaign file: {campaign_file}")
            content_loaded = load_campaign_content(str(campaign_file))
            if not content_loaded:
                print(f"[runner] Skipping {campaign_file.name} due to load failure")
                continue

            if isinstance(content_loaded, dict):
                subject = content_loaded.get("subject") or f"Campaign: {campaign_name}"
                content = content_loaded.get("content", "")
                from_name = content_loaded.get("from_name", "Campaign System")
                config = content_loaded.get("config", {}) or {}
                contact_mapping = config.get("contact_mapping", {}) or {}
            else:
                subject = extract_subject_from_content(content_loaded) or f"Campaign: {campaign_name}"
                content = content_loaded
                from_name = "Campaign System"
                config = {}
                contact_mapping = {}

            tracking_id = generate_tracking_id(domain, campaign_name, campaign_file.name)

            # Build recipients enriched with metadata
            recipients: List[Dict] = []
            for i, c in enumerate(contacts):
                rc = dict(c)
                rc["recipient_id"] = f"{domain}_{campaign_name.replace('/', '_')}_{i+1}"
                rc["campaign_id"] = campaign_name
                rc["domain"] = domain
                rc["tracking_id"] = tracking_id
                recipients.append(rc)

            try:
                result = emailer.send_campaign(
                    campaign_name=f"{domain}/{campaign_name}",
                    subject_template=subject,
                    content_template=content,
                    recipients=recipients,
                    from_name=from_name,
                    tracking_id=tracking_id,
                    contact_mapping=contact_mapping,
                    batch_size=batch_size,
                    delay=delay
                )
                campaigns_processed += 1
                total_sent += result.get("sent", 0)
                total_queued += result.get("queued", 0)
                total_failed += result.get("failed", 0)
                campaign_results.append(result)

                tracking_data = {
                    "tracking_id": tracking_id,
                    "domain": domain,
                    "campaign_file": str(campaign_file),
                    "subject": subject,
                    "from_name": from_name,
                    "timestamp": datetime.now().isoformat(),
                    "total_recipients": result.get("total_recipients", 0),
                    "sent": result.get("sent", 0),
                    "queued": result.get("queued", 0),
                    "failed": result.get("failed", 0),
                    "config": config
                }
                save_tracking_data(tracking_root, domain, tracking_id, tracking_data)
                log_file = Path(tracking_root) / "campaign_execution.log"
                with open(log_file, "a", encoding="utf-8") as fh:
                    fh.write(json.dumps(tracking_data) + "\n")

            except Exception as e:
                print(f"[runner] Exception while processing {campaign_file.name}: {e}")
                traceback.print_exc()
                continue

    # Summary
    print("\n" + "=" * 60)
    print("[runner] CAMPAIGN RUN SUMMARY")
    print("=" * 60)
    print(f"Campaigns processed: {campaigns_processed}")
    print(f"Emails sent: {total_sent}")
    print(f"Emails queued: {total_queued}")
    print(f"Failures: {total_failed}")

    if queue_emails and emailer.queue_dir:
        create_github_actions_summary(emailer.queue_dir, emailer.queued_count, campaign_results, tracking_root)

    try:
        summary_subject = f"Campaign Summary: {campaigns_processed} campaigns, {total_sent + total_queued + total_failed} emails"
        summary_body = f"Processed {campaigns_processed} campaigns\nSent: {total_sent}\nQueued: {total_queued}\nFailed: {total_failed}\nTime: {datetime.now().isoformat()}\n"
        emailer.send_alert(summary_subject, summary_body)
    except Exception as e:
        print("[runner] Warning: failed to send summary alert:", e)

# --------------------------
# CLI
# --------------------------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Domain-Aware Email Campaign Processor (with Personalization & Compliance)")
    parser.add_argument("--contacts", required=True, help="Contacts directory path")
    parser.add_argument("--scheduled", required=True, help="Templates directory (campaign-templates/)")
    parser.add_argument("--tracking", required=True, help="Tracking directory path")
    parser.add_argument("--alerts", required=True, help="Alerts email address")
    parser.add_argument("--template-file", help="Specific template file to process")
    parser.add_argument("--dry-run", action="store_true", help="Print personalized emails instead of sending")
    parser.add_argument("--queue-emails", action="store_true", help="Queue emails for later sending (GitHub Actions mode)")
    parser.add_argument("--debug", action="store_true", help="Enable debug mode")
    parser.add_argument("--batch-size", type=int, default=50, help="Batch size for processing")
    parser.add_argument("--delay", type=int, default=1, help="Delay (seconds) between sends in live mode")
    parser.add_argument("--process-replies", action="store_true", help="Run reply_handler to process unsubscribe replies before run")
    args = parser.parse_args()

    campaign_main(
        contacts_root=args.contacts,
        scheduled_root=args.scheduled,
        tracking_root=args.tracking,
        alerts_email=args.alerts,
        dry_run=args.dry_run,
        queue_emails=args.queue_emails,
        specific_template=args.template_file,
        debug=args.debug,
        batch_size=args.batch_size,
        delay=args.delay,
        process_replies_before_run=args.process_replies
    )
