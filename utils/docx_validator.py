#!/usr/bin/env python3
"""
Enhanced DOCX Parser with File Validation
Adds robust validation and error handling for DOCX files
"""

import os
import sys
import zipfile
import traceback
from pathlib import Path
from docx import Document

class DOCXValidator:
    """Validates DOCX file integrity before processing"""
    
    @staticmethod
    def is_valid_zip(filepath):
        """Check if file is a valid ZIP archive"""
        try:
            return zipfile.is_zipfile(filepath)
        except Exception as e:
            print(f"  ❌ Error checking ZIP validity: {e}")
            return False
    
    @staticmethod
    def validate_docx_structure(filepath):
        """Validate DOCX file structure (must contain required XML files)"""
        required_files = [
            'word/document.xml',
            '[Content_Types].xml'
        ]
        
        try:
            with zipfile.ZipFile(filepath, 'r') as zip_file:
                file_list = zip_file.namelist()
                
                for required_file in required_files:
                    if required_file not in file_list:
                        print(f"  ⚠️  Missing required file: {required_file}")
                        return False
                
                return True
        except zipfile.BadZipFile:
            print(f"  ❌ File is not a valid ZIP archive")
            return False
        except Exception as e:
            print(f"  ❌ Error validating structure: {e}")
            return False
    
    @staticmethod
    def get_file_info(filepath):
        """Get detailed file information for diagnostics"""
        try:
            stat = os.stat(filepath)
            return {
                'size': stat.st_size,
                'size_kb': stat.st_size / 1024,
                'exists': True,
                'readable': os.access(filepath, os.R_OK)
            }
        except Exception as e:
            return {
                'error': str(e),
                'exists': False
            }
    
    @staticmethod
    def validate_docx_file(filepath, verbose=True):
        """
        Comprehensive DOCX file validation
        Returns: (is_valid, error_message, file_info)
        """
        filepath = Path(filepath)
        
        # Check 1: File exists
        if not filepath.exists():
            return False, f"File does not exist: {filepath}", None
        
        # Check 2: Get file info
        file_info = DOCXValidator.get_file_info(filepath)
        
        if verbose:
            print(f"  📄 File: {filepath.name}")
            print(f"  📊 Size: {file_info.get('size_kb', 0):.2f} KB")
        
        # Check 3: File is not empty
        if file_info.get('size', 0) == 0:
            return False, "File is empty (0 bytes)", file_info
        
        # Check 4: File is readable
        if not file_info.get('readable', False):
            return False, "File is not readable (permission denied)", file_info
        
        # Check 5: File is a valid ZIP
        if not DOCXValidator.is_valid_zip(filepath):
            return False, "File is not a valid ZIP archive (DOCX files are ZIP-based)", file_info
        
        # Check 6: DOCX structure is valid
        if not DOCXValidator.validate_docx_structure(filepath):
            return False, "Invalid DOCX structure (missing required XML files)", file_info
        
        if verbose:
            print(f"  ✅ Validation passed")
        
        return True, None, file_info


def load_campaign_content_safe(campaign_path, verbose=True):
    """
    Enhanced version of load_campaign_content with validation
    """
    try:
        file_ext = os.path.splitext(campaign_path)[1].lower()
        campaign_path = Path(campaign_path)
        
        if verbose:
            print(f"\n🔍 Processing: {campaign_path.name}")
        
        if file_ext == '.docx':
            # Validate DOCX before attempting to open
            is_valid, error_msg, file_info = DOCXValidator.validate_docx_file(
                campaign_path, 
                verbose=verbose
            )
            
            if not is_valid:
                print(f"  ❌ DOCX Validation Failed: {error_msg}")
                if file_info:
                    print(f"  📋 File Info: {file_info}")
                
                # Try to provide helpful suggestions
                if file_info and file_info.get('size', 0) == 0:
                    print(f"  💡 Suggestion: File is empty - regenerate or replace it")
                elif "ZIP" in error_msg:
                    print(f"  💡 Suggestion: File may be corrupted or not a real DOCX")
                    print(f"     Try opening it in Word/LibreOffice to verify")
                
                return None
            
            # File is valid, proceed with opening
            try:
                doc = Document(campaign_path)
                content = ""
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content += paragraph.text + "\n"
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content += cell.text + " "
                        content += "\n"
                
                content = content.strip()
                
                if not content:
                    print(f"  ⚠️  Warning: DOCX file is valid but contains no extractable text")
                    return None
                
                if verbose:
                    print(f"  ✅ Extracted {len(content)} characters")
                
                return content
                
            except Exception as e:
                print(f"  ❌ Error reading DOCX content: {str(e)}")
                traceback.print_exc()
                return None
        
        elif file_ext == '.json':
            # Handle JSON files (existing logic)
            import json
            with open(campaign_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        elif file_ext in ['.txt', '.html', '.md']:
            # Handle text files with multiple encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(campaign_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    if verbose:
                        print(f"  ✅ Loaded with {encoding} encoding")
                    return content
                except UnicodeDecodeError:
                    continue
            
            # Fallback to binary read
            with open(campaign_path, 'rb') as f:
                raw_content = f.read()
                content = raw_content.decode('utf-8', errors='ignore')
                print(f"  ⚠️  Used fallback encoding (may have character issues)")
                return content
        
        else:
            print(f"  ❌ Unsupported file format: {file_ext}")
            return None
        
    except Exception as e:
        print(f"  ❌ Error loading campaign content: {str(e)}")
        traceback.print_exc()
        return None


def validate_campaign_directory(templates_dir, verbose=True):
    """
    Validate all DOCX files in a campaign directory before processing
    Returns: (valid_files, invalid_files, summary)
    """
    print(f"\n🔍 Validating campaign directory: {templates_dir}")
    print("=" * 70)
    
    templates_path = Path(templates_dir)
    
    if not templates_path.exists():
        print(f"❌ Directory does not exist: {templates_dir}")
        return [], [], None
    
    valid_files = []
    invalid_files = []
    
    # Find all DOCX files
    docx_files = list(templates_path.rglob("*.docx"))
    
    if not docx_files:
        print(f"⚠️  No DOCX files found in {templates_dir}")
        return [], [], None
    
    print(f"📁 Found {len(docx_files)} DOCX file(s)")
    print()
    
    for docx_file in docx_files:
        relative_path = docx_file.relative_to(templates_path)
        print(f"Checking: {relative_path}")
        
        is_valid, error_msg, file_info = DOCXValidator.validate_docx_file(
            docx_file,
            verbose=verbose
        )
        
        if is_valid:
            valid_files.append(docx_file)
        else:
            invalid_files.append({
                'file': docx_file,
                'error': error_msg,
                'info': file_info
            })
        
        print()
    
    # Summary
    summary = {
        'total': len(docx_files),
        'valid': len(valid_files),
        'invalid': len(invalid_files),
        'success_rate': (len(valid_files) / len(docx_files) * 100) if docx_files else 0
    }
    
    print("=" * 70)
    print(f"📊 VALIDATION SUMMARY")
    print(f"  Total files: {summary['total']}")
    print(f"  ✅ Valid: {summary['valid']}")
    print(f"  ❌ Invalid: {summary['invalid']}")
    print(f"  Success rate: {summary['success_rate']:.1f}%")
    
    if invalid_files:
        print(f"\n⚠️  INVALID FILES:")
        for item in invalid_files:
            print(f"  • {item['file'].name}")
            print(f"    Error: {item['error']}")
    
    return valid_files, invalid_files, summary


# CLI Interface for standalone validation
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Validate DOCX files for campaign processing")
    parser.add_argument("path", help="Path to DOCX file or directory")
    parser.add_argument("--verbose", action="store_true", help="Verbose output")
    parser.add_argument("--fix", action="store_true", help="Attempt to fix common issues")
    
    args = parser.parse_args()
    path = Path(args.path)
    
    if path.is_file():
        # Validate single file
        print(f"🔍 Validating single file: {path.name}\n")
        is_valid, error_msg, file_info = DOCXValidator.validate_docx_file(
            path,
            verbose=True
        )
        
        if is_valid:
            print(f"\n✅ File is valid and ready for processing")
            
            # Try to load content
            print(f"\n📄 Attempting to load content...")
            content = load_campaign_content_safe(path, verbose=True)
            
            if content:
                print(f"\n✅ Successfully loaded {len(content)} characters")
                print(f"Preview: {content[:200]}...")
            else:
                print(f"\n❌ Could not load content")
                sys.exit(1)
        else:
            print(f"\n❌ Validation failed: {error_msg}")
            sys.exit(1)
    
    elif path.is_dir():
        # Validate directory
        valid_files, invalid_files, summary = validate_campaign_directory(
            path,
            verbose=args.verbose
        )
        
        if invalid_files:
            sys.exit(1)
        else:
            print(f"\n✅ All files are valid!")
    
    else:
        print(f"❌ Path does not exist: {path}")
        sys.exit(1)
