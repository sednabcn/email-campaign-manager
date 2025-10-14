#!/usr/bin/env python3
import argparse
import os
import sys
import traceback
import json
import re
import hashlib
from datetime import datetime
from pathlib import Path
import hashlib
import base64
from pathlib import Path
import json
import traceback

# ============================================================================
# UNSUBSCRIBE SYSTEM
# ============================================================================

class UnsubscribeManager:
    """Manages unsubscribe list and link generation"""
    
    def __init__(self, tracking_dir="tracking", base_url="https://sednabcn.github.io/unsubscribe"):
        self.tracking_dir = Path(tracking_dir)
        self.tracking_dir.mkdir(exist_ok=True)
        self.unsubscribe_file = self.tracking_dir / "unsubscribed.json"
        self.base_url = base_url.rstrip('/')
        self.unsubscribed = self._load_unsubscribed()
        
        print(f"✅ UnsubscribeManager initialized")
        print(f"   Base URL: {self.base_url}")
        print(f"   Unsubscribed count: {len(self.unsubscribed)}")
    
    def _load_unsubscribed(self):
        """Load unsubscribe list from file"""
        if not self.unsubscribe_file.exists():
            return {}
        
        try:
            with open(self.unsubscribe_file, 'r') as f:
                return json.load(f)
        except Exception as e:
            print(f"⚠️  Warning: Error loading unsubscribe list: {e}")
            return {}
    
    def _save_unsubscribed(self):
        """Save unsubscribe list to file"""
        try:
            with open(self.unsubscribe_file, 'w') as f:
                json.dump(self.unsubscribed, f, indent=2)
        except Exception as e:
            print(f"⚠️  Error saving unsubscribe list: {e}")
    
    def is_unsubscribed(self, email, campaign_id=None):
        """
        Check if email is unsubscribed
        
        Args:
            email: Email address to check
            campaign_id: Optional campaign ID to check specific campaign unsubscribe
            
        Returns:
            True if unsubscribed, False otherwise
        """
        email_lower = email.lower().strip()
        
        if email_lower not in self.unsubscribed:
            return False
        
        unsub_data = self.unsubscribed[email_lower]
        
        # Check if unsubscribed from all campaigns
        if 'all' in unsub_data.get('campaigns', []):
            return True
        
        # Check if unsubscribed from specific campaign
        if campaign_id and campaign_id in unsub_data.get('campaigns', []):
            return True
        
        return False
    
    def add_unsubscribe(self, email, campaign_id='all', reason=None):
        """
        Add email to unsubscribe list
        
        Args:
            email: Email address
            campaign_id: Campaign ID or 'all' for global unsubscribe
            reason: Optional reason for unsubscribe
        """
        email_lower = email.lower().strip()
        
        if email_lower not in self.unsubscribed:
            self.unsubscribed[email_lower] = {
                'campaigns': [],
                'unsubscribed_at': datetime.now().isoformat(),
                'reason': reason
            }
        
        if campaign_id not in self.unsubscribed[email_lower]['campaigns']:
            self.unsubscribed[email_lower]['campaigns'].append(campaign_id)
        
        self._save_unsubscribed()
        print(f"📛 Added {email} to unsubscribe list (campaign: {campaign_id})")
    
    def generate_unsubscribe_link(self, email, campaign_id='general'):
        """
        Generate unsubscribe link with encoded email
        
        Args:
            email: Recipient email address
            campaign_id: Campaign identifier
            
        Returns:
            Complete unsubscribe URL
        """
        # URL encode email and campaign
        from urllib.parse import quote
        
        email_encoded = quote(email)
        campaign_encoded = quote(campaign_id)
        
        # Generate verification token (optional security measure)
        token = self._generate_token(email, campaign_id)
        
        return f"{self.base_url}?email={email_encoded}&campaign={campaign_encoded}&token={token}"
    
    def _generate_token(self, email, campaign_id):
        """Generate verification token for unsubscribe link"""
        # Simple token generation - you can make this more secure
        seed = f"{email.lower()}{campaign_id}unsubscribe_salt"
        return hashlib.md5(seed.encode()).hexdigest()[:8]
    
    def verify_token(self, email, campaign_id, token):
        """Verify unsubscribe token"""
        expected_token = self._generate_token(email, campaign_id)
        return token == expected_token
    
    def get_stats(self):
        """Get unsubscribe statistics"""
        total = len(self.unsubscribed)
        
        if total == 0:
            return {
                'total_unsubscribed': 0,
                'global_unsubscribes': 0,
                'campaign_specific': 0
            }
        
        global_unsub = sum(1 for data in self.unsubscribed.values() 
                          if 'all' in data.get('campaigns', []))
        campaign_specific = total - global_unsub
        
        return {
            'total_unsubscribed': total,
            'global_unsubscribes': global_unsub,
            'campaign_specific': campaign_specific
        }


def add_unsubscribe_footer(content, unsubscribe_link, is_html=True):
    """
    Add unsubscribe footer to email content
    
    Args:
        content: Email body content
        unsubscribe_link: Unsubscribe URL
        is_html: Whether content is HTML or plain text
        
    Returns:
        Content with unsubscribe footer appended
    """
    if is_html:
        footer = f"""
<hr style="margin-top: 30px; border: none; border-top: 1px solid #ccc;">
<div style="font-size: 11px; color: #666; margin-top: 15px; text-align: center;">
    <p><strong>Professional Outreach</strong></p>
    <p>You received this email as part of professional networking outreach.</p>
    <p>
        If you prefer not to receive future emails, you can 
        <a href="{unsubscribe_link}" style="color: #0066cc; text-decoration: underline;">unsubscribe here</a>.
    </p>
    <p style="font-size: 10px; color: #999; margin-top: 10px;">
        This is professional outreach. We respect your preferences and will honor all opt-out requests immediately.
    </p>
</div>
"""
    else:
        footer = f"""

---
Professional Outreach

You received this email as part of professional networking outreach.

If you prefer not to receive future emails, please visit:
{unsubscribe_link}

Or reply with "UNSUBSCRIBE" in the subject line.

This is professional outreach. We respect your preferences 
and will honor all opt-out requests immediately.
"""
    
    return content + footer


# ============================================================================
# ENHANCED EmailSender WITH UNSUBSCRIBE SUPPORT
# ============================================================================

def enhance_email_sender_with_unsubscribe():
    """
    Patch to add unsubscribe functionality to EmailSender class.
    Add this method call at the beginning of campaign_main()
    """
    
    # Store original send_campaign method
    original_send_campaign = EmailSender.send_campaign
    
    def send_campaign_with_unsubscribe(self, campaign_name, subject, content, recipients, 
                                      from_name="Campaign System", tracking_id=None, 
                                      contact_mapping=None):
        """Enhanced send_campaign with unsubscribe filtering and footer injection"""
        
        # Initialize unsubscribe manager if not exists
               
        # Filter out unsubscribed recipients
        filtered_recipients = []
        skipped_count = 0
        
        for recipient in recipients:
            email = recipient.get('email', '').strip()
            
            if not email:
                skipped_count += 1
                continue
            
            # Check if unsubscribed
            if self.unsubscribe_manager.is_unsubscribed(email, campaign_name):
                print(f"  Skipping {email} - unsubscribed")
                skipped_count += 1
                continue
            
            # Generate unsubscribe link for this recipient
            unsubscribe_link = self.unsubscribe_manager.generate_unsubscribe_link(
                email, 
                campaign_name
            )
            
            # Add unsubscribe link to recipient data
            recipient['unsubscribe_link'] = unsubscribe_link
            
            filtered_recipients.append(recipient)
        
        if skipped_count > 0:
            print(f"  Filtered out {skipped_count} unsubscribed recipients")
        
        if not filtered_recipients:
            print(f"  No valid recipients remaining after filtering")
            return {
                'campaign_name': campaign_name,
                'tracking_id': tracking_id,
                'total_recipients': len(recipients),
                'sent': 0,
                'queued': 0,
                'failed': 0,
                'skipped_unsubscribed': skipped_count,
                'duration_seconds': 0
            }
        
        # Process content with unsubscribe footer
        enhanced_recipients = []
        for recipient in filtered_recipients:
            recipient_copy = recipient.copy()
            
            # Personalize content first
            personalized_content = self.substitute_variables(
                content, 
                recipient, 
                contact_mapping=contact_mapping
            )
            
            # Detect if HTML
            is_html = '<html' in content.lower() or '<body' in content.lower() or '<p>' in content.lower()
            
            # Add unsubscribe footer
            personalized_content = add_unsubscribe_footer(
                personalized_content,
                recipient['unsubscribe_link'],
                is_html=is_html
            )
            
            # Store enhanced content back
            recipient_copy['_enhanced_content'] = personalized_content
            enhanced_recipients.append(recipient_copy)
        
        # Call original method with enhanced recipients
        # We'll need to handle the content separately
        result = original_send_campaign(
            self,
            campaign_name,
            subject,
            content,  # Original content - will be re-personalized
            enhanced_recipients,
            from_name,
            tracking_id,
            contact_mapping
        )
        
        # Add skipped count to result
        result['skipped_unsubscribed'] = skipped_count
        
        return result
    
    # Apply the patch
    EmailSender.send_campaign = send_campaign_with_unsubscribe
    print("✅ EmailSender enhanced with unsubscribe support")


# ============================================================================
# INTEGRATION INSTRUCTIONS
# ============================================================================

"""
INTEGRATION STEPS:

1. Add UnsubscribeManager class and helper functions to docx_parser.py

2. In campaign_main(), add at the beginning (after emailer initialization):

    # Initialize unsubscribe system
    unsubscribe_manager = UnsubscribeManager(
        tracking_dir=tracking_root,
        base_url="https://sednabcn.github.io/unsubscribe"
    )
    
    # Attach to emailer
    emailer.unsubscribe_manager = unsubscribe_manager
    
    # Enhance emailer with unsubscribe
    enhance_email_sender_with_unsubscribe()

3. Update your email sending loop to filter unsubscribed:

    # Before sending campaign
    valid_recipients = []
    for contact in all_contacts:
        email = contact.get('email', '').strip()
        
        # Check unsubscribe status
        if unsubscribe_manager.is_unsubscribed(email, campaign_name):
            print(f"  Skipping {email} - unsubscribed")
            continue
        
        # Add unsubscribe link
        contact['unsubscribe_link'] = unsubscribe_manager.generate_unsubscribe_link(
            email, 
            campaign_name
        )
        
        valid_recipients.append(contact)

4. Update email content to include unsubscribe footer:

    # For each email
    personalized_content = emailer.substitute_variables(content, contact)
    
    # Add unsubscribe footer
    is_html = '<html' in content.lower() or '<p>' in content.lower()
    final_content = add_unsubscribe_footer(
        personalized_content,
        contact['unsubscribe_link'],
        is_html=is_html
    )

5. Add unsubscribe statistics to final summary:

    unsub_stats = unsubscribe_manager.get_stats()
    print(f"Unsubscribe list: {unsub_stats['total_unsubscribed']} emails")

6. IMPORTANT: Create the GitHub Pages unsubscribe page at:
   sednabcn.github.io/unsubscribe/index.html
   
   (Use the HTML template I provided earlier)

7. Manual unsubscribe processing:
   - Users visit your unsubscribe page
   - They submit their email
   - You manually add them to tracking/unsubscribed.json:
   
   {
     "user@example.com": {
       "campaigns": ["all"],
       "unsubscribed_at": "2025-10-10T10:00:00",
       "reason": "User request"
     }
   }

8. Automated unsubscribe (future enhancement):
   - Set up Google Forms backend
   - Form submissions go to Google Sheets
   - Use Google Sheets API to sync to unsubscribed.json
   - Or manually export CSV and update JSON file

TESTING:

1. Test unsubscribe link generation:
   python3 -c "from docx_parser import UnsubscribeManager; um = UnsubscribeManager(); print(um.generate_unsubscribe_link('test@example.com', 'campaign1'))"

2. Test unsubscribe filtering:
   # Add test entry to tracking/unsubscribed.json
   # Run campaign with --dry-run
   # Verify unsubscribed emails are skipped

3. Test footer injection:
   # Run with --dry-run
   # Check dryrun.log for unsubscribe footer in content
"""

# ============================================================================
# EXAMPLE USAGE
# ============================================================================

if __name__ == "__main__":
    # Example: Initialize unsubscribe manager
    manager = UnsubscribeManager(
        tracking_dir="tracking",
        base_url="https://sednabcn.github.io/unsubscribe"
    )
    
    # Example: Generate link
    link = manager.generate_unsubscribe_link("test@example.com", "campaign1")
    print(f"Unsubscribe link: {link}")
    
    # Example: Check if unsubscribed
    is_unsub = manager.is_unsubscribed("test@example.com", "campaign1")
    print(f"Is unsubscribed: {is_unsub}")
    
    # Example: Add to unsubscribe list
    manager.add_unsubscribe("test@example.com", "all", "User request")
    
    # Example: Get statistics
    stats = manager.get_stats()
    print(f"Unsubscribe stats: {stats}")
    
    # Example: Add footer to content
    html_content = "<html><body><h1>Test</h1><p>Content here</p></body></html>"
    enhanced = add_unsubscribe_footer(html_content, link, is_html=True)
    print("\nEnhanced HTML:")
    print(enhanced)
    
# Add current directory to Python path for imports
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if SCRIPT_DIR not in sys.path:
    sys.path.insert(0, SCRIPT_DIR)

# Environment detection
IS_REMOTE = os.getenv('GITHUB_ACTIONS') is not None or os.getenv('CI') is not None

# GitHub Actions email integration
try:
    from email_sender import GitHubActionsEmailSender
    GITHUB_ACTIONS_EMAIL_AVAILABLE = True
    print("GitHub Actions email sender available")
except ImportError:
    GITHUB_ACTIONS_EMAIL_AVAILABLE = False
    print("GitHub Actions email sender not available - using standard EmailSender")

# Import the professional data loader
try:
    from data_loader import load_contacts_directory, validate_contact_data
    DATA_LOADER_AVAILABLE = True
    print("Using professional data_loader module")
except ImportError:
    DATA_LOADER_AVAILABLE = False
    print("Warning: data_loader module not found, using fallback functions")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx library not available")


# ============================================================================
# EMAIL QUEUE FUNCTIONS
# ============================================================================

def save_email_to_queue(batch_dir, email_index, to_email, subject, body, from_name="Campaign System"):
    """Save individual email to queue directory"""
    email_data = {
        'to': to_email,
        'subject': subject,
        'body': body,
        'from_name': from_name,
        'queued_at': datetime.now().isoformat(),
        'ready_to_send': True
    }
    
    email_file = batch_dir / f"email_{email_index}.json"
    with open(email_file, 'w') as f:
        json.dump(email_data, f, indent=2)
    
    return email_file


def create_github_actions_summary(batch_dir, total_queued, campaign_results):
    """Create summary file for GitHub Actions email sender"""
    summary = {
        'batch_directory': str(batch_dir),
        'total_emails': total_queued,
        'queued_at': datetime.now().isoformat(),
        'campaigns': len(campaign_results),
        'campaign_details': campaign_results
    }
    
    with open('github_actions_email_summary.json', 'w') as f:
        json.dump(summary, f, indent=2)
    
    print(f"Created GitHub Actions email summary: {total_queued} emails queued")
    return summary


def create_campaign_summary_email(alerts_email, campaigns_count, total_queued, campaign_results):
    """Create summary email to be sent after batch processing"""
    subject = f"Campaign Summary: {campaigns_count} campaigns, {total_queued} emails queued"
    
    body = f"""Campaign Execution Summary
=========================

Campaigns Processed: {campaigns_count}
Total Emails Queued: {total_queued}
Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

Campaign Details:
"""
    
    for result in campaign_results:
        body += f"\n- {result['campaign_name']}: {result.get('queued', result.get('sent', 0))} emails"
        if result.get('tracking_id'):
            body += f" [ID: {result['tracking_id']}]"
    
    body += "\n\nEmails will be sent by the GitHub Actions email sender job."
    
    summary_data = {
        'to': alerts_email,
        'subject': subject,
        'body': body,
        'from_email': os.getenv('SMTP_USER', 'noreply@example.com')
    }
    
    with open('campaign_summary_email.json', 'w') as f:
        json.dump(summary_data, f, indent=2)
    
    return summary_data


# ============================================================================
# EMAIL SENDER WITH QUEUE SUPPORT
# ============================================================================

try:
    from email_sender import EmailSender as BaseEmailSender
    EMAIL_SENDER_AVAILABLE = True
    
    class EmailSender(BaseEmailSender):
        """Enhanced EmailSender with template variable substitution and queue support"""
        
        def __init__(self, smtp_host=None, smtp_port=None, smtp_user=None, smtp_password=None, 
                     alerts_email=None, dry_run=False, queue_emails=False):
            if not queue_emails:
                super().__init__(smtp_host, smtp_port, smtp_user, smtp_password, alerts_email, dry_run)
            else:
                # Minimal init for queue mode
                self.smtp_host = smtp_host
                self.smtp_port = smtp_port
                self.smtp_user = smtp_user
                self.smtp_password = smtp_password
                self.alerts_email = alerts_email
                self.dry_run = False  # Queue mode is never dry-run
            
            self.queue_emails = queue_emails
            self.queue_batch_dir = None
            self.queued_count = 0
            
            if queue_emails:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                self.queue_batch_dir = Path(f"email_batch_{timestamp}")
                self.queue_batch_dir.mkdir(exist_ok=True)
                print(f"Email queue mode enabled - saving to {self.queue_batch_dir}")
        
        def substitute_variables(self, content, contact, additional_vars=None, contact_mapping=None):
            """
            Enhanced variable substitution with contact_mapping support
            
            Supports multiple placeholder formats:
            - [Recipient Name] with contact_mapping
            - {{variable}}, {variable} for direct field access
            """
            if not isinstance(content, str):
                return str(content)
            
            result = content
            variables = {}
            
            # Phase 1: Handle contact_mapping for [Placeholder] format
            if contact_mapping and isinstance(contact_mapping, dict) and isinstance(contact, dict):
                for template_placeholder, csv_field in contact_mapping.items():
                    value = contact.get(csv_field, '')
                    if value and str(value).lower() not in ['nan', 'none', '']:
                        placeholder_pattern = f'[{template_placeholder}]'
                        result = result.replace(placeholder_pattern, str(value))
                # DEBUG: Print what's happening
                print(f"\n🔍 DEBUG substitute_variables:")
                print(f"   Contact fields: {list(contact.keys())}")
                print(f"   Contact mapping: {contact_mapping}")
                print(f"   Sample values: {dict(list(contact.items())[:3])}")
                print(f"   Content sample: {content[:200]}")
            # Phase 2: Build variables dictionary for generic patterns
            if isinstance(contact, dict):
                for key, value in contact.items():
                    if value is not None:
                        str_value = str(value).strip()
                        variables[key.lower()] = str_value
                        variables[key] = str_value
                        variables[key.replace('_', ' ').title()] = str_value
                        
                        if key.lower() in ['name', 'full_name']:
                            variables['Contact Name'] = str_value
                            variables['contact name'] = str_value
                            variables['name'] = str_value
                        elif key.lower() in ['email', 'email_address']:
                            variables['Contact Email'] = str_value
                            variables['contact email'] = str_value
                            variables['email'] = str_value
                        elif key.lower() in ['company', 'organization']:
                            variables['Company'] = str_value
                            variables['company'] = str_value
                            variables['organization'] = str_value
            
            if additional_vars and isinstance(additional_vars, dict):
                variables.update(additional_vars)
            
            if 'name' not in variables:
                variables['name'] = contact.get('email', '').split('@')[0] if contact.get('email') else 'Friend'
                variables['Contact Name'] = variables['name']
            
            if 'company' not in variables:
                variables['company'] = 'your organization'
                variables['Company'] = 'your organization'
            
            # Phase 3: Handle {{variable}} format
            pattern1 = re.compile(r'\{\{([^}]+)\}\}')
            for match in pattern1.finditer(result):
                var_name = match.group(1).strip()
                if var_name in variables:
                    result = result.replace(match.group(0), variables[var_name])
                else:
                    var_lower = var_name.lower()
                    if var_lower in variables:
                        result = result.replace(match.group(0), variables[var_lower])
            
            # Phase 4: Handle {variable} format
            pattern2 = re.compile(r'\{([^}]+)\}')
            for match in pattern2.finditer(result):
                var_name = match.group(1).strip().lower()
                if var_name in variables:
                    result = result.replace(match.group(0), variables[var_name])
            
            return result
        
        def send_campaign(self, campaign_name, subject, content, recipients, from_name="Campaign System", 
                         tracking_id=None, contact_mapping=None):
            print(f"\n=== CAMPAIGN: {campaign_name} ===")
            if tracking_id:
                print(f"Tracking ID: {tracking_id}")
            print(f"Subject Template: {subject}")
            print(f"From: {from_name}")
            print(f"Recipients: {len(recipients)}")
            
            if self.queue_emails:
                print(f"Queue mode: Saving emails to {self.queue_batch_dir}")
            
            processed_recipients = []
            sent_count = 0
            queued_count = 0
            failed_count = 0
            
            for i, recipient in enumerate(recipients):
                if not isinstance(recipient, dict):
                    print(f"Skipping invalid recipient {i+1}: not a dictionary")
                    failed_count += 1
                    continue
                
                email = recipient.get('email', '').strip()
                if not email or '@' not in email:
                    print(f"Skipping recipient {i+1}: invalid email '{email}'")
                    failed_count += 1
                    continue
                
                try:
                    personalized_subject = self.substitute_variables(subject, recipient, contact_mapping=contact_mapping)
                    personalized_content = self.substitute_variables(content, recipient, contact_mapping=contact_mapping)
                    # ADD THIS: Inject unsubscribe footer
                    if recipient.get('unsubscribe_link'):
                         is_html = '<html' in content.lower() or '<body' in content.lower() or '<p>' in content.lower()
                         personalized_content = add_unsubscribe_footer(
                             personalized_content,
                             recipient['unsubscribe_link'],
                             is_html=is_html
                         )
    
                    
                    if self.queue_emails:
                        # QUEUE MODE: Save to file instead of sending
                        save_email_to_queue(
                            self.queue_batch_dir,
                            self.queued_count,
                            email,
                            personalized_subject,
                            personalized_content,
                            from_name
                        )
                        queued_count += 1
                        self.queued_count += 1
                        
                        if (i + 1) % 10 == 0:
                            print(f"  Queued {queued_count}/{len(recipients)} emails...")
                    
                    elif self.dry_run:
                        processed_recipient = {
                            **recipient,
                            'personalized_subject': personalized_subject,
                            'personalized_content': personalized_content,
                            'original_subject': subject,
                            'original_content': content[:100] + "..." if len(content) > 100 else content,
                            'tracking_id': tracking_id
                        }
                        
                        processed_recipients.append(processed_recipient)
                        sent_count += 1
                        
                        print(f"  {i+1}. {recipient.get('name', 'N/A')} <{email}>")
                        print(f"      Subject: {personalized_subject}")
                        if len(personalized_content) > 150:
                            print(f"      Content: {personalized_content[:150]}...")
                        else:
                            print(f"      Content: {personalized_content}")
                    
                    else:
                        # ACTUAL SEND MODE: Use parent class method
                        processed_recipient = {
                            **recipient,
                            'personalized_subject': personalized_subject,
                            'personalized_content': personalized_content,
                            'tracking_id': tracking_id
                        }
                        processed_recipients.append(processed_recipient)
                        sent_count += 1
                    
                except Exception as e:
                    print(f"Error processing recipient {i+1} ({email}): {str(e)}")
                    failed_count += 1
                    continue
            
            if self.dry_run and not self.queue_emails:
                print("DRY-RUN MODE: No emails sent")
                if len(processed_recipients) > 3:
                    print(f"  ... and {len(processed_recipients) - 3} more recipients with personalized content")
            elif self.queue_emails:
                print(f"Queued {queued_count} emails for campaign {campaign_name}")
            
            return {
                'campaign_name': campaign_name,
                'tracking_id': tracking_id,
                'total_recipients': len(recipients),
                'sent': sent_count if not (self.dry_run or self.queue_emails) else 0,
                'queued': queued_count if self.queue_emails else 0,
                'failed': failed_count,
                'duration_seconds': 1.5,
                'template_substitution': True
            }
        
        def send_alert(self, subject, body):
            if self.queue_emails:
                print(f"ALERT (queued): {subject}")
                print(f"Body: {body[:200]}...")
            else:
                super().send_alert(subject, body)
                
except ImportError:
    print("Warning: email_sender module not found, using fallback")
    EMAIL_SENDER_AVAILABLE = False
    
    class EmailSender:
        def __init__(self, smtp_host=None, smtp_port=None, smtp_user=None, smtp_password=None, 
                     alerts_email=None, dry_run=False, queue_emails=False):
            self.smtp_host = smtp_host
            self.smtp_port = smtp_port
            self.smtp_user = smtp_user
            self.smtp_password = smtp_password
            self.alerts_email = alerts_email
            self.dry_run = dry_run
            self.queue_emails = queue_emails
            self.queue_batch_dir = None
            self.queued_count = 0
            
            if queue_emails:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                self.queue_batch_dir = Path(f"email_batch_{timestamp}")
                self.queue_batch_dir.mkdir(exist_ok=True)
                print(f"Fallback EmailSender in queue mode - saving to {self.queue_batch_dir}")
            else:
                print(f"Fallback EmailSender initialized - dry_run: {dry_run}, alerts: {alerts_email}")
        
        def substitute_variables(self, content, contact, additional_vars=None, contact_mapping=None):
            if not isinstance(content, str):
                return str(content)
            
            result = content
            
            # Handle contact_mapping for [Placeholder] format
            if contact_mapping and isinstance(contact_mapping, dict) and isinstance(contact, dict):
                for template_placeholder, csv_field in contact_mapping.items():
                    value = contact.get(csv_field, '')
                    if value:
                        placeholder_pattern = f'[{template_placeholder}]'
                        result = result.replace(placeholder_pattern, str(value))
            
            # Handle generic patterns
            if isinstance(contact, dict):
                name = contact.get('name', contact.get('email', '').split('@')[0] if contact.get('email') else 'Friend')
                email = contact.get('email', '')
                company = contact.get('company', contact.get('organization', 'your organization'))
                
                replacements = {
                    '{{Contact Name}}': name, '{{contact name}}': name, '{{name}}': name, '{name}': name,
                    '{{Contact Email}}': email, '{{contact email}}': email, '{{email}}': email, '{email}': email,
                    '{{Company}}': company, '{{company}}': company, '{company}': company,
                }
                
                for pattern, value in replacements.items():
                    result = result.replace(pattern, str(value))
            
            return result
        
        def send_campaign(self, campaign_name, subject, content, recipients, from_name="Campaign System", 
                         tracking_id=None, contact_mapping=None):
            print(f"\n=== CAMPAIGN: {campaign_name} ===")
            if tracking_id:
                print(f"Tracking ID: {tracking_id}")
            print(f"Subject Template: {subject}")
            print(f"From: {from_name}")
            print(f"Recipients: {len(recipients)}")
            
            sent_count = 0
            queued_count = 0
            failed_count = 0
            
            if self.queue_emails:
                print(f"Queue mode: Saving to {self.queue_batch_dir}")
                for i, recipient in enumerate(recipients):
                    if isinstance(recipient, dict):
                        email = recipient.get('email', '').strip()
                        if email and '@' in email:
                            personalized_subject = self.substitute_variables(subject, recipient, contact_mapping=contact_mapping)
                            personalized_content = self.substitute_variables(content, recipient, contact_mapping=contact_mapping)

                            
                            save_email_to_queue(
                                self.queue_batch_dir,
                                self.queued_count,
                                email,
                                personalized_subject,
                                personalized_content,
                                from_name
                            )
                            queued_count += 1
                            self.queued_count += 1
                        else:
                            failed_count += 1
                    else:
                        failed_count += 1
                
                print(f"Queued {queued_count} emails")
            
            elif self.dry_run:
                print("DRY-RUN MODE: No emails sent")
                for i, recipient in enumerate(recipients[:3]):
                    if isinstance(recipient, dict):
                        personalized_subject = self.substitute_variables(subject, recipient, contact_mapping=contact_mapping)
                        personalized_content = self.substitute_variables(content, recipient, contact_mapping=contact_mapping)

                        # ADD UNSUBSCRIBE FOOTER (fallback mode)
                        if recipient.get('unsubscribe_link'):
                            is_html = '<html' in content.lower() or '<body' in content.lower() or '<p>' in content.lower()
                            personalized_content = add_unsubscribe_footer(
                                personalized_content,
                                recipient['unsubscribe_link'],
                                is_html=is_html
                            )
                        email = recipient.get('email', 'N/A')
                        name = recipient.get('name', 'N/A')
                        
                        print(f"  {i+1}. {name} <{email}>")
                        print(f"      Subject: {personalized_subject}")
                        print(f"      Content: {personalized_content[:100]}...")
                        sent_count += 1
                    else:
                        failed_count += 1
                
                if len(recipients) > 3:
                    remaining = len(recipients) - 3
                    sent_count += max(0, remaining - failed_count)
                    print(f"  ... and {remaining} more recipients with personalized content")
            
            return {
                'campaign_name': campaign_name,
                'tracking_id': tracking_id,
                'total_recipients': len(recipients),
                'sent': sent_count if not (self.dry_run or self.queue_emails) else 0,
                'queued': queued_count if self.queue_emails else 0,
                'failed': failed_count,
                'duration_seconds': 1.5,
                'template_substitution': True
            }
        
        def send_alert(self, subject, body):
            print(f"ALERT: {subject}")
            print(f"Body: {body[:200]}...")


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def fallback_load_contacts_from_directory(contacts_dir):
    """Fallback contact loader when data_loader module is not available"""
    print("Warning: Using fallback contact loading - data_loader module not found")
    
    all_contacts = []
    contacts_path = Path(contacts_dir)
    
    if not contacts_path.exists():
        print(f"Contacts directory not found: {contacts_dir}")
        return all_contacts
    
    import csv
    csv_files = list(contacts_path.glob('*.csv'))
    
    for csv_file in csv_files:
        try:
            with open(csv_file, 'r', newline='', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                for row in reader:
                    contact = {}
                    for key, value in row.items():
                        if key and value:
                            clean_key = key.strip().lower()
                            clean_value = value.strip()
                            
                            if clean_key in ['email', 'email_address']:
                                contact['email'] = clean_value
                            elif clean_key in ['name', 'full_name']:
                                contact['name'] = clean_value
                            else:
                                contact[clean_key] = clean_value
                    
                    if contact.get('email') and '@' in contact['email']:
                        all_contacts.append(contact)
            
            print(f"Loaded {len(all_contacts)} contacts from {csv_file}")
        except Exception as e:
            print(f"Error loading {csv_file}: {e}")
    
    return all_contacts


def generate_tracking_id(domain, campaign_name, template_name):
    """Generate unique tracking ID for campaign"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    seed_string = f"{domain}_{campaign_name}_{template_name}_{timestamp}"
    hash_object = hashlib.md5(seed_string.encode())
    short_hash = hash_object.hexdigest()[:8]
    tracking_id = f"{domain.upper()}_{short_hash}_{timestamp}"
    return tracking_id


"""
Replace the load_campaign_content function in your docx_parser.py with this enhanced version
"""

import zipfile
from pathlib import Path
from docx import Document

def is_valid_docx(filepath):
    """Check if file is a valid DOCX (ZIP-based) file"""
    try:
        if not zipfile.is_zipfile(filepath):
            return False, "Not a valid ZIP archive"
        
        # Check for required DOCX structure
        with zipfile.ZipFile(filepath, 'r') as zip_file:
            required_files = ['word/document.xml', '[Content_Types].xml']
            file_list = zip_file.namelist()
            
            for required_file in required_files:
                if required_file not in file_list:
                    return False, f"Missing required file: {required_file}"
        
        return True, None
    except zipfile.BadZipFile:
        return False, "Corrupted ZIP archive"
    except Exception as e:
        return False, f"Validation error: {str(e)}"

def validate_campaign_schedule(config):
    """
    Validate campaign schedule based on mode and date
    
    Args:
        config: Campaign configuration dictionary
        
    Returns:
        tuple: (is_valid, reason)
    """
    from datetime import datetime, date
    
    try:
        mode = config.get('mode', 'immediate')
        campaign_date_str = config.get('date')
        
        print(f"  📅 Validating schedule: mode={mode}, date={campaign_date_str}")
        
        # Parse campaign date
        if campaign_date_str:
            try:
                campaign_date = datetime.strptime(campaign_date_str, "%Y-%m-%d").date()
            except ValueError:
                return False, f"Invalid date format: {campaign_date_str} (expected YYYY-MM-DD)"
        else:
            # No date specified - allow immediate mode only
            if mode == 'immediate':
                return True, "Immediate mode with no date restriction"
            else:
                return False, "Date required for non-immediate mode"
        
        today = date.today()
        
        # Mode-based validation
        if mode == 'immediate':
            # Immediate mode: date must be today or future
            if campaign_date < today:
                return False, f"Campaign date {campaign_date} is in the past (today: {today})"
            return True, f"Immediate mode: valid for {campaign_date}"
        
        elif mode == 'schedule':
            # Schedule mode: date must be today or future
            if campaign_date < today:
                return False, f"Scheduled date {campaign_date} has passed (today: {today})"
            elif campaign_date > today:
                return False, f"Scheduled for future date {campaign_date} (today: {today})"
            return True, f"Scheduled for today: {campaign_date}"
        
        elif mode == 'schedule_now':
            # Schedule now: date must be today
            if campaign_date < today:
                return False, f"Campaign date {campaign_date} is in the past"
            elif campaign_date > today:
                return False, f"Scheduled for future: {campaign_date} (not today)"
            return True, f"Schedule now: executing for {campaign_date}"
        
        else:
            return False, f"Unknown mode: {mode}"
        
    except Exception as e:
        return False, f"Schedule validation error: {e}"


def prepare_campaign_isolation(config, campaign_file):
    """
    Prepare isolated tracking and archiving for each campaign
    
    Args:
        config: Campaign configuration
        campaign_file: Path to campaign JSON
        
    Returns:
        dict with isolated paths
    """
    from pathlib import Path
    from datetime import datetime
    
    campaign_name = config.get('name', Path(campaign_file).stem)
    sector = config.get('sector', 'general')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Create isolated paths
    isolation = {
        'campaign_id': f"{sector}_{campaign_name}_{timestamp}".replace(' ', '_'),
        'tracking_dir': Path('tracking') / sector / campaign_name,
        'archive_dir': Path('contact_details_used') / sector / campaign_name,
        'log_file': f"campaign_{sector}_{campaign_name}_{timestamp}.log"
    }
    
    # Create directories
    isolation['tracking_dir'].mkdir(parents=True, exist_ok=True)
    isolation['archive_dir'].mkdir(parents=True, exist_ok=True)
    
    return isolation
    
def load_campaign_content(campaign_path):
    """
    Load campaign content from various file formats with validation
    Enhanced version with DOCX integrity checking
    """
    try:
        file_ext = os.path.splitext(campaign_path)[1].lower()
        campaign_path = Path(campaign_path)
        
        # Handle DOCX files with validation
        if file_ext == '.docx':
            if not DOCX_AVAILABLE:
                print(f"⚠️  Warning: python-docx not available, skipping {campaign_path}")
                return None
            
            # Pre-flight validation
            print(f"  🔍 Validating: {campaign_path.name}")
            
            # Check file exists and has content
            if not campaign_path.exists():
                print(f"  ❌ Error: File does not exist")
                return None
            
            file_size = campaign_path.stat().st_size
            if file_size == 0:
                print(f"  ❌ Error: File is empty (0 bytes)")
                return None
            
            print(f"  📊 File size: {file_size / 1024:.2f} KB")
            
            # Validate DOCX structure
            is_valid, error_msg = is_valid_docx(campaign_path)
            if not is_valid:
                print(f"  ❌ DOCX Validation Failed: {error_msg}")
                print(f"  💡 Suggestion: Regenerate or replace this file")
                return None
            
            print(f"  ✅ DOCX structure valid")
            
            # Attempt to load with python-docx
            try:
                doc = Document(campaign_path)
                content = ""
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content += cell.text + " "
                        content += "\n"
                
                content = content.strip()
                
                if not content:
                    print(f"  ⚠️  Warning: DOCX is valid but contains no text")
                    return None
                
                print(f"  ✅ Extracted {len(content)} characters")
                return content
                
            except Exception as e:
                print(f"  ❌ Error reading DOCX: {str(e)}")
                print(f"  💡 File may be password-protected or use unsupported features")
                return None
        
        # Handle JSON files
        elif file_ext == '.json':
            return load_json_campaign(campaign_path)
        
        # Handle text-based files
        elif file_ext in ['.txt', '.html', '.md']:
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(campaign_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return content
                except UnicodeDecodeError:
                    continue
            
            # Fallback
            with open(campaign_path, 'rb') as f:
                raw_content = f.read()
                return raw_content.decode('utf-8', errors='ignore')
        
        return None
        
    except Exception as e:
        print(f"  ❌ Error loading campaign content from {campaign_path}: {str(e)}")
        import traceback as tb
        tb.print_exc()
        return None


# PATCH INSTRUCTIONS:
# 1. Add the is_valid_docx function before load_campaign_content
# 2. Replace your existing load_campaign_content function with this version
# 3. The enhanced version will:
#    - Check file exists and is not empty
#    - Validate DOCX ZIP structure
#    - Verify required XML files are present
#    - Provide helpful error messages
#    - Continue processing other files if one fails

def extract_subject_from_content(content):
    """Extract subject line from content if present"""
    try:
        if isinstance(content, dict):
            return content.get('subject', 'Campaign')
            
        lines = str(content).split('\n')
        for line in lines[:10]:
            if line.lower().startswith('subject:'):
                return line.split(':', 1)[1].strip()
            elif line.lower().startswith('# '):
                return line[2:].strip()
        return None
    except:
        return None


def scan_domain_campaigns(templates_dir, specific_file=None):
    """Scan for campaigns in domain-based directory structure or process a specific file"""
    templates_path = Path(templates_dir)
    domain_campaigns = {}
    
    # Handle specific file mode
    if specific_file:
        specific_path = Path(specific_file)
        
        # Check if file exists
        if not specific_path.exists():
            print(f"ERROR: Specified template file not found: {specific_file}")
            return domain_campaigns
        
        # Validate file extension
        if specific_path.suffix.lower() not in ['.docx', '.txt', '.html', '.md', '.json']:
            print(f"ERROR: Unsupported file format: {specific_path.suffix}")
            return domain_campaigns
        
        # Determine domain from path structure
        try:
            # Try to extract domain from path (e.g., campaign-templates/education/template.docx)
            relative_to_templates = specific_path.relative_to(templates_path)
            if len(relative_to_templates.parts) > 1:
                domain_name = relative_to_templates.parts[0]
            else:
                domain_name = 'default'
        except ValueError:
            # File is not inside templates_dir, use 'default' domain
            domain_name = 'default'
        
        domain_campaigns[domain_name] = [specific_path]
        print(f"Processing specific template file: {specific_file} (domain: {domain_name})")
        return domain_campaigns
    
    # Original directory scanning logic
    if not templates_path.exists():
        print(f"Templates directory not found: {templates_dir}")
        return domain_campaigns
    
    has_subdomains = False
    for domain_dir in templates_path.iterdir():
        if domain_dir.is_dir() and not domain_dir.name.startswith('.'):
            domain_name = domain_dir.name
            campaigns = []
            
            for ext in ['.docx', '.txt', '.html', '.md', '.json']:
                campaigns.extend(list(domain_dir.rglob(f"*{ext}")))
            
            if campaigns:
                has_subdomains = True
                domain_campaigns[domain_name] = campaigns
                print(f"Found {len(campaigns)} campaign(s) in {domain_name}/ (including subdirectories)")
                
                subdirs = set()
                for campaign in campaigns:
                    relative_path = campaign.relative_to(domain_dir)
                    if len(relative_path.parts) > 1:
                        subdirs.add(relative_path.parts[0])
                
                if subdirs:
                    print(f"  Subdirectories in {domain_name}/: {', '.join(sorted(subdirs))}")
    
    if not has_subdomains:
        print("No domain subdirectories found, using flat directory structure")
        campaigns = []
        for ext in ['.docx', '.txt', '.html', '.md', '.json']:
            campaigns.extend(list(templates_path.glob(f"*{ext}")))
        
        if campaigns:
            domain_campaigns['default'] = campaigns
            print(f"Found {len(campaigns)} campaign(s) in flat structure (using 'default' domain)")
    
    return domain_campaigns


def save_tracking_data(tracking_dir, domain, tracking_id, campaign_data):
    """Save tracking data to JSON file"""
    domain_tracking = Path(tracking_dir) / domain / "campaigns"
    domain_tracking.mkdir(parents=True, exist_ok=True)
    
    tracking_file = domain_tracking / f"{tracking_id}.json"
    
    try:
        with open(tracking_file, 'w') as f:
            json.dump(campaign_data, f, indent=2)
        print(f"Tracking data saved: {tracking_file}")
    except Exception as e:
        print(f"Warning: Could not save tracking data: {e}")


def send_summary_alert(emailer, campaigns_count, sent_count, failed_count, campaign_results):
    """Send summary alert email"""
    try:
        total_emails = sent_count + failed_count
        success_rate = (sent_count / max(1, total_emails)) * 100
        subject = f"Campaign Summary: {campaigns_count} campaigns, {total_emails} emails"
        
        body = f"""
Campaign Execution Summary
=========================

Campaigns Processed: {campaigns_count}
Total Emails: {total_emails}
Successful: {sent_count}
Failed: {failed_count}
Success Rate: {success_rate:.1f}%
Template Processing: ENABLED

Campaign Details:
"""
        
        for result in campaign_results:
            body += f"\n- {result['campaign_name']}: {result.get('sent', 0)}/{result['total_recipients']} sent"
            if result.get('tracking_id'):
                body += f" [ID: {result['tracking_id']}]"
            if result['failed'] > 0:
                body += f" ({result['failed']} failed)"
            if result.get('template_substitution'):
                body += " [Personalized]"
        
        body += f"\n\nExecution completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        emailer.send_alert(subject, body)
        print("Summary alert sent")
        
    except Exception as e:
        print(f"Warning: Could not send summary alert: {e}")


def archive_used_contacts(contacts_path, archive_root="contact_details_used"):
    """
    Archive processed contact files or directories
    
    Args:
        contacts_path: Path to contact file/directory that was processed
        archive_root: Root directory for archived contacts
    """
    import shutil
    from pathlib import Path
    from datetime import datetime
    
    try:
        source = Path(contacts_path)
        
        if not source.exists():
            print(f"  ⚠️ Source not found: {contacts_path}")
            return False
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        archive_base = Path(archive_root)
        
        # Handle both files and directories
        if source.is_file():
            # Archive single file
            relative_path = source.relative_to(Path.cwd()) if source.is_absolute() else source
            
            # Preserve directory structure
            target_dir = archive_base / relative_path.parent
            target_dir.mkdir(parents=True, exist_ok=True)
            
            archived_name = f"{source.stem}_used_{timestamp}{source.suffix}"
            target_file = target_dir / archived_name
            
            shutil.move(str(source), str(target_file))
            print(f"  ✅ Archived file: {source.name}")
            print(f"     → {target_file}")
            
        elif source.is_dir():
            # Archive entire directory
            target_dir = archive_base / f"{source.name}_used_{timestamp}"
            
            shutil.move(str(source), str(target_dir))
            print(f"  ✅ Archived directory: {source}")
            print(f"     → {target_dir}")
            
            # Create empty directory to replace archived one (optional)
            source.mkdir(parents=True, exist_ok=True)
            print(f"  📁 Created empty replacement: {source}")
        
        return True
        
    except Exception as e:
        print(f"  ❌ Error archiving: {e}")
        import traceback
        traceback.print_exc()
        return False


def load_json_campaign(campaign_path):
    """Load and process JSON campaign file - supports both content and config formats"""
    try:
        with open(campaign_path, 'r', encoding='utf-8') as f:
            campaign_data = json.load(f)
        
        print(f"📄 Loaded JSON campaign: {campaign_path}")
        
        # ===== VALIDATE SCHEDULE FIRST =====
        is_valid, reason = validate_campaign_schedule(campaign_data)
        
        if not is_valid:
            print(f"  ⏭️  SKIPPING: {reason}")
            return None
        else:
            print(f"  ✅ {reason}")
        
        # Format 1: Simple content format with subject/content
        if 'subject' in campaign_data and 'content' in campaign_data:
            return {
                'subject': campaign_data.get('subject'),
                'content': campaign_data['content'],
                'from_name': campaign_data.get('from_name', 'Campaign System'),
                'content_type': campaign_data.get('content_type', 'html'),
                'metadata': campaign_data.get('metadata', {}),
                'config': campaign_data  # Include full config
            }
        
        # Format 2: Config format with templates array
        elif 'templates' in campaign_data and isinstance(campaign_data['templates'], list):
            if not campaign_data['templates']:
                print(f"  ⚠️  No templates specified in config")
                return None
            
            # Get the first template file
            template_file = campaign_data['templates'][0]
            template_path = Path(template_file)
            
            # If path is relative, resolve it
            if not template_path.is_absolute():
                template_path = Path(campaign_path).parent.parent / template_file
            
            if not template_path.exists():
                print(f"  ⚠️  Template file not found: {template_path}")
                return None
            
            # Load the actual template content
            template_content = load_campaign_content(str(template_path))
            if not template_content:
                print(f"  ⚠️  Could not load template content from {template_path}")
                return None
            
            # Build campaign data from config
            return {
                'subject': campaign_data.get('subject', 'Campaign'),
                'content': template_content if isinstance(template_content, str) else template_content.get('content', ''),
                'from_name': campaign_data.get('from_name', 'Campaign System'),
                'content_type': 'html',
                'config': campaign_data,
                'template_source': str(template_path)
            }
        
        # Format 3: campaigns array
        elif 'campaigns' in campaign_data:
            if campaign_data['campaigns']:
                first_campaign = campaign_data['campaigns'][0]
                return {
                    'subject': first_campaign.get('subject', 'Campaign'),
                    'content': first_campaign.get('content', ''),
                    'from_name': first_campaign.get('from_name', 'Campaign System'),
                    'content_type': first_campaign.get('content_type', 'html'),
                    'metadata': campaign_data.get('metadata', {}),
                    'config': campaign_data
                }
        
        print(f"  ⚠️  Unknown JSON campaign format")
        return None
        
    except Exception as e:
        print(f"  ❌ Error loading JSON campaign: {e}")
        traceback.print_exc()
        return None


    
def campaign_main(contacts_root, scheduled_root, tracking_root, alerts_email, 
                  dry_run=False, queue_emails=False, specific_template=None, 
                  feedback_email=None, target_domain=None, campaign_filter=None, 
                  debug=False, compliance_mode=False, daily_limit=0, 
                  per_domain_limit=0, suppression_file=None, 
                  batch_size=50, delay=5, **kwargs):
    """
    Main campaign execution function with compliance support and rate limiting.
    
    Args:
        contacts_root: Path to contacts directory
        scheduled_root: Path to scheduled campaigns/templates directory
        tracking_root: Path to tracking directory
        alerts_email: Email for system alerts
        dry_run: If True, preview without sending
        queue_emails: If True, queue emails for batch sending
        specific_template: Path to specific template file to process
        feedback_email: Email for feedback responses
        target_domain: Filter to specific domain
        campaign_filter: Filter campaigns by pattern
        debug: Enable debug logging
        compliance_mode: Enable compliance checks and rate limiting
        daily_limit: Maximum emails per day (0 = no limit)
        per_domain_limit: Maximum emails per domain (0 = no limit)
        suppression_file: Path to suppression list JSON
        batch_size: Batch size for processing
        delay: Delay between batches in seconds
        **kwargs: Additional arguments for future expansion
    """
    campaigns_processed = 0
    total_emails_sent = 0
    total_emails_queued = 0
    total_failures = 0
    campaign_results = []

    try:
        # ===== INITIALIZATION & STARTUP LOGGING =====
        print(f"Starting domain-aware campaign system")
        print(f"GitHub Actions detected: {os.getenv('GITHUB_ACTIONS') is not None}")
        print(f"Queue mode: {queue_emails}")
        print(f"Dry run mode: {dry_run}")
        print(f"Compliance mode: {compliance_mode}")
        print(f"Contacts: {contacts_root}")
        print(f"Templates (domain-based): {scheduled_root}")
        
        if specific_template:
            print(f"Specific template file: {specific_template}")
        
        print(f"Tracking: {tracking_root}")
        print(f"Alerts: {alerts_email}")
        
        os.makedirs(tracking_root, exist_ok=True)
        
        # ===== COMPLIANCE & RATE LIMITING SETUP =====
        suppression_list = set()
        rate_limit_file = Path(tracking_root) / 'rate_limits.json'
        rate_data = {
            'daily_sent': 0,
            'last_reset': datetime.now().date().isoformat(),
            'domain_counts': {}
        }
        
        if compliance_mode:
            print(f"\n🔒 Compliance Mode Active:")
            print(f"   Daily limit: {daily_limit if daily_limit > 0 else 'No limit'}")
            print(f"   Per-domain limit: {per_domain_limit if per_domain_limit > 0 else 'No limit'}")
            
            # Check if rate limit file needs reset (new day)
            if rate_limit_file.exists():
                try:
                    with open(rate_limit_file, 'r') as f:
                        existing_data = json.load(f)
                        last_reset = existing_data.get('last_reset')
                        today = datetime.now().date().isoformat()
                        
                        if last_reset != today:
                            # New day, reset counters
                            rate_data = {
                                'daily_sent': 0,
                                'last_reset': today,
                                'domain_counts': {}
                            }
                            print(f"   ↻ Rate limits reset for new day")
                        else:
                            # Same day, load existing data
                            rate_data = existing_data
                except Exception as e:
                    print(f"   ⚠️  Warning: Could not load rate limits: {e}")
            
            print(f"   Current daily sent: {rate_data.get('daily_sent', 0)}/{daily_limit if daily_limit > 0 else '∞'}")
            
            # Load suppression list
            if suppression_file:
                suppression_path = Path(suppression_file)
                if suppression_path.exists():
                    try:
                        with open(suppression_path, 'r') as f:
                            suppression_data = json.load(f)
                            suppression_list = set(suppression_data.get('suppressed_emails', []))
                        print(f"   ✅ Loaded {len(suppression_list)} suppressed emails")
                    except Exception as e:
                        print(f"   ⚠️  Warning: Could not load suppression file: {e}")
                else:
                    print(f"   ℹ️  Suppression file will be created: {suppression_file}")
                    suppression_path.parent.mkdir(parents=True, exist_ok=True)
                    with open(suppression_path, 'w') as f:
                        json.dump({
                            'suppressed_emails': [],
                            'last_updated': datetime.now().isoformat(),
                            'count': 0
                        }, f, indent=2)
            
            print()
        
        # ===== LOAD CONTACTS =====
        print("Loading contacts for validation...\n")
        if DATA_LOADER_AVAILABLE:
            all_contacts = load_contacts_directory(contacts_root)
            if all_contacts:
                stats, all_contacts = validate_contact_data(all_contacts)
                original_contact_count = len(all_contacts)
                print(f"✅ Loaded {original_contact_count} contacts globally for validation")
            else:
                print("⚠️  No contacts loaded globally")
                original_contact_count = 0
                all_contacts = []
        else:
            all_contacts = fallback_load_contacts_from_directory(contacts_root)
            original_contact_count = len(all_contacts)
            print(f"✅ Loaded {original_contact_count} contacts (fallback loader)")
         # ===== APPLY COMPLIANCE FILTERS =====
        print("⚠️  Global compliance filtering disabled - applied per-campaign\n")      
        # ===== INITIALIZE EMAIL SYSTEM =====
        smtp_host = os.getenv('SMTP_HOST') or os.getenv('SMTP_SERVER')
        smtp_port = os.getenv('SMTP_PORT')
        smtp_user = os.getenv('SMTP_USER') or os.getenv('SMTP_USERNAME')
        smtp_pass = os.getenv('SMTP_PASS') or os.getenv('SMTP_PASSWORD')
        
        # Initialize unsubscribe system
        unsubscribe_manager = UnsubscribeManager(
            tracking_dir=tracking_root,
            base_url="https://sednabcn.github.io/unsubscribe"
        )
        print(f"✅ Unsubscribe system initialized: {unsubscribe_manager.get_stats()}\n")
        
        # Initialize emailer
        if queue_emails:
            print("Initializing email queue mode")
            emailer = EmailSender(
                smtp_host=None,
                smtp_port=None,
                smtp_user=None,
                smtp_password=None,
                alerts_email=alerts_email,
                dry_run=False,
                queue_emails=True
            )
            emailer.unsubscribe_manager = unsubscribe_manager
                
        elif GITHUB_ACTIONS_EMAIL_AVAILABLE and os.getenv('GITHUB_ACTIONS'):
            print("Using GitHubActionsEmailSender - SMTP timeouts bypassed")
            emailer = GitHubActionsEmailSender(
                smtp_host=smtp_host,
                smtp_port=smtp_port,
                smtp_user=smtp_user,
                smtp_pass=smtp_pass,
                alerts_email=alerts_email,
                dry_run=dry_run
            )
        else:
            print("Using standard EmailSender")
            emailer = EmailSender(
                smtp_host=smtp_host,
                smtp_port=smtp_port,
                smtp_user=smtp_user,
                smtp_password=smtp_pass,
                alerts_email=alerts_email,
                dry_run=dry_run
            )

        # ===== SCAN & LOAD CAMPAIGNS WITH SCHEDULE FILTERING =====
        domain_campaigns = scan_domain_campaigns(scheduled_root, specific_file=specific_template)
        
        if not domain_campaigns:
            print(f"ERROR: No campaigns found")
            return
        
        # ===== ORGANIZE CAMPAIGNS BY SCHEDULE PRIORITY =====
        scheduled_campaigns = {
            'immediate': [],      # mode=immediate, date=today
            'schedule_now': [],   # mode=schedule_now, date=today
            'scheduled': []       # mode=schedule, date=today
        }
        
        skipped_campaigns = []

        # Track processed campaigns to avoid duplicates
        processed_campaign_files = set()

        
        print(f"\n{'='*70}")
        print(f"SCANNING AND VALIDATING CAMPAIGNS")
        print(f"{'='*70}\n")
        
        for domain, campaign_files in domain_campaigns.items():
            print(f"Domain: {domain.upper()}")
            
            for campaign_file in campaign_files:
                campaign_name = campaign_file.stem
                
                # Load and validate
                campaign_content = load_campaign_content(str(campaign_file))
                
                if not campaign_content:
                    print(f"  ⏭️  Skipped: {campaign_name} (failed to load)")
                    skipped_campaigns.append((domain, campaign_name, "Failed to load"))
                    continue
                
                # Extract config
                if isinstance(campaign_content, dict):
                    config = campaign_content.get('config', {})
                else:
                    config = {}
                
                # Validate schedule
                is_valid, reason = validate_campaign_schedule(config)
                
                if not is_valid:
                    print(f"  ⏭️  Skipped: {campaign_name} - {reason}")
                    skipped_campaigns.append((domain, campaign_name, reason))
                    continue
                
                # Categorize by mode
                mode = config.get('mode', 'immediate')
                
                campaign_info = {
                    'domain': domain,
                    'file': campaign_file,
                    'name': campaign_name,
                    'content': campaign_content,
                    'config': config,
                    'sector': config.get('sector', domain)
                }
                
                if mode == 'immediate':
                    scheduled_campaigns['immediate'].append(campaign_info)
                    print(f"  ✅ Queued (IMMEDIATE): {campaign_name}")
                elif mode == 'schedule_now':
                    scheduled_campaigns['schedule_now'].append(campaign_info)
                    print(f"  ✅ Queued (SCHEDULE NOW): {campaign_name}")
                elif mode == 'schedule':
                    scheduled_campaigns['scheduled'].append(campaign_info)
                    print(f"  ✅ Queued (SCHEDULED): {campaign_name}")
            
            print()
        
        # Summary
        total_queued = (len(scheduled_campaigns['immediate']) + 
                       len(scheduled_campaigns['schedule_now']) + 
                       len(scheduled_campaigns['scheduled']))
        
        print(f"{'='*70}")
        print(f"CAMPAIGN QUEUE SUMMARY")
        print(f"{'='*70}")
        print(f"Immediate:     {len(scheduled_campaigns['immediate'])} campaigns")
        print(f"Schedule Now:  {len(scheduled_campaigns['schedule_now'])} campaigns")
        print(f"Scheduled:     {len(scheduled_campaigns['scheduled'])} campaigns")
        print(f"Total Queued:  {total_queued} campaigns")
        print(f"Skipped:       {len(skipped_campaigns)} campaigns")
        print(f"{'='*70}\n")
        
        if skipped_campaigns:
            print("Skipped Campaigns:")
            for domain, name, reason in skipped_campaigns:
                print(f"  - {domain}/{name}: {reason}")
            print()
        
        if total_queued == 0:
            print("No valid campaigns to process")
            return
                
        if not domain_campaigns:
            print(f"ERROR: No campaigns found in {scheduled_root}")
            print("Supported formats: .docx, .txt, .html, .md, .json")
            print("Structures supported:")
            print("  1. Domain-based: {scheduled_root}/{domain}/*.docx")
            print("  2. Flat: {scheduled_root}/*.docx")
            print("  3. Specific file: --template-file path/to/file.docx")
            return
        # ===== PROCESS CAMPAIGNS IN PRIORITY ORDER =====
        processing_order = ['immediate', 'schedule_now', 'scheduled']
             
        for priority in processing_order:
            campaigns = scheduled_campaigns[priority]
    
            if not campaigns:
                continue
    
            print(f"\n{'='*70}")
            print(f"PROCESSING {priority.upper()} CAMPAIGNS ({len(campaigns)})")
            print(f"{'='*70}\n")

            for campaign_info in campaigns:
                domain = campaign_info['domain']
                campaign_name = campaign_info['name']
                campaign_content = campaign_info['content']
                config = campaign_info['config']
                campaign_file = campaign_info['file']
    
                # Track this campaign as processed
                processed_campaign_files.add(str(campaign_file))  # ← ADD THIS LINE

                print(f"\n--- Campaign: {domain}/{campaign_name} ---")

                # CREATE ISOLATION - ADD THESE LINES:
                isolation = prepare_campaign_isolation(config, campaign_file)
                print(f"  Campaign ID: {isolation['campaign_id']}")
                
                print(f"  Tracking: {isolation['tracking_dir']}")
                print(f"  Archive: {isolation['archive_dir']}")
        
                # ===== CRITICAL FIX: Load contacts specific to THIS campaign ONLY =====
                contacts_path = config.get('contacts', contacts_root)
                print(f"  📂 Loading contacts from: {contacts_path}")



                # ALWAYS reload contacts for each campaign - NEVER use global all_contacts
                if DATA_LOADER_AVAILABLE:
                    campaign_contacts = load_contacts_directory(contacts_path)
                    if campaign_contacts:
                        stats, campaign_contacts = validate_contact_data(campaign_contacts)
                        print(f"    ✅ Loaded: {len(campaign_contacts)} valid contacts")
                    else:
                        print(f"    ⚠️ No contacts loaded from {contacts_path}")
                        campaign_contacts = []
                else:
                    campaign_contacts = fallback_load_contacts_from_directory(contacts_path)
                    if campaign_contacts:
                        print(f"    ✅ Loaded: {len(campaign_contacts)} contacts")
                    else:
                        print(f"    ⚠️ No contacts loaded from {contacts_path}")
                        campaign_contacts = []

                # ===== VALIDATION: Ensure we have campaign-specific contacts =====
                if not campaign_contacts:
                    print(f"  ⚠️ No contacts loaded from {contacts_path} - SKIPPING")
                    skipped_campaigns.append((domain, campaign_name, "No contacts found"))
                    continue
                
               
                # ===== VALIDATION: Ensure we have campaign-specific contacts =====
                if not campaign_contacts:
                    print(f"  ⚠️  No contacts loaded from {contacts_path} - SKIPPING")
                    continue
        
                # Verify contacts are from expected subdirectory
                expected_subdir = os.path.basename(contacts_path)
                print(f"  📍 Expected subdirectory: {expected_subdir}")
                print(f"  📧 Campaign-specific contacts: {len(campaign_contacts)}")

                
                # Debug: Show sample contact
                if campaign_contacts and len(campaign_contacts) > 0:
                    sample = campaign_contacts[0]
                    print(f"  📋 Sample contact: {sample.get('email', 'NO EMAIL')} - {sample.get('name', 'NO NAME')}")
                    print(f"  📋 Contact fields: {list(sample.keys())}")
                
                # ===== Apply compliance filters FOR THIS CAMPAIGN ONLY =====
                original_count = len(campaign_contacts)
        
                if compliance_mode:
                    print(f"  🔒 Applying compliance filters...")
            
                    # Filter suppressed
                    if suppression_list:
                        before = len(campaign_contacts)
                        campaign_contacts = [c for c in campaign_contacts 
                                             if c.get('email', '').lower() not in suppression_list]
                        suppressed_count = before - len(campaign_contacts)
                        if suppressed_count > 0:
                            print(f"     - Filtered {suppressed_count} suppressed contacts")
            
                    # Check per-domain limit
                    if per_domain_limit > 0:
                        domain_count = rate_data['domain_counts'].get(domain, 0)
                        available = per_domain_limit - domain_count
                
                        if available <= 0:
                            print(f"  ⚠️  Per-domain limit reached for {domain} ({per_domain_limit})")
                            continue
                        elif len(campaign_contacts) > available:
                            print(f"     - Limiting to {available} contacts (per-domain limit)")
                            campaign_contacts = campaign_contacts[:available]
            
                    # Check daily limit
                    if daily_limit > 0:
                        current_daily = rate_data.get('daily_sent', 0)
                        remaining_daily = daily_limit - current_daily
                
                        if remaining_daily <= 0:
                            print(f"  ⚠️  Daily limit reached ({daily_limit})")
                            break  # Stop processing ALL campaigns
                        elif len(campaign_contacts) > remaining_daily:
                            print(f"     - Limiting to {remaining_daily} contacts (daily limit)")
                            campaign_contacts = campaign_contacts[:remaining_daily]
            
                    filtered_count = original_count - len(campaign_contacts)
                    if filtered_count > 0:
                        print(f"  🔒 Total filtered: {filtered_count} contacts")
        
                # Final validation
                if not campaign_contacts:
                    print(f"  ⚠️  No eligible contacts after filtering - SKIPPING")
                    continue
        
                # Extract campaign details
                if isinstance(campaign_content, dict):
                    subject = campaign_content.get('subject', f"Campaign: {campaign_name}")
                    content = campaign_content.get('content', '')
                    from_name = campaign_content.get('from_name', 'Campaign System')
                else:
                    subject = extract_subject_from_content(campaign_content) or f"Campaign: {campaign_name}"
                    content = str(campaign_content)
                    from_name = "Campaign System"
        
                # Generate tracking ID
                tracking_id = generate_tracking_id(domain, campaign_name, campaign_file.name)
        
                # ===== CRITICAL: Use ONLY campaign_contacts, NEVER all_contacts =====
               
                contacts_with_ids = []
                skipped_unsub = 0

                print(f"  📄 Processing {len(campaign_contacts)} campaign-specific contacts...")

                if not campaign_contacts:
                    print(f"  ⚠️ ERROR: campaign_contacts is empty before processing!")
                    continue
    
                for i, contact in enumerate(campaign_contacts):  # <- CRITICAL: campaign_contacts
                    email = contact.get('email', '').strip()
                    
                    if not email:
                        print(f"    ⚠️ Skipping contact {i+1}: no email field")
                        continue
    
                    if '@' not in email:
                        print(f"    ⚠️ Skipping contact {i+1}: invalid email '{email}'")
                        continue
    
                    # Check unsubscribe (only in live mode, not dry-run)
                    if not dry_run and unsubscribe_manager.is_unsubscribed(email, campaign_name):
                        skipped_unsub += 1
                        continue
    
                    contact_copy = contact.copy()
                    contact_copy['recipient_id'] = f"{isolation['campaign_id']}_{i+1}"
                    contact_copy['campaign_id'] = campaign_name
                    contact_copy['domain'] = domain
                    contact_copy['tracking_id'] = tracking_id
                    contact_copy['unsubscribe_link'] = unsubscribe_manager.generate_unsubscribe_link(
                        email, campaign_name
                    )
    
                    contacts_with_ids.append(contact_copy)

                if skipped_unsub > 0:
                    print(f"  🔕 Skipped {skipped_unsub} unsubscribed contacts")

                if not contacts_with_ids:
                        print(f"  ⚠️ No eligible contacts after all filtering - SKIPPING")
                        print(f"     Original: {len(campaign_contacts)}, After filtering: {len(contacts_with_ids)}")
                        skipped_campaigns.append((domain, campaign_name, "All contacts filtered out"))
                        continue

                print(f"  ✅ Ready to send: {len(contacts_with_ids)} contacts")
                print(f"  📁 Source: {contacts_path}")
                
                # Send campaign
                try:
                        campaign_result = emailer.send_campaign(
                            campaign_name=f"{domain}/{campaign_name}",
                            subject=subject,
                            content=content,
                            recipients=contacts_with_ids,
                            from_name=from_name,
                            tracking_id=tracking_id,
                            contact_mapping=config.get('contact_mapping', {})
                        )
            
                        campaigns_processed += 1
                        sent_count = campaign_result.get('sent', 0)
                        queued_count = campaign_result.get('queued', 0)
                        
                        total_emails_sent += sent_count
                        total_emails_queued += queued_count
                        total_failures += campaign_result['failed']
            
                        campaign_results.append(campaign_result)
            
                        # Update rate limits
                        if compliance_mode:
                            rate_data['daily_sent'] = rate_data.get('daily_sent', 0) + sent_count
                            rate_data['domain_counts'][domain] = rate_data['domain_counts'].get(domain, 0) + sent_count
                            
                        # Save tracking data with SOURCE PATH
                        tracking_data = {
                            'tracking_id': tracking_id,
                            'campaign_id': isolation['campaign_id'],
                            'domain': domain,
                            'campaign_name': campaign_name,
                            'sector': config.get('sector', domain),
                            'mode': config.get('mode'),
                            'date': config.get('date'),
                            'subject': subject,
                            'from_name': from_name,
                            'timestamp': datetime.now().isoformat(),
                            'contacts_source': contacts_path,  # ← CRITICAL: Record source
                            'contacts_loaded': len(campaign_contacts),
                            'contacts_after_filtering': len(contacts_with_ids),
                            'total_recipients': len(contacts_with_ids),
                            'sent': sent_count,
                            'queued': queued_count,
                            'failed': campaign_result['failed'],
                            'isolation': {
                                'tracking_dir': str(isolation['tracking_dir']),
                                'archive_dir': str(isolation['archive_dir']),
                                'log_file': isolation['log_file']
                            }
                        }
            
                        # Save to isolated tracking directory
                        tracking_file = isolation['tracking_dir'] / f"{tracking_id}.json"
                        with open(tracking_file, 'w') as f:
                            json.dump(tracking_data, f, indent=2)
            
                        print(f"  ✅ Campaign complete:")
                        print(f"     - Sent: {sent_count}")
                        print(f"     - Queued: {queued_count}")
                        print(f"     - Failed: {campaign_result['failed']}")
                        print(f"     - Source verified: {contacts_path}")
            
                        # Archive contacts after successful send
                        if (sent_count > 0 or queued_count > 0) and not dry_run and not queue_emails:
                            print(f"  📦 Archiving processed contacts...")
                            archive_success = archive_used_contacts(
                                contacts_path, 
                                archive_root=str(isolation['archive_dir'])
                            )
                            if archive_success:
                                print(f"  ✅ Contacts archived to {isolation['archive_dir']}")
        
                except Exception as e:
                        print(f"  ❌ Error processing campaign: {e}")
                        import traceback as tb
                        tb.print_exc()
                        continue

        # ===== INITIALIZE TRACKING & LOGGING =====
        log_file = "dryrun.log" if dry_run else "campaign_execution.log"
        with open(log_file, 'w') as f:
            f.write("Domain-Aware Campaign Log\n")
            f.write(f"GitHub Actions mode: {os.getenv('GITHUB_ACTIONS') is not None}\n")
            f.write(f"Queue mode: {queue_emails}\n")
            f.write(f"Compliance mode: {compliance_mode}\n")
            f.write(f"Specific template: {specific_template if specific_template else 'None'}\n")
            # Note: Contacts loaded per-campaign, not globally
            f.write(f"Domains found: {len(domain_campaigns)}\n")
            f.write(f"Timestamp: {datetime.now().isoformat()}\n\n")
        
        # ===== PROCESS CAMPAIGNS BY DOMAIN (SKIP IF ALREADY PROCESSED) =
        for domain, campaign_files in domain_campaigns.items():
            # Skip campaigns already processed by priority system
            unprocessed = [f for f in campaign_files if str(f) not in processed_campaign_files]
    
            if not unprocessed:
                print(f"ℹ️  Domain {domain}: All campaigns already processed")
                continue
    
            print(f"\n{'='*70}")
            print(f"PROCESSING DOMAIN: {domain.upper()}")
            print(f"{'='*70}")
                
            # Create domain tracking structure
            domain_tracking = Path(tracking_root) / domain
            (domain_tracking / "campaigns").mkdir(parents=True, exist_ok=True)
            (domain_tracking / "responses").mkdir(parents=True, exist_ok=True)
            (domain_tracking / "analytics").mkdir(parents=True, exist_ok=True)
            
            domain_emails_sent = 0
            
            # Process each campaign
            for campaign_file in unprocessed:
                campaign_name = campaign_file.stem
                campaign_path = str(campaign_file)
                
                # Determine subdirectory structure
                relative_path = campaign_file.relative_to(Path(scheduled_root) / domain) if domain != 'default' else campaign_file.relative_to(Path(scheduled_root))
                subdirectory = relative_path.parent.name if len(relative_path.parts) > 1 else None
                full_campaign_name = f"{subdirectory}/{campaign_name}" if subdirectory else campaign_name
                
                print(f"\n--- Processing Campaign: {domain}/{full_campaign_name} ---")
                
                # Load campaign content
                campaign_content = load_campaign_content(campaign_path)
                if not campaign_content:
                    print(f"Warning: Could not load content for {campaign_file.name}")
                    continue
                
                # Extract campaign details
                if isinstance(campaign_content, dict):
                    subject = campaign_content.get('subject', f"Campaign: {campaign_name}")
                    content = campaign_content.get('content', '')
                    from_name = campaign_content.get('from_name', 'Campaign System')
                    config = campaign_content.get('config', {})
                    
                    if config:
                        print(f"Using config-based campaign settings from JSON")
                        from_name = config.get('from_name', from_name)
                        subject = config.get('subject', subject)
                        
                        contact_mapping = config.get('contact_mapping', {})
                        if contact_mapping:
                            print(f"Contact field mapping active: {list(contact_mapping.keys())}")
                        
                        if config.get('sector'):
                            print(f"Sector: {config['sector']}")
                        if config.get('feedback', {}).get('auto_inject'):
                            print(f"Feedback auto-injection: ENABLED")
                        if config.get('tracking', {}).get('enabled'):
                            print(f"Enhanced tracking: ENABLED")
                else:
                    subject = extract_subject_from_content(campaign_content) or f"Campaign: {campaign_name}"
                    content = str(campaign_content)
                    from_name = "Campaign System"
                    config = {}
                
                # Generate tracking ID
                tracking_id = generate_tracking_id(domain, full_campaign_name, campaign_file.name)
                
                # ===== PREPARE CONTACTS FOR CAMPAIGN =====
                contacts_with_ids = []
                skipped_unsub = 0
                skipped_per_domain = 0
                
                for i, contact in enumerate(campaign_contacts):
                    email = contact.get('email', '').strip()
                    
                    # Check unsubscribe status
                    if unsubscribe_manager.is_unsubscribed(email, full_campaign_name):
                        skipped_unsub += 1
                        continue
                    
                    # Check per-domain limit (if enabled)
                    if compliance_mode and per_domain_limit > 0:
                        domain_count = rate_data['domain_counts'].get(domain, 0)
                        if domain_count >= per_domain_limit:
                            skipped_per_domain += 1
                            continue
                    
                    # Prepare contact record
                    contact_copy = contact.copy()
                    contact_copy['recipient_id'] = f"{domain}_{full_campaign_name.replace('/', '_')}_{i+1}"
                    contact_copy['campaign_id'] = full_campaign_name
                    contact_copy['domain'] = domain
                    contact_copy['tracking_id'] = tracking_id
                    contact_copy['unsubscribe_link'] = unsubscribe_manager.generate_unsubscribe_link(
                        email, 
                        full_campaign_name
                    )
                    
                    contacts_with_ids.append(contact_copy)
                
                if skipped_unsub > 0:
                    print(f"  Filtered out {skipped_unsub} unsubscribed contacts")
                if skipped_per_domain > 0:
                    print(f"  Filtered out {skipped_per_domain} contacts (per-domain limit)")
                
                # Skip campaign if no contacts
                if not contacts_with_ids:
                    print(f"  No eligible contacts for this campaign")
                    continue
                
                # ===== SEND CAMPAIGN =====
                try:
                    campaign_result = emailer.send_campaign(
                        campaign_name=f"{domain}/{full_campaign_name}",
                        subject=subject,
                        content=content,
                        recipients=contacts_with_ids,
                        from_name=from_name,
                        tracking_id=tracking_id,
                        contact_mapping=config.get('contact_mapping', {})
                    )
                    
                    campaigns_processed += 1
                    sent_count = campaign_result.get('sent', 0)
                    queued_count = campaign_result.get('queued', 0)

                    # Archive contacts after successful send
                    if (sent_count > 0 or queued_count > 0) and not dry_run and not queue_emails:
                         # Get contacts file from config
                         if isinstance(campaign_content, dict) and campaign_content.get('config'):
                                contacts_file = campaign_content['config'].get('contacts')
                                if contacts_file:
                                    print(f"\n📦 Archiving processed contacts...")
                                    archive_used_contacts(contacts_file)
                                else:
                                    print(f"  ℹ️  No contacts file path in config - skipping archive")
                    total_emails_sent += sent_count
                    total_emails_queued += queued_count
                    total_failures += campaign_result['failed']
                    domain_emails_sent += sent_count
                    
                    campaign_results.append(campaign_result)
                    
                    # ===== UPDATE RATE LIMITS =====
                    if compliance_mode:
                        rate_data['daily_sent'] = rate_data.get('daily_sent', 0) + sent_count
                        rate_data['domain_counts'][domain] = rate_data['domain_counts'].get(domain, 0) + sent_count
                    
                    # ===== CREATE TRACKING DATA =====
                    tracking_data = {
                        'tracking_id': tracking_id,
                        'domain': domain,
                        'campaign_name': campaign_name,
                        'full_path': full_campaign_name,
                        'subdirectory': subdirectory,
                        'template_file': campaign_file.name,
                        'subject': subject,
                        'from_name': from_name,
                        'timestamp': datetime.now().isoformat(),
                        'total_recipients': campaign_result['total_recipients'],
                        'sent': sent_count,
                        'queued': queued_count,
                        'failed': campaign_result['failed'],
                        'dry_run': dry_run,
                        'queue_mode': queue_emails,
                        'specific_template': specific_template if specific_template else None,
                        'compliance_mode': compliance_mode
                    }
                    
                    # Add config metadata if available
                    if config:
                        tracking_data['config_based'] = True
                        tracking_data['sector'] = config.get('sector')
                        tracking_data['feedback_enabled'] = config.get('feedback', {}).get('auto_inject', False)
                        tracking_data['tracking_enabled'] = config.get('tracking', {}).get('enabled', False)
                        tracking_data['template_source'] = campaign_content.get('template_source')
                        if 'contact_mapping' in config:
                            tracking_data['contact_mapping'] = config['contact_mapping']
                    
                    # Save tracking data
                    save_tracking_data(tracking_root, domain, tracking_id, tracking_data)
                    
                    # Append to log
                    with open(log_file, 'a') as f:
                        f.write(f"Domain: {domain}\n")
                        f.write(f"Campaign: {full_campaign_name}\n")
                        f.write(f"Tracking ID: {tracking_id}\n")
                        f.write(f"Recipients: {campaign_result['total_recipients']}\n")
                        f.write(f"Sent: {sent_count}\n")
                        f.write(f"Queued: {queued_count}\n")
                        f.write(f"Failed: {campaign_result['failed']}\n\n")
                    
                    print(f"  ✅ Sent: {sent_count}, Queued: {queued_count}, Failed: {campaign_result['failed']}")
                    
                except Exception as e:
                    print(f"Error processing campaign '{domain}/{full_campaign_name}': {str(e)}")
                    traceback.print_exc()
                    continue
            
            if domain_emails_sent > 0:
                print(f"\nDomain {domain.upper()} total: {domain_emails_sent} emails sent")
        
        # ===== SAVE RATE LIMITS =====
        if compliance_mode:
            rate_data['last_updated'] = datetime.now().isoformat()
            with open(rate_limit_file, 'w') as f:
                json.dump(rate_data, f, indent=2)
        
        # ===== SEND SUMMARY =====
        if queue_emails and emailer.queued_count > 0:
            create_github_actions_summary(
                emailer.queue_batch_dir,
                emailer.queued_count,
                campaign_results
            )
            
            create_campaign_summary_email(
                alerts_email,
                campaigns_processed,
                emailer.queued_count,
                campaign_results
            )
            
            print(f"\n{'='*70}")
            print(f"EMAIL QUEUE SUMMARY")
            print(f"{'='*70}")
            print(f"Total emails queued: {emailer.queued_count}")
            print(f"Queue directory: {emailer.queue_batch_dir}")
            print(f"Campaigns: {campaigns_processed}")
            print(f"Domains: {len(domain_campaigns)}")
            print("Emails will be sent by GitHub Actions email sender job")
            print(f"Summary files created:")
            print("  - github_actions_email_summary.json")
            print("  - campaign_summary_email.json")
            

        elif hasattr(emailer, 'send_batch_summary') and os.getenv('GITHUB_ACTIONS') and not dry_run:
            emailer.send_batch_summary(campaigns_processed, total_emails_sent, total_failures, campaign_results)
            print("Campaign summary saved for GitHub Actions email delivery")
           
        elif not dry_run and not queue_emails and campaigns_processed > 0:
            send_summary_alert(emailer, campaigns_processed, total_emails_sent, total_failures, campaign_results)
        
        # ===== FINAL LOG ENTRY =====
        with open(log_file, 'a') as f:
            f.write("=== CAMPAIGN SUMMARY ===\n")
            f.write(f"Domains processed: {len(domain_campaigns)}\n")
            f.write(f"Campaigns processed: {campaigns_processed}\n")
            f.write(f"Total emails: {total_emails_sent + total_emails_queued + total_failures}\n")
            f.write(f"Successful: {total_emails_sent}\n")
            f.write(f"Queued: {total_emails_queued}\n")
            f.write(f"Failed: {total_failures}\n")
            f.write(f"Tracking system: DOMAIN-BASED\n")
            f.write(f"Queue mode: {queue_emails}\n")
            f.write(f"Compliance mode: {compliance_mode}\n")
            if compliance_mode:
                f.write(f"Daily limit: {daily_limit if daily_limit > 0 else 'No limit'}\n")
                f.write(f"Per-domain limit: {per_domain_limit if per_domain_limit > 0 else 'No limit'}\n")
                f.write(f"Total suppressed: {len(suppression_list)}\n")
            f.write(f"Specific template: {specific_template if specific_template else 'None'}\n")
            f.write(f"Completed: {datetime.now().isoformat()}\n")
        
        # ===== FINAL SUMMARY PRINT =====
        print(f"\n{'='*70}")
        print(f"FINAL SUMMARY")
        print(f"{'='*70}")
        print(f"Domains processed: {len(domain_campaigns)}")
        print(f"Campaigns processed: {campaigns_processed}")
        print(f"Emails sent: {total_emails_sent}")
        print(f"Emails queued: {total_emails_queued}")
        print(f"Failures: {total_failures}")
        print(f"Tracking system: DOMAIN-BASED")
        print(f"Template processing: ENABLED")
        
        if compliance_mode:
            print(f"\n🔒 COMPLIANCE SUMMARY:")
            print(f"   Daily sent: {rate_data.get('daily_sent', 0)}/{daily_limit if daily_limit > 0 else '∞'}")
            print(f"   Suppressed contacts: {len(suppression_list)}")
            print(f"   Domain distribution: {dict(rate_data.get('domain_counts', {}))}")
        
        print("\nCampaign completed successfully")
        
    except Exception as e:
        print(f"ERROR: {str(e)}")
        import traceback as tb
        tb.print_exc()
        sys.exit(1)


        
# ============================================================================
# MAIN ENTRY POINT - COMPLETE FIXED VERSION
# ============================================================================

if __name__ == "__main__":
    print("Domain-Aware Campaign System Script started")
    print(f"Remote environment: {IS_REMOTE}")
    print(f"Available modules: docx={DOCX_AVAILABLE}, email_sender={EMAIL_SENDER_AVAILABLE}, data_loader={DATA_LOADER_AVAILABLE}")
    
    # SINGLE CONSOLIDATED ARGUMENT PARSER
    parser = argparse.ArgumentParser(description='Domain-Aware Email Campaign System with Compliance')
    
    # === REQUIRED CORE ARGUMENTS ===
    parser.add_argument("--contacts", required=True, help="Contacts directory path")
    parser.add_argument("--scheduled", required=True, help="Domain-based templates directory")
    parser.add_argument("--tracking", required=True, help="Tracking directory path")
    parser.add_argument("--alerts", required=True, help="Alerts email address")
    
    # === OPTIONAL CORE ARGUMENTS ===
    parser.add_argument("--feedback", help="Feedback email address")
    parser.add_argument("--templates", help="Templates directory (alias for scheduled)")
    
    # === CAMPAIGN CONTROL ===
    parser.add_argument("--template-file", help="Specific template file to process")
    parser.add_argument("--domain", help="Process only specific domain")
    parser.add_argument("--filter-domain", help="Filter campaigns by domain pattern")
    
    # === EXECUTION MODES ===
    parser.add_argument("--dry-run", action="store_true", 
                       help="Preview emails without sending")
    parser.add_argument("--queue-emails", action="store_true", 
                       help="Queue emails for batch sending (GitHub Actions mode)")
    parser.add_argument("--debug", action="store_true", 
                       help="Enable debug logging")
    parser.add_argument("--remote-only", action="store_true", 
                       help="Force remote-only mode")
    
    # === FEATURE FLAGS ===
    parser.add_argument("--no-feedback", action="store_true", 
                       help="Skip feedback injection")
    parser.add_argument("--enhanced-mode", action="store_true", 
                       help="Enable enhanced processing")
    parser.add_argument("--template-variables", action="store_true", 
                       help="Enable template variable processing")
    parser.add_argument("--comprehensive-tracking", action="store_true", 
                       help="Enable comprehensive tracking")
    
    # === COMPLIANCE ARGUMENTS (CRITICAL - THESE WERE MISSING) ===
    parser.add_argument("--compliance", action="store_true", 
                       help="Enable compliance mode with rate limiting")
    parser.add_argument("--daily-limit", type=int, default=0, 
                       help="Daily send limit (0 = no limit)")
    parser.add_argument("--per-domain-limit", type=int, default=0, 
                       help="Per-domain send limit (0 = no limit)")
    parser.add_argument("--suppression-file", type=str, 
                       help="Path to suppression list JSON file")
    
    # === PROCESSING CONTROL ===
    parser.add_argument("--batch-size", type=int, default=50, 
                       help="Batch size for email processing")
    parser.add_argument("--delay", type=int, default=5, 
                       help="Delay between batches in seconds")
    
    print("Parsing arguments...")
    args = parser.parse_args()
    
    # === VALIDATE AND PROCESS ARGUMENTS ===
    
    # Handle remote-only mode
    if args.remote_only:
        IS_REMOTE = True
        print("Forced remote-only mode enabled")
    
    # Display compliance settings if enabled
    if args.compliance:
        print(f"\n🔒 Compliance Mode Enabled:")
        print(f"   Daily limit: {args.daily_limit}")
        print(f"   Per-domain limit: {args.per_domain_limit}")
        if args.suppression_file:
            print(f"   Suppression file: {args.suppression_file}")
            # Verify suppression file exists
            if not Path(args.suppression_file).exists():
                print(f"   ⚠️  Warning: Suppression file not found, will be created")
        print()
    
    # Handle templates alias
    scheduled_path = args.templates if args.templates else args.scheduled
    
    # Display all parsed arguments
    print(f"\nArguments parsed successfully:")
    print(f"  --contacts: {args.contacts}")
    print(f"  --scheduled: {scheduled_path}")
    print(f"  --tracking: {args.tracking}")
    print(f"  --alerts: {args.alerts}")
    print(f"  --feedback: {args.feedback}")
    print(f"  --template-file: {args.template_file}")
    print(f"  --domain: {args.domain}")
    print(f"  --dry-run: {args.dry_run}")
    print(f"  --queue-emails: {args.queue_emails}")
    print(f"  --compliance: {args.compliance}")
    print(f"  --debug: {args.debug}")
    
    if args.template_file:
        print(f"\n📄 Using specific template file: {args.template_file}")
    
    # === CALL MAIN CAMPAIGN FUNCTION ===
    print("\nCalling domain-aware campaign_main...")
    
    try:
        campaign_main(
            contacts_root=args.contacts,
            scheduled_root=scheduled_path,
            tracking_root=args.tracking,
            alerts_email=args.alerts,
            dry_run=args.dry_run,
            queue_emails=args.queue_emails,
            specific_template=args.template_file,
            feedback_email=args.feedback,
            target_domain=args.domain,
            campaign_filter=args.filter_domain,
            debug=args.debug,
            compliance_mode=args.compliance,
            daily_limit=args.daily_limit,
            per_domain_limit=args.per_domain_limit,
            suppression_file=args.suppression_file,
            batch_size=args.batch_size,
            delay=args.delay
        )
        print("\n✅ Domain-aware campaign system completed successfully")
        sys.exit(0)
        
    except Exception as e:
        print(f"\n❌ Campaign execution failed: {str(e)}")
        import traceback as tb
        tb.print_exc()
        sys.exit(1)




















        
