#!/usr/bin/env python3
"""
Campaign Template Diagnostic & Fix Tool
Analyzes template and contact data to identify substitution issues
"""

import json
import csv
import sys
from pathlib import Path
from docx import Document

def analyze_docx_template(template_path):
    """Extract placeholders from DOCX template"""
    print(f"\n{'='*70}")
    print(f"ANALYZING TEMPLATE: {template_path}")
    print(f"{'='*70}")
    
    try:
        doc = Document(template_path)
        all_text = []
        placeholders = set()
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text = para.text
            all_text.append(text)
            # Find [placeholder] style
            import re
            found = re.findall(r'\[([^\]]+)\]', text)
            placeholders.update(found)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    all_text.append(text)
                    found = re.findall(r'\[([^\]]+)\]', text)
                    placeholders.update(found)
        
        print(f"\n📄 Template contains {len(placeholders)} unique placeholders:")
        for ph in sorted(placeholders):
            print(f"   - [{ph}]")
        
        print(f"\n📝 First 500 characters of template:")
        combined_text = '\n'.join(all_text[:5])
        print(combined_text[:500])
        
        return placeholders, all_text
        
    except Exception as e:
        print(f"❌ Error reading template: {e}")
        return set(), []

def find_csv_file(config_csv_path, contacts_dir):
    """Try to find the actual CSV file"""
    print(f"\n🔍 Looking for CSV file...")
    print(f"   Config specifies: {config_csv_path}")
    
    # Check if exact path exists
    if Path(config_csv_path).exists():
        print(f"   ✓ Found at specified path")
        return config_csv_path
    
    # Look for similar files in contacts directory
    contacts_path = Path(contacts_dir)
    if contacts_path.exists():
        csv_files = list(contacts_path.glob("*.csv"))
        print(f"   ⚠️  File not found, but found {len(csv_files)} CSV files in {contacts_dir}:")
        for f in csv_files:
            print(f"      - {f.name}")
        
        # Look for adult education files
        adult_files = [f for f in csv_files if 'adult' in f.name.lower() or 'edu' in f.name.lower()]
        if adult_files:
            latest = max(adult_files, key=lambda p: p.stat().st_mtime)
            print(f"   💡 Using most recent matching file: {latest.name}")
            return str(latest)
    
    return None

def analyze_csv_contacts(csv_path):
    """Extract field names from CSV"""
    print(f"\n{'='*70}")
    print(f"ANALYZING CONTACTS: {csv_path}")
    print(f"{'='*70}")
    
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            fieldnames = reader.fieldnames
            
            print(f"\n📋 CSV contains {len(fieldnames)} fields:")
            for field in fieldnames:
                print(f"   - '{field}'")
            
            # Read first contact
            contacts = list(reader)
            print(f"\n📊 Total contacts in CSV: {len(contacts)}")
            
            if contacts:
                print(f"\n👤 Sample contact (first row):")
                for key, value in contacts[0].items():
                    display_value = value[:50] + "..." if len(value) > 50 else value
                    print(f"   {key}: {display_value}")
            
            return fieldnames, contacts
            
    except Exception as e:
        print(f"❌ Error reading CSV: {e}")
        return [], []

def analyze_mapping(config_path):
    """Analyze the contact mapping configuration"""
    print(f"\n{'='*70}")
    print(f"ANALYZING CONFIG: {config_path}")
    print(f"{'='*70}")
    
    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
        
        mapping = config.get('contact_mapping', {})
        print(f"\n🔗 Contact mapping configuration:")
        for key, value in mapping.items():
            print(f"   {key} → '{value}'")
        
        return config, mapping
        
    except Exception as e:
        print(f"❌ Error reading config: {e}")
        return {}, {}

def create_placeholder_mapping(template_placeholders, csv_fields, current_mapping):
    """Create a mapping from template placeholders to CSV fields"""
    print(f"\n{'='*70}")
    print("CREATING PLACEHOLDER-TO-CSV MAPPING")
    print(f"{'='*70}")
    
    # This is what the code needs: template placeholder → CSV field
    suggested_mapping = {}
    
    # Common mappings
    placeholder_to_csv = {
        'Recipient Name': 'name',
        'Title / Position': 'position',
        'Department / School': 'organization',
        'University / Institution': 'organization',
        'Address or Campus, if needed': 'organization',
        'Department Name': 'organization',
        'Month Year': None,  # This needs to be generated
    }
    
    print("\n📋 Suggested template placeholder → CSV field mapping:")
    for placeholder in sorted(template_placeholders):
        csv_field = placeholder_to_csv.get(placeholder)
        if csv_field and csv_field in csv_fields:
            suggested_mapping[placeholder] = csv_field
            print(f"   [{placeholder}] → CSV field '{csv_field}' ✓")
        elif csv_field:
            print(f"   [{placeholder}] → '{csv_field}' ❌ (CSV doesn't have this field)")
        else:
            print(f"   [{placeholder}] → ??? (needs manual mapping)")
    
    return suggested_mapping

def analyze_code_issue(template_placeholders, current_mapping):
    """Identify the code issue"""
    print(f"\n{'='*70}")
    print("ROOT CAUSE ANALYSIS")
    print(f"{'='*70}")
    
    print("\n🔍 THE PROBLEM:")
    print("   1. Template has placeholders like: [Recipient Name]")
    print("   2. Config mapping has keys like: 'name', 'email', 'position'")
    print("   3. These DON'T MATCH! The code is looking for [name], [email], [position]")
    print("      but the template has [Recipient Name], [Title / Position], etc.")
    
    print("\n💡 WHAT'S HAPPENING:")
    print("   The code is doing something like:")
    print("      content.replace('[name]', contact['name'])")
    print("      content.replace('[email]', contact['email'])")
    print("      content.replace('[position]', contact['position'])")
    print()
    print("   But the template needs:")
    print("      content.replace('[Recipient Name]', contact['name'])")
    print("      content.replace('[Title / Position]', contact['position'])")
    print("      content.replace('[Department / School]', contact['organization'])")

def suggest_fixes(template_placeholders, csv_fields, mapping):
    """Suggest fixes for identified issues"""
    print(f"\n{'='*70}")
    print("SOLUTION OPTIONS")
    print(f"{'='*70}")
    
    print("\n✅ OPTION 1: Fix the Config Mapping (RECOMMENDED)")
    print("   Update test_campaign.json contact_mapping to match template placeholders:")
    print()
    print('   "contact_mapping": {')
    print('     "Recipient Name": "name",')
    print('     "Title / Position": "position",')
    print('     "Department / School": "organization",')
    print('     "University / Institution": "organization",')
    print('     "Address or Campus, if needed": "organization"')
    print('   }')
    
    print("\n✅ OPTION 2: Fix the Code")
    print("   In docx_parser.py, change the substitution logic from:")
    print("      content.replace(f'[{mapping_key}]', value)")
    print("   To:")
    print("      # Use mapping_key as the template placeholder name directly")
    print("      content.replace(f'[{mapping_key}]', value)")
    
    print("\n✅ OPTION 3: Update the Template")
    print("   Change template placeholders to match mapping keys:")
    print("      [Recipient Name] → [name]")
    print("      [Title / Position] → [position]")
    print("      [Department / School] → [organization]")
    
    print("\n🎯 RECOMMENDED ACTION:")
    print("   Use OPTION 1 - it requires the least changes and is most maintainable")

def generate_fixed_config(template_placeholders, csv_fields):
    """Generate a corrected config file"""
    print(f"\n{'='*70}")
    print("GENERATING FIXED CONFIG")
    print(f"{'='*70}")
    
    mapping = {
        "Recipient Name": "name",
        "Title / Position": "position", 
        "Department / School": "organization",
        "University / Institution": "organization",
        "Address or Campus, if needed": "organization",
        "Department Name": "organization",
    }
    
    # Only include mappings where CSV field exists
    valid_mapping = {k: v for k, v in mapping.items() if v in csv_fields}
    
    print("\n📝 Corrected contact_mapping:")
    print(json.dumps(valid_mapping, indent=2))
    
    return valid_mapping

if __name__ == "__main__":
    print("\n" + "="*70)
    print("CAMPAIGN TEMPLATE DIAGNOSTIC TOOL")
    print("="*70)
    
    # Paths
    config_path = "scheduled-campaigns/test_campaign.json"
    contacts_dir = "contacts"
    
    # 1. Analyze config
    config, mapping = analyze_mapping(config_path)
    
    if not config:
        print("\n❌ Cannot proceed without config file")
        sys.exit(1)
    
    # 2. Get template path from config
    templates = config.get('templates', [])
    if templates:
        template_path = templates[0]
        template_placeholders, _ = analyze_docx_template(template_path)
    else:
        print("\n❌ No templates found in config")
        template_placeholders = set()
    
    # 3. Find and analyze contacts CSV
    config_csv_path = config.get('contacts', '')
    actual_csv_path = find_csv_file(config_csv_path, contacts_dir)
    
    if actual_csv_path:
        csv_fields, contacts = analyze_csv_contacts(actual_csv_path)
    else:
        print("\n❌ Cannot find contacts CSV file")
        csv_fields = []
    
    # 4. Analyze the mismatch
    if template_placeholders and csv_fields:
        analyze_code_issue(template_placeholders, mapping)
        create_placeholder_mapping(template_placeholders, csv_fields, mapping)
    
    # 5. Suggest fixes
    suggest_fixes(template_placeholders, csv_fields, mapping)
    
    # 6. Generate fixed config
    if csv_fields:
        fixed_mapping = generate_fixed_config(template_placeholders, csv_fields)
    
    print("\n" + "="*70)
    print("DIAGNOSTIC COMPLETE")
    print("="*70)
