#!/usr/bin/env python3
"""
Campaign Template Diagnostic Tool
Analyzes template and contact data to identify substitution issues
"""

import json
import csv
import sys
from pathlib import Path
from docx import Document

def analyze_docx_template(template_path):
    """Extract placeholders from DOCX template"""
    print(f"\n{'='*70}")
    print(f"ANALYZING TEMPLATE: {template_path}")
    print(f"{'='*70}")
    
    try:
        doc = Document(template_path)
        all_text = []
        placeholders = set()
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text = para.text
            all_text.append(text)
            # Find [placeholder] style
            import re
            found = re.findall(r'\[([^\]]+)\]', text)
            placeholders.update(found)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    all_text.append(text)
                    found = re.findall(r'\[([^\]]+)\]', text)
                    placeholders.update(found)
        
        print(f"\n📄 Template contains {len(placeholders)} unique placeholders:")
        for ph in sorted(placeholders):
            print(f"   - [{ph}]")
        
        print(f"\n📝 First 500 characters of template:")
        combined_text = '\n'.join(all_text[:5])
        print(combined_text[:500])
        
        return placeholders, all_text
        
    except Exception as e:
        print(f"❌ Error reading template: {e}")
        return set(), []

def analyze_csv_contacts(csv_path):
    """Extract field names from CSV"""
    print(f"\n{'='*70}")
    print(f"ANALYZING CONTACTS: {csv_path}")
    print(f"{'='*70}")
    
    try:
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            fieldnames = reader.fieldnames
            
            print(f"\n📋 CSV contains {len(fieldnames)} fields:")
            for field in fieldnames:
                print(f"   - {field}")
            
            # Read first contact
            contacts = list(reader)
            if contacts:
                print(f"\n👤 Sample contact (first row):")
                for key, value in contacts[0].items():
                    print(f"   {key}: {value}")
            
            return fieldnames, contacts
            
    except Exception as e:
        print(f"❌ Error reading CSV: {e}")
        return [], []

def analyze_mapping(config_path):
    """Analyze the contact mapping configuration"""
    print(f"\n{'='*70}")
    print(f"ANALYZING CONFIG: {config_path}")
    print(f"{'='*70}")
    
    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
        
        mapping = config.get('contact_mapping', {})
        print(f"\n🔗 Contact mapping configuration:")
        for key, value in mapping.items():
            print(f"   {key} → {value}")
        
        return config, mapping
        
    except Exception as e:
        print(f"❌ Error reading config: {e}")
        return {}, {}

def find_mismatches(template_placeholders, csv_fields, mapping):
    """Identify mismatches between template and data"""
    print(f"\n{'='*70}")
    print("MISMATCH ANALYSIS")
    print(f"{'='*70}")
    
    issues = []
    
    # Check if template uses bracket notation
    print("\n🔍 Template placeholder format:")
    if template_placeholders:
        print("   ✓ Uses [Placeholder] bracket notation")
    else:
        print("   ⚠️  No placeholders found in template")
        issues.append("No placeholders found in template")
    
    # Check mapping
    print("\n🔍 Mapping analysis:")
    for mapped_key, csv_field in mapping.items():
        if csv_field not in csv_fields:
            print(f"   ❌ Mapping error: '{mapped_key}' maps to '{csv_field}' but CSV has no such field")
            issues.append(f"Missing CSV field: {csv_field}")
        else:
            print(f"   ✓ '{mapped_key}' → '{csv_field}' (exists in CSV)")
    
    # Check if placeholders match mapped keys
    print("\n🔍 Template-to-mapping match:")
    mapped_keys = set(mapping.keys())
    
    for placeholder in template_placeholders:
        # Normalize placeholder name
        normalized = placeholder.lower().replace(' ', '_').replace('/', '_')
        
        if placeholder not in mapped_keys and normalized not in mapped_keys:
            print(f"   ❌ Template has '[{placeholder}]' but no mapping exists")
            issues.append(f"Unmapped placeholder: [{placeholder}]")
        else:
            print(f"   ✓ '[{placeholder}]' can be mapped")
    
    return issues

def suggest_fixes(template_placeholders, csv_fields, mapping):
    """Suggest fixes for identified issues"""
    print(f"\n{'='*70}")
    print("SUGGESTED FIXES")
    print(f"{'='*70}")
    
    print("\n💡 The issue is likely:")
    print("   Your template uses [Bracket Placeholders] but the substitution")
    print("   code expects a different format like {name} or {{name}}")
    
    print("\n🔧 Recommended fixes:")
    print("\n1. UPDATE THE TEMPLATE to use proper Python format strings:")
    print("   Replace: [Recipient Name] → {name}")
    print("   Replace: [Title / Position] → {position}")
    print("   Replace: [Department / School] → {organization}")
    
    print("\n2. OR UPDATE THE CODE to handle [Bracket] notation:")
    print("   In docx_parser.py, find the substitution logic and ensure it replaces")
    print("   [placeholder] patterns with actual data")
    
    print("\n3. CHECK the substitution code looks like this:")
    print("   ```python")
    print("   for key, csv_field in contact_mapping.items():")
    print("       value = contact_data.get(csv_field, '')")
    print("       content = content.replace(f'[{key}]', value)")
    print("   ```")

if __name__ == "__main__":
    print("\n" + "="*70)
    print("CAMPAIGN TEMPLATE DIAGNOSTIC TOOL")
    print("="*70)
    
    # Paths
    config_path = "scheduled-campaigns/test_campaign.json"
    
    # 1. Analyze config
    config, mapping = analyze_mapping(config_path)
    
    if not config:
        print("\n❌ Cannot proceed without config file")
        sys.exit(1)
    
    # 2. Get template path from config
    templates = config.get('templates', [])
    if templates:
        template_path = templates[0]
        template_placeholders, _ = analyze_docx_template(template_path)
    else:
        print("\n❌ No templates found in config")
        template_placeholders = set()
    
    # 3. Get contacts path from config
    contacts_path = config.get('contacts', '')
    if contacts_path:
        csv_fields, contacts = analyze_csv_contacts(contacts_path)
    else:
        print("\n❌ No contacts file specified in config")
        csv_fields = []
    
    # 4. Find mismatches
    if template_placeholders and csv_fields:
        issues = find_mismatches(template_placeholders, csv_fields, mapping)
        
        if issues:
            print(f"\n⚠️  Found {len(issues)} issues:")
            for issue in issues:
                print(f"   - {issue}")
        else:
            print("\n✓ No obvious mismatches found")
            print("  The issue must be in the substitution code itself")
    
    # 5. Suggest fixes
    suggest_fixes(template_placeholders, csv_fields, mapping)
    
    print("\n" + "="*70)
    print("DIAGNOSTIC COMPLETE")
    print("="*70)
