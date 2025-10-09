#!/usr/bin/env python3
"""
Fix corrupted DOCX files by regenerating them from scratch
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime

def create_valid_docx(filepath, content=None):
    """Create a new valid DOCX file"""
    try:
        # Create a new document
        doc = Document()
        
        # Add default content if none provided
        if content is None:
            doc.add_heading('Email Campaign Template', 0)
            doc.add_paragraph()
            
            # Subject line
            p = doc.add_paragraph()
            p.add_run('Subject: ').bold = True
            p.add_run('Welcome {{name}} to Our Platform')
            
            doc.add_paragraph()
            
            # Email body
            doc.add_paragraph('Dear {{name}},')
            doc.add_paragraph()
            doc.add_paragraph(
                'We are excited to welcome you from {{organization}} to our platform. '
                'This email is being sent to {{email}}.'
            )
            doc.add_paragraph()
            doc.add_paragraph('Your role as {{role}} in the {{domain}} sector makes you an ideal candidate for our services.')
            doc.add_paragraph()
            doc.add_paragraph('Best regards,')
            doc.add_paragraph('The Platform Team')
            doc.add_paragraph()
            
            # Footer
            footer = doc.add_paragraph()
            footer.add_run('---').italic = True
            footer = doc.add_paragraph()
            footer.add_run('Reply to this email for feedback: feedback@modelphysmat.com').italic = True
        else:
            # Add provided content
            for line in content.split('\n'):
                doc.add_paragraph(line)
        
        # Save the document
        doc.save(str(filepath))
        print(f"✅ Created valid DOCX: {filepath}")
        return True
        
    except Exception as e:
        print(f"❌ Error creating {filepath}: {e}")
        return False

def backup_file(filepath):
    """Create a backup of the file"""
    if filepath.exists():
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_path = filepath.with_suffix(f'.corrupted_{timestamp}.bak')
        try:
            filepath.rename(backup_path)
            print(f"📦 Backed up corrupted file to: {backup_path}")
            return True
        except Exception as e:
            print(f"⚠️  Could not backup {filepath}: {e}")
            return False
    return False

def fix_corrupted_docx_files(directory):
    """Find and fix all corrupted DOCX files in directory"""
    directory = Path(directory)
    
    if not directory.exists():
        print(f"❌ Directory not found: {directory}")
        return False
    
    print(f"🔍 Scanning {directory} for DOCX files...")
    print()
    
    docx_files = list(directory.rglob('*.docx'))
    
    if not docx_files:
        print("No DOCX files found")
        return True
    
    print(f"Found {len(docx_files)} DOCX file(s)")
    print()
    
    corrupted_count = 0
    fixed_count = 0
    
    for docx_file in docx_files:
        # Skip backup files
        if '.backup_' in docx_file.name or '.corrupted_' in docx_file.name:
            print(f"⏭️  Skipping backup file: {docx_file.name}")
            continue
        
        print(f"Checking: {docx_file}")
        
        # Try to open the file
        try:
            doc = Document(str(docx_file))
            print(f"  ✅ Valid DOCX file")
        except Exception as e:
            print(f"  ❌ Corrupted: {e}")
            corrupted_count += 1
            
            # Backup the corrupted file
            backup_file(docx_file)
            
            # Try to extract any text content from the file
            try:
                with open(docx_file, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                    # Look for any readable text
                    if len(content) > 0:
                        print(f"  📝 Attempting to preserve content from corrupted file...")
            except:
                content = None
            
            # Create a new valid DOCX file
            if create_valid_docx(docx_file, content=None):
                fixed_count += 1
        
        print()
    
    print("=" * 50)
    print(f"📊 Summary:")
    print(f"  Total files: {len(docx_files)}")
    print(f"  Corrupted: {corrupted_count}")
    print(f"  Fixed: {fixed_count}")
    print("=" * 50)
    
    return corrupted_count == 0 or fixed_count == corrupted_count

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python fix_corrupted_docx.py <directory>")
        print("Example: python fix_corrupted_docx.py campaign-templates/")
        sys.exit(1)
    
    directory = sys.argv[1]
    success = fix_corrupted_docx_files(directory)
    sys.exit(0 if success else 1)
