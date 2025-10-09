#!/usr/bin/env python3
"""
Quick DOCX File Diagnostic Tool
Run this to diagnose issues with specific DOCX files
"""

import os
import sys
import zipfile
from pathlib import Path

def diagnose_docx(filepath):
    """Comprehensive DOCX file diagnosis"""
    print("=" * 70)
    print(f"🔍 DOCX FILE DIAGNOSTIC REPORT")
    print("=" * 70)
    print(f"File: {filepath}")
    print()
    
    filepath = Path(filepath)
    
    # Test 1: File Exists
    print("📋 Test 1: File Existence")
    if not filepath.exists():
        print("  ❌ FAILED: File does not exist")
        return False
    print("  ✅ PASSED: File exists")
    print()
    
    # Test 2: File Size
    print("📋 Test 2: File Size")
    try:
        size = filepath.stat().st_size
        size_kb = size / 1024
        size_mb = size_kb / 1024
        
        print(f"  Size: {size} bytes ({size_kb:.2f} KB, {size_mb:.2f} MB)")
        
        if size == 0:
            print("  ❌ FAILED: File is empty (0 bytes)")
            print("  💡 Fix: Regenerate or replace the file")
            return False
        elif size < 1024:
            print("  ⚠️  WARNING: File is very small (< 1 KB)")
            print("  💡 May be incomplete or corrupted")
        else:
            print("  ✅ PASSED: File has content")
    except Exception as e:
        print(f"  ❌ FAILED: Cannot read file size - {e}")
        return False
    print()
    
    # Test 3: File Permissions
    print("📋 Test 3: File Permissions")
    readable = os.access(filepath, os.R_OK)
    if readable:
        print("  ✅ PASSED: File is readable")
    else:
        print("  ❌ FAILED: File is not readable (permission denied)")
        print("  💡 Fix: Check file permissions with 'ls -la'")
        return False
    print()
    
    # Test 4: ZIP Structure
    print("📋 Test 4: ZIP Archive Validity")
    if not zipfile.is_zipfile(filepath):
        print("  ❌ FAILED: File is not a valid ZIP archive")
        print("  💡 DOCX files are ZIP archives containing XML files")
        print("  💡 Fix: File may be corrupted - try:")
        print("     - Opening in Microsoft Word or LibreOffice")
        print("     - Regenerating from original source")
        print("     - Restoring from backup")
        
        # Try to determine actual file type
        try:
            with open(filepath, 'rb') as f:
                header = f.read(8)
            print(f"\n  File header (hex): {header.hex()}")
            
            # Common file signatures
            signatures = {
                b'PK\x03\x04': 'ZIP archive (expected for DOCX)',
                b'\x50\x4b\x03\x04': 'ZIP archive (expected for DOCX)',
                b'{\rtf': 'RTF document',
                b'<html': 'HTML document',
                b'<?xml': 'XML document',
            }
            
            for sig, desc in signatures.items():
                if header.startswith(sig):
                    print(f"  Detected: {desc}")
                    break
            else:
                print(f"  Unknown file type")
        except:
            pass
        
        return False
    
    print("  ✅ PASSED: File is a valid ZIP archive")
    print()
    
    # Test 5: DOCX Structure
    print("📋 Test 5: DOCX Structure (Required Files)")
    required_files = {
        'word/document.xml': 'Main document content',
        '[Content_Types].xml': 'Content type definitions',
        'word/_rels/document.xml.rels': 'Document relationships (optional)',
        'docProps/core.xml': 'Core properties (optional)'
    }
    
    try:
        with zipfile.ZipFile(filepath, 'r') as zip_file:
            file_list = zip_file.namelist()
            
            print(f"  Total files in archive: {len(file_list)}")
            print()
            
            all_required_present = True
            for required_file, description in required_files.items():
                if required_file in file_list:
                    print(f"  ✅ {required_file}")
                    print(f"     └─ {description}")
                else:
                    if 'optional' in description.lower():
                        print(f"  ⚠️  {required_file} (missing, but optional)")
                    else:
                        print(f"  ❌ {required_file} (MISSING - REQUIRED)")
                        all_required_present = False
                    print(f"     └─ {description}")
            
            print()
            if not all_required_present:
                print("  ❌ FAILED: Missing required DOCX files")
                print("  💡 Fix: File structure is corrupted - regenerate file")
                return False
            
            print("  ✅ PASSED: All required files present")
    except zipfile.BadZipFile:
        print("  ❌ FAILED: ZIP archive is corrupted")
        return False
    except Exception as e:
        print(f"  ❌ FAILED: Error reading ZIP - {e}")
        return False
    print()
    
    # Test 6: python-docx Compatibility
    print("📋 Test 6: python-docx Library Compatibility")
    try:
        from docx import Document
        
        try:
            doc = Document(filepath)
            
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            # Extract text
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            total_text = '\n'.join(text_content)
            char_count = len(total_text)
            
            print(f"  ✅ PASSED: Successfully opened with python-docx")
            print(f"  Paragraphs: {para_count}")
            print(f"  Tables: {table_count}")
            print(f"  Text characters: {char_count}")
            
            if char_count == 0:
                print(f"  ⚠️  WARNING: Document contains no extractable text")
                print(f"  💡 May contain only images or special content")
            else:
                print(f"\n  📄 Content Preview (first 200 chars):")
                print(f"  {total_text[:200]}...")
            
        except Exception as e:
            print(f"  ❌ FAILED: python-docx cannot open file")
            print(f"  Error: {e}")
            print(f"  💡 Possible causes:")
            print(f"     - Password-protected document")
            print(f"     - Unsupported DOCX features")
            print(f"     - Corrupted XML content")
            return False
            
    except ImportError:
        print("  ⚠️  SKIPPED: python-docx not installed")
        print("  Install with: pip install python-docx")
    print()
    
    # Final Verdict
    print("=" * 70)
    print("🎉 DIAGNOSIS COMPLETE: File appears valid")
    print("=" * 70)
    print("✅ All tests passed - file should work with campaign system")
    print()
    
    return True


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python docx_diagnostic.py <path-to-docx-file>")
        print("\nExample:")
        print("  python docx_diagnostic.py campaign-templates/education/outreach.docx")
        sys.exit(1)
    
    filepath = sys.argv[1]
    
    result = diagnose_docx(filepath)
    
    if not result:
        print("\n❌ FILE FAILED VALIDATION")
        print("=" * 70)
        print("🔧 RECOMMENDED ACTIONS:")
        print("=" * 70)
        print("1. Try opening the file in Microsoft Word or LibreOffice")
        print("2. If it opens, save it again to repair any corruption")
        print("3. If it doesn't open, regenerate the file from scratch")
        print("4. Check if the file was correctly uploaded/transferred")
        print("5. Verify the file is not password-protected")
        print("\n💡 QUICK FIXES:")
        print("  - Empty file: Create new content and save as DOCX")
        print("  - Not a ZIP: File may have been renamed incorrectly")
        print("  - Missing XML: Recreate file in Word/LibreOffice")
        print("  - Permission denied: Run 'chmod 644 <filename>'")
        sys.exit(1)
    else:
        print("✅ File is ready for use in campaign system")
        sys.exit(0)
